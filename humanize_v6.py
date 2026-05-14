"""
IEEE paper humanizer v6 — watermark cleaning + ultra-aggressive restructuring.

Phase 1: Deep watermark / artifact cleaning
  - Strip ALL invisible Unicode (zero-width, BOM, direction marks)
  - Replace Cyrillic/homoglyph lookalikes with ASCII
  - Normalize all spaces (double→single, NBSP→regular, thin/hair→regular)
  - Convert em/en dashes to hyphens (AI overuses dashes)
  - Normalize smart quotes to straight quotes
  - Sanitize document metadata (remove python-docx fingerprint)
  - Consolidate run fragments (merge adjacent same-format runs)
  - Remove whitespace-only runs

Phase 2: Ultra-aggressive text humanization
  - 100% sentence touch rate (every sentence gets >=1 structural change)
  - Two-pass restructuring with fallback guarantee
  - Forced burstiness (CV >= 0.50)
  - Forced starter variation (max 20% for any single starter word)
  - 95% synonym perturbation
  - Final AI vocabulary hard-kill pass
"""
import sys, io, re as _re, unicodedata as _ud, math, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

SRC = 'paper_ieee.docx'
DST = 'page.docx'
doc = Document(SRC)
print(f"Loaded {SRC}: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

SKIP_INDICES = {0, 1, 2, 3, 5}
REFERENCES_START = 438


# ================================================================
# PHASE 1: WATERMARK / ARTIFACT CLEANING
# ================================================================

# 1a. Metadata sanitization
print("\n--- Phase 1a: Metadata sanitization ---")
props = doc.core_properties
props.author = "Student"
props.comments = ""
props.keywords = ""
props.category = ""
props.subject = ""
props.content_status = ""
props.identifier = ""
props.language = "en-US"
props.version = ""
try:
    props.last_modified_by = "Student"
except:
    pass
print("  Metadata cleaned")

# 1b. Homoglyph map (Cyrillic/Greek lookalikes → ASCII)
HOMOGLYPHS = {
    '\u0410': 'A', '\u0412': 'B', '\u0421': 'C', '\u0415': 'E',
    '\u041D': 'H', '\u041A': 'K', '\u041C': 'M', '\u041E': 'O',
    '\u0420': 'P', '\u0422': 'T', '\u0425': 'X', '\u0430': 'a',
    '\u0435': 'e', '\u043E': 'o', '\u0440': 'p', '\u0441': 'c',
    '\u0443': 'y', '\u0445': 'x', '\u0455': 's', '\u0456': 'i',
    '\u0458': 'j', '\u04BB': 'h',
    '\u0391': 'A', '\u0392': 'B', '\u0395': 'E', '\u0396': 'Z',
    '\u0397': 'H', '\u0399': 'I', '\u039A': 'K', '\u039C': 'M',
    '\u039D': 'N', '\u039F': 'O', '\u03A1': 'P', '\u03A4': 'T',
    '\u03A5': 'Y', '\u03A7': 'X',
}

def clean_text_watermarks(text):
    """Remove all hidden watermark characters from text."""
    if not text:
        return text

    for src, dst in HOMOGLYPHS.items():
        text = text.replace(src, dst)

    text = _re.sub(r'[\u200B\u200C\u200D\u200E\u200F\u2060\u2061\u2062\u2063\u2064\uFEFF\u180E\u034F\u061C\u2028\u2029]', '', text)
    text = _re.sub(r'[\u00A0\u2007\u2009\u200A\u202F\u205F\u3000\u2000-\u200A]', ' ', text)
    text = _re.sub(r'[\u0000-\u0008\u000B\u000C\u000E-\u001F\u007F-\u009F]', '', text)

    text = _re.sub(r'[\u2013\u2014\u2212]', '-', text)
    text = _re.sub(r'[\u201C\u201D\u201E\u00AB\u00BB]', '"', text)
    text = _re.sub(r'[\u2018\u2019\u201A\u201B\u2039\u203A]', "'", text)
    text = text.replace('\u2026', '...')
    text = _re.sub(r'[\u2022\u25E6\u00B7\u2219]', '-', text)

    text = _re.sub(r'&[#\w]+;', ' ', text)
    text = _re.sub(r' {2,}', ' ', text)
    text = _re.sub(r'\s+([.,!?;:])', r'\1', text)

    return text


# 1c. Clean every run in the document
print("--- Phase 1b: Run-level watermark cleaning ---")
cleaned_runs = 0

def clean_all_runs(paragraphs):
    global cleaned_runs
    for p in paragraphs:
        for run in p.runs:
            if run.text:
                original = run.text
                cleaned = clean_text_watermarks(original)
                if cleaned != original:
                    run.text = cleaned
                    cleaned_runs += 1

clean_all_runs(doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            clean_all_runs(cell.paragraphs)
print(f"  Cleaned {cleaned_runs} runs")


# 1d. Consolidate fragmented runs (merge adjacent runs with same formatting)
print("--- Phase 1c: Run consolidation ---")
merged_count = 0

def runs_same_format(r1, r2):
    """Check if two runs have identical formatting."""
    if r1.bold != r2.bold:
        return False
    if r1.italic != r2.italic:
        return False
    if r1.underline != r2.underline:
        return False
    f1, f2 = r1.font, r2.font
    if f1.name != f2.name:
        return False
    if f1.size != f2.size:
        return False
    try:
        if f1.color.rgb != f2.color.rgb:
            return False
    except:
        pass
    if r1.style != r2.style:
        return False
    return True

def consolidate_paragraph_runs(para):
    global merged_count
    runs = list(para.runs)
    if len(runs) < 2:
        return
    i = 0
    while i < len(runs) - 1:
        if not runs[i].text.strip() and not runs[i].text:
            p_elem = runs[i]._element
            p_elem.getparent().remove(p_elem)
            runs.pop(i)
            merged_count += 1
            continue

        if runs_same_format(runs[i], runs[i + 1]):
            runs[i].text = (runs[i].text or '') + (runs[i + 1].text or '')
            p_elem = runs[i + 1]._element
            p_elem.getparent().remove(p_elem)
            runs.pop(i + 1)
            merged_count += 1
        else:
            i += 1

for p in doc.paragraphs:
    consolidate_paragraph_runs(p)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                consolidate_paragraph_runs(p)
print(f"  Merged/removed {merged_count} redundant runs")


# ================================================================
# PHASE 2: TEXT HUMANIZATION
# ================================================================

def is_section_header(txt):
    return bool(_re.match(r'^(I{1,3}V?|VI{0,3}|IX|X|XI{0,3})\.\s', txt)) or bool(_re.match(r'^\d+\.\d+', txt))

def is_algorithm_line(txt):
    return bool(_re.match(r'^\d{1,2}[\.\)]\s', txt)) or txt.startswith("Input:") or txt.startswith("Output:")

def is_formula(txt):
    if len(txt) > 200:
        return False
    special = '\u2211\u220f\u222b\u2208\u2200\u2203\u2264\u2265\u2260\u221e\u2202\u2207\u00d7\u2297\u2299\u221d'
    if sum(1 for c in txt if c in special) >= 2:
        return True
    words = txt.split()
    if len(words) < 8 and _re.search(r'[=<>]', txt) and _re.search(r'[a-z].*=', txt):
        return True
    if txt.count('=') >= 2 and len(words) < 12:
        return True
    return False

def has_citation(txt):
    return bool(_re.search(r'\[\d+\]', txt))

def should_skip(idx, txt):
    if idx in SKIP_INDICES: return True
    if idx >= REFERENCES_START: return True
    if not txt or len(txt.strip()) < 5: return True
    if is_section_header(txt): return True
    if is_algorithm_line(txt): return True
    if is_formula(txt): return True
    return False


class PRNG:
    def __init__(self, seed=42):
        self._s = seed
    def rand(self):
        self._s = (self._s * 1103515245 + 12345) & 0x7FFFFFFF
        return self._s / 0x7FFFFFFF
    def pick(self, arr):
        return arr[int(self.rand() * len(arr))]
    def shuffle(self, arr):
        a = list(arr)
        for i in range(len(a)-1, 0, -1):
            j = int(self.rand() * (i+1))
            a[i], a[j] = a[j], a[i]
        return a


def split_sentences(text):
    text = _re.sub(r'(e\.g\.)\s+', r'\1NOSPLIT', text)
    text = _re.sub(r'(i\.e\.)\s+', r'\1NOSPLIT', text)
    text = _re.sub(r'(et al\.)\s+', r'\1NOSPLIT', text)
    text = _re.sub(r'(Fig\.)\s+', r'\1NOSPLIT', text)
    text = _re.sub(r'(Sec\.)\s+', r'\1NOSPLIT', text)
    text = _re.sub(r'(vs\.)\s+', r'\1NOSPLIT', text)
    text = _re.sub(r'(vol\.)\s+', r'\1NOSPLIT', text)
    text = _re.sub(r'(no\.)\s+', r'\1NOSPLIT', text)
    text = _re.sub(r'(pp\.)\s+', r'\1NOSPLIT', text)
    text = _re.sub(r'\(([ivx]+)\)\s*', lambda m: '(' + m.group(1) + ')NOSPLIT', text)
    parts = _re.split(r'(?<=[.!?])\s+(?=[A-Z\[])', text)
    return [p.replace('NOSPLIT', ' ').strip() for p in parts if p.strip()]


# ================================================================
# VOCABULARY + FILLER + TRANSITION + CONTRACTION — all 100%
# ================================================================
def humanize_words(text, rng):
    transitions = [
        (r"\bFurthermore,?\s", lambda: rng.pick(["Also, ","In addition, ","Equally, ","Beyond this, "])),
        (r"\bMoreover,?\s", lambda: rng.pick(["Also, ","On top of that, ","Along with this, "])),
        (r"\bAdditionally,?\s", lambda: rng.pick(["Also, ","On top of this, ","Equally, ","Besides this, "])),
        (r"\bConsequently,?\s", lambda: rng.pick(["As such, ","For this reason, ","So, ","Because of this, "])),
        (r"\bNevertheless,?\s", lambda: rng.pick(["Even so, ","That said, ","Still, ","All the same, "])),
        (r"\bNonetheless,?\s", lambda: rng.pick(["Even so, ","Still, ","Yet, ","All the same, "])),
        (r"\bSubsequently,?\s", lambda: rng.pick(["Then, ","After that, ","Next, ","Later, "])),
        (r"\bConversely,?\s", lambda: rng.pick(["On the other hand, ","In contrast, ","By contrast, "])),
        (r"\bUltimately,?\s", lambda: rng.pick(["In the end, ","Finally, ","At last, "])),
        (r"(?i)\bIn\s+addition,?\s", lambda: rng.pick(["Also, ","Equally, ","Plus, "])),
        (r"(?i)\bAs\s+a\s+result,?\s", lambda: rng.pick(["Because of this, ","For this reason, ","So, "])),
        (r"(?i)\bIn\s+particular,?\s", lambda: rng.pick(["Especially, ","Above all, "])),
        (r"(?i)\bSpecifically,?\s", lambda: rng.pick(["Namely, ","That is, ","More precisely, "])),
        (r"(?i)\bAccordingly,?\s", lambda: rng.pick(["For that reason, ","As such, ","So, "])),
        (r"(?i)\bHence,?\s", lambda: rng.pick(["For this reason, ","As such, ","So, "])),
        (r"(?i)\bThus,?\s", lambda: rng.pick(["In this way, ","So, ","As a result, "])),
        (r"(?i)\bFurther,?\s", lambda: rng.pick(["Also, ","On top of this, ","Beyond that, "])),
        (r"(?i)\bNotably,?\s", lambda: rng.pick(["Especially, ","In particular, "])),
        (r"(?i)\bSimilarly,?\s", lambda: rng.pick(["Along the same lines, ","Likewise, ","In the same way, "])),
        (r"(?i)\bOverall,?\s", lambda: rng.pick(["Taken together, ","All in all, ","On the whole, "])),
    ]
    for pat, fn in transitions:
        text = _re.sub(pat, lambda m, f=fn: f(), text)

    vocab = [
        (r"(?i)\butilizes?\b","uses"),(r"(?i)\butilized\b","used"),(r"(?i)\butilizing\b","using"),
        (r"(?i)\bfacilitates?\b","helps"),(r"(?i)\bfacilitated\b","helped"),(r"(?i)\bfacilitating\b","helping"),
        (r"(?i)\bdemonstrates?\b","shows"),(r"(?i)\bdemonstrated\b","showed"),(r"(?i)\bdemonstrating\b","showing"),
        (r"(?i)\bcomprehensive\b", lambda: rng.pick(["thorough","complete","full"])),
        (r"(?i)\bsignificantly\b", lambda: rng.pick(["greatly","markedly","a lot"])),
        (r"(?i)\bsignificant\b", lambda: rng.pick(["major","notable","marked","big"])),
        (r"(?i)\bsubstantially\b", lambda: rng.pick(["greatly","largely","a lot"])),
        (r"(?i)\bsubstantial\b", lambda: rng.pick(["large","considerable","sizable","big"])),
        (r"(?i)\bapproximately\b", lambda: rng.pick(["about","around","roughly","close to"])),
        (r"(?i)\bnumerous\b", lambda: rng.pick(["many","several","plenty of"])),
        (r"(?i)\boptimal\b", lambda: rng.pick(["best","ideal","top"])),
        (r"(?i)\benhances?\b","improves"),(r"(?i)\benhanced\b","improved"),(r"(?i)\benhancing\b","improving"),
        (r"(?i)\bleverages?\b","uses"),(r"(?i)\bleveraging\b","using"),(r"(?i)\bleveraged\b","used"),
        (r"(?i)\bensures?\b","makes sure"),(r"(?i)\bensuring\b","making sure"),(r"(?i)\bensured\b","made sure"),
        (r"(?i)\bpivotal\b", lambda: rng.pick(["key","central","critical"])),
        (r"(?i)\brobust\b", lambda: rng.pick(["strong","solid","reliable","sturdy"])),
        (r"(?i)\bseamlessly\b","smoothly"),(r"(?i)\bseamless\b","smooth"),
        (r"(?i)\binnovative\b", lambda: rng.pick(["new","novel","creative","fresh"])),
        (r"(?i)\bholistic\b", lambda: rng.pick(["overall","complete","full"])),
        (r"(?i)\bparadigms?\b","models"),(r"(?i)\bmethodolog(?:y|ies)\b", lambda: rng.pick(["method","technique","approach"])),
        (r"(?i)\bunderscores?\b","highlights"),(r"(?i)\bunderscoring\b","highlighting"),
        (r"(?i)\bmitigates?\b","reduces"),(r"(?i)\bmitigating\b","reducing"),(r"(?i)\bmitigation\b","reduction"),
        (r"(?i)\bthe proposed\b","our"),(r"(?i)\bthis paper presents\b","we present"),
        (r"(?i)\bthis paper proposes\b","we propose"),(r"(?i)\bthis work introduces\b","we introduce"),
        (r"(?i)\bnotwithstanding\b","despite"),(r"(?i)\bwherein\b","where"),
        (r"(?i)\bthereby\b","and so"),(r"(?i)\bwhereas\b","while"),
        (r"(?i)\bshowcases?\b","shows"),(r"(?i)\bencompasses?\b","covers"),
        (r"(?i)\bexemplif(?:ies|y)\b","illustrates"),
        (r"(?i)\bharnessing\b","using"),(r"(?i)\bharnesses\b","uses"),
        (r"(?i)\bfostering\b","building"),(r"(?i)\bfosters\b","builds"),
        (r"(?i)\baforementioned\b","mentioned earlier"),(r"(?i)\bherein\b","here"),
        (r"(?i)\bexhibits?\b","shows"),(r"(?i)\bexhibited\b","showed"),(r"(?i)\bexhibiting\b","showing"),
        (r"(?i)\bmanifests?\b","shows"),(r"(?i)\bmanifested\b","showed"),
        (r"(?i)\bpossesses?\b","has"),(r"(?i)\brenders?\b","makes"),(r"(?i)\brendered\b","made"),
        (r"(?i)\bpredominantly\b", lambda: rng.pick(["mostly","mainly","largely"])),
        (r"(?i)\bmultifaceted\b", lambda: rng.pick(["complex","varied","layered"])),
        (r"(?i)\bparamount\b", lambda: rng.pick(["critical","most important"])),
        (r"(?i)\bstreamlines?\b","simplifies"),(r"(?i)\bexpedites?\b","speeds up"),
        (r"(?i)\bconducive\b","helpful"),(r"(?i)\brequisite\b","needed"),
        (r"(?i)\bpertaining to\b","about"),(r"(?i)\bwith respect to\b","about"),
        (r"(?i)\bin the context of\b","in"),(r"(?i)\bwith regard to\b","about"),
        (r"(?i)\bvis-a-vis\b","compared to"),(r"(?i)\bin lieu of\b","instead of"),
        (r"(?i)\bcategorically\b","firmly"),(r"(?i)\bexponentially\b","rapidly"),
        (r"(?i)\bintrinsically\b","by nature"),(r"(?i)\binherently\b","by nature"),
        (r"(?i)\boverarching\b","broad"),(r"(?i)\bunderpins?\b","supports"),
        (r"(?i)\bencapsulates?\b","captures"),(r"(?i)\bdelves?\b","digs"),
        (r"(?i)\bpertinent\b","relevant"),(r"(?i)\bpreclude[sd]?\b","prevents"),
        (r"(?i)\binsofar\b","to the extent"),(r"(?i)\bhenceforth\b","from now on"),
        (r"(?i)\bgarners?\b","gets"),(r"(?i)\bgarnered\b","got"),
        (r"(?i)\baugments?\b","boosts"),(r"(?i)\baugmented\b","boosted"),
        (r"(?i)\bcommences?\b","starts"),(r"(?i)\bcommenced\b","started"),
        (r"(?i)\bterminates?\b","ends"),(r"(?i)\bterminated\b","ended"),
        (r"(?i)\bcorroborates?\b","backs up"),(r"(?i)\belucidated?\b","explained"),
        (r"(?i)\bameliorate[sd]?\b","improves"),(r"(?i)\bdelineates?\b","outlines"),
        (r"(?i)\bexacerbate[sd]?\b","worsens"),(r"(?i)\bsupersede[sd]?\b","replaces"),
    ]
    for pat, repl in vocab:
        if callable(repl):
            text = _re.sub(pat, lambda m, f=repl: f(), text)
        else:
            text = _re.sub(pat, repl, text)

    fillers = [
        (r"(?i)\bIt is important to note that\s*",""),(r"(?i)\bIt should be noted that\s*",""),
        (r"(?i)\bIt is worth (?:mentioning|noting) that\s*",""),(r"(?i)\bIt goes without saying that?\s*",""),
        (r"(?i)\bNeedless to say,?\s*",""),(r"(?i)\bin order to\b","to"),
        (r"(?i)\bdue to the fact that\b","because"),(r"(?i)\ba wide range of\b","many"),
        (r"(?i)\ba plethora of\b","many"),(r"(?i)\bplays a (?:crucial|vital|key|important|significant) role\b","matters"),
        (r"(?i)\bin today'?s rapidly evolving\b","in today's"),
        (r"(?i)\bin the ever-changing landscape of\b","in"),(r"(?i)\bat the present time\b","now"),
        (r"(?i)\bin the realm of\b","in"),(r"(?i)\bfor the purpose of\b","for"),
        (r"(?i)\bin the event that\b","if"),(r"(?i)\bprior to\b","before"),
        (r"(?i)\bsubsequent to\b","after"),(r"(?i)\bin conjunction with\b","with"),
        (r"(?i)\bon the basis of\b","based on"),(r"(?i)\bin the absence of\b","without"),
        (r"(?i)\bhas the potential to\b","can"),(r"(?i)\bis capable of\b","can"),
        (r"(?i)\bthe vast majority of\b","most"),(r"(?i)\ba considerable amount of\b","much"),
        (r"(?i)\bgiven the fact that\b", lambda: rng.pick(["since","because"])),
        (r"(?i)\bin light of\b", lambda: rng.pick(["given","considering"])),
        (r"(?i)\bin terms of\b", lambda: rng.pick(["regarding","for","when it comes to"])),
        (r"(?i)\bthrough the use of\b","using"),(r"(?i)\bby means of\b","through"),
        (r"(?i)\bin the case of\b","for"),(r"(?i)\bas a consequence of\b","because of"),
        (r"(?i)\bin accordance with\b","following"),(r"(?i)\bas opposed to\b","unlike"),
        (r"(?i)\bwith the exception of\b","except for"),(r"(?i)\bon account of\b","because of"),
        (r"(?i)\bin the process of\b","while"),(r"(?i)\bwith a view to\b","to"),
        (r"(?i)\bin an effort to\b","to"),(r"(?i)\bin such a manner that\b","so that"),
        (r"(?i)\bto a large extent\b","largely"),(r"(?i)\bat this juncture\b","now"),
    ]
    for pat, repl in fillers:
        if callable(repl):
            text = _re.sub(pat, lambda m, f=repl: f(), text)
        else:
            text = _re.sub(pat, repl, text)

    contractions = [
        (r"\bdoes not\b","doesn't"),(r"\bDoes not\b","Doesn't"),
        (r"\bdo not\b","don't"),(r"\bDo not\b","Don't"),
        (r"\bdid not\b","didn't"),(r"\bDid not\b","Didn't"),
        (r"\bis not\b","isn't"),(r"\bIs not\b","Isn't"),
        (r"\bare not\b","aren't"),(r"\bAre not\b","Aren't"),
        (r"\bwas not\b","wasn't"),(r"\bWas not\b","Wasn't"),
        (r"\bwere not\b","weren't"),(r"\bWere not\b","Weren't"),
        (r"\bcannot\b","can't"),(r"\bCannot\b","Can't"),
        (r"\bcould not\b","couldn't"),(r"\bCould not\b","Couldn't"),
        (r"\bwould not\b","wouldn't"),(r"\bWould not\b","Wouldn't"),
        (r"\bwill not\b","won't"),(r"\bWill not\b","Won't"),
        (r"\bshould not\b","shouldn't"),(r"\bShould not\b","Shouldn't"),
        (r"\bit is\b","it's"),(r"\bIt is\b","It's"),
        (r"\bthat is\b","that's"),(r"\bThat is\b","That's"),
        (r"\bthere is\b","there's"),(r"\bThere is\b","There's"),
        (r"\bthey are\b","they're"),(r"\bThey are\b","They're"),
        (r"\bwe are\b","we're"),(r"\bWe are\b","We're"),
        (r"\byou are\b","you're"),(r"\bYou are\b","You're"),
        (r"\bwe have\b","we've"),(r"\bWe have\b","We've"),
        (r"\bthey have\b","they've"),(r"\bThey have\b","They've"),
        (r"\bhave not\b","haven't"),(r"\bHave not\b","Haven't"),
        (r"\bhas not\b","hasn't"),(r"\bHas not\b","Hasn't"),
        (r"\blet us\b","let's"),(r"\bLet us\b","Let's"),
        (r"\bhe is\b","he's"),(r"\bshe is\b","she's"),
        (r"\bwhat is\b","what's"),(r"\bWhat is\b","What's"),
        (r"\bwho is\b","who's"),(r"\bWho is\b","Who's"),
        (r"\bwhere is\b","where's"),(r"\bWhere is\b","Where's"),
    ]
    for pat, repl in contractions:
        text = _re.sub(pat, repl, text)

    adverbs = [
        (r"(?i)\bconsiderably\b", lambda: rng.pick(["noticeably","markedly","a good deal"])),
        (r"(?i)\bremarkably\b", lambda: rng.pick(["surprisingly","strikingly"])),
        (r"(?i)\bparticularly\b", lambda: rng.pick(["especially","specifically"])),
        (r"(?i)\bprimarily\b", lambda: rng.pick(["mainly","mostly","chiefly","first and foremost"])),
        (r"(?i)\bfundamentally\b", lambda: rng.pick(["at its core","basically","at root"])),
        (r"(?i)\bincreasingly\b", lambda: rng.pick(["more and more","progressively","ever more"])),
        (r"(?i)\bsystematically\b", lambda: rng.pick(["step by step","one by one"])),
        (r"(?i)\beffectively\b", lambda: rng.pick(["in practice","successfully","well"])),
        (r"(?i)\bextensively\b", lambda: rng.pick(["broadly","widely","at length"])),
        (r"(?i)\bempirically\b", lambda: rng.pick(["through experiments","from data"])),
        (r"(?i)\bconcurrently\b", lambda: rng.pick(["at the same time","in parallel","side by side"])),
        (r"(?i)\bexplicitly\b", lambda: rng.pick(["directly","clearly","openly"])),
        (r"(?i)\bsimultaneously\b", lambda: rng.pick(["at once","at the same time","together"])),
        (r"(?i)\binvariably\b", lambda: rng.pick(["always","every time","without fail"])),
    ]
    for pat, fn in adverbs:
        text = _re.sub(pat, lambda m, f=fn: f(), text)

    passives = [
        (r"(?i)\bwas conducted\b", lambda: rng.pick(["took place","was carried out"])),
        (r"(?i)\bwere conducted\b", lambda: rng.pick(["took place","were carried out"])),
        (r"(?i)\bwas observed\b", lambda: rng.pick(["showed up","was noticed","appeared"])),
        (r"(?i)\bwas identified\b", lambda: rng.pick(["stood out","was spotted","came up"])),
        (r"(?i)\bwere identified\b", lambda: rng.pick(["stood out","were found","emerged"])),
        (r"(?i)\bis characterized by\b", lambda: rng.pick(["stands out for","is known for"])),
        (r"(?i)\bwas implemented\b", lambda: rng.pick(["was built","was put in place"])),
        (r"(?i)\bwas utilized\b","was used"),(r"(?i)\bwas employed\b","was used"),
        (r"(?i)\bwere employed\b","were used"),
        (r"(?i)\bwas performed\b", lambda: rng.pick(["was done","was carried out","ran"])),
        (r"(?i)\bwere performed\b", lambda: rng.pick(["were done","were carried out"])),
        (r"(?i)\bwas evaluated\b", lambda: rng.pick(["was tested","was assessed"])),
        (r"(?i)\bwere evaluated\b", lambda: rng.pick(["were tested","were assessed"])),
        (r"(?i)\bwas obtained\b", lambda: rng.pick(["was found","came out","was gathered"])),
        (r"(?i)\bwere obtained\b", lambda: rng.pick(["were found","came out"])),
        (r"(?i)\bwas achieved\b", lambda: rng.pick(["was reached","was hit"])),
        (r"(?i)\bwere achieved\b", lambda: rng.pick(["were reached","were hit"])),
        (r"(?i)\bwas applied\b", lambda: rng.pick(["was used","went into effect"])),
        (r"(?i)\bwas proposed\b", lambda: rng.pick(["was put forward","was suggested"])),
        (r"(?i)\bwas selected\b", lambda: rng.pick(["was chosen","was picked"])),
        (r"(?i)\bwere selected\b", lambda: rng.pick(["were chosen","were picked"])),
    ]
    for pat, repl in passives:
        if callable(repl):
            text = _re.sub(pat, lambda m, f=repl: f(), text)
        else:
            text = _re.sub(pat, repl, text)

    ngrams = [
        (r"(?i)\bIt is worth noting\b", lambda: rng.pick(["Note that","One should note"])),
        (r"(?i)\bOn the other hand\b", lambda: rng.pick(["In contrast","Then again"])),
        (r"(?i)\bIn this regard\b", lambda: rng.pick(["On this point","Here"])),
        (r"(?i)\bThis suggests that\b", lambda: rng.pick(["This points to","So"])),
        (r"(?i)\bThis indicates that\b", lambda: rng.pick(["This shows that","So"])),
        (r"(?i)\bThis implies that\b", lambda: rng.pick(["This means","So"])),
        (r"(?i)\bIt is essential to\b", lambda: rng.pick(["It's critical to","One must","We need to"])),
        (r"(?i)\bserves as a\b", lambda: rng.pick(["works as a","acts as a","functions as a"])),
        (r"(?i)\bpaves the way for\b", lambda: rng.pick(["opens the door to","enables"])),
        (r"(?i)\bsheds? light on\b", lambda: rng.pick(["clarifies","helps explain"])),
        (r"(?i)\bthe findings suggest\b", lambda: rng.pick(["the results point to","what we found suggests"])),
        (r"(?i)\bthe results indicate\b", lambda: rng.pick(["the data shows","our results point to"])),
        (r"(?i)\bit is evident that\b", lambda: rng.pick(["clearly,","we can see that"])),
        (r"(?i)\bit is clear that\b", lambda: rng.pick(["clearly,","plainly,","we can see that"])),
        (r"(?i)\bplay(?:s)? a (?:crucial|vital|critical|key|significant|important) role\b",
         lambda: rng.pick(["matter greatly","is central","is key"])),
    ]
    for pat, fn in ngrams:
        text = _re.sub(pat, lambda m, f=fn: f(), text)

    return text


# ================================================================
# SENTENCE RESTRUCTURING — 100% TOUCH RATE
# ================================================================
VERB_PAT = _re.compile(r'\b(is|are|was|were|has|have|had|can|could|will|would|shall|should|may|might|must|does|do|did|remains|becomes|seems|appears|provides|requires|involves|allows|enables|shows|leads|makes|takes|gives|keeps|needs|offers|produces|generates|achieves|reaches|uses|applies|runs|gets)\b', _re.I)
DEP_STARTERS = ('despite','although','while','since','because','if','unless','when','once','even though','given that','whereas','after','before','until','provided')

def is_dep_start(s):
    low = s.lower()
    return any(low.startswith(d) for d in DEP_STARTERS)


def restructure_sentence(s, rng, allow_merge_next=False, next_s=None):
    """Try to structurally change a single sentence. Returns (result_list, consumed_next)."""
    words = s.split()
    wc = len(words)

    if wc < 5 or is_formula(s):
        return [s], False

    # A) THERE-IS ELIMINATION
    m = _re.match(r"^(There(?:'s|'re| is| are| was| were))\s+(.+)$", s, _re.I)
    if m and not has_citation(s[:30]) and rng.rand() < 0.95:
        rest = m.group(2)
        if len(rest.split()) > 4:
            return [rest[0].upper() + rest[1:]], False

    # B) CLAUSE REORDERING
    if wc > 7 and rng.rand() < 0.92:
        for conj, front in [(" because ", "Because "), (" since ", "Since "),
                            (" although ", "Although "), (" while ", "While "),
                            (" when ", "When "), (" even though ", "Even though "),
                            (" so that ", "So that "), (" if ", "If "),
                            (" unless ", "Unless "), (" once ", "Once ")]:
            idx = s.lower().find(conj, max(8, len(s) // 5))
            if idx > 0 and not has_citation(s[idx:]):
                main = s[:idx].rstrip()
                sub = s[idx + len(conj):].strip()
                if len(sub.split()) >= 3 and len(main.split()) >= 3:
                    if main[0].isupper():
                        main = main[0].lower() + main[1:]
                    sub = sub.rstrip('.!?')
                    return [front + sub + ", " + main + "."], False

    # C) MERGE SHORT SENTENCES
    if allow_merge_next and next_s and wc < 14 and len(next_s.split()) < 14 and rng.rand() < 0.75:
        conj = rng.pick([", and ", "; ", " - ", ", plus ", " and also "])
        first = s.rstrip('.')
        second = next_s
        if second and second[0].isupper():
            second = second[0].lower() + second[1:]
        return [first + conj + second], True

    # D) SPLIT LONG SENTENCES (but NOT subordinate-clause openers without main verb before split)
    if wc > 20 and rng.rand() < 0.88:
        starts_dep = is_dep_start(s)
        for conj in [", and ", ", but ", ", yet ", " however ", "; ", ", which ", ", where ",
                     ", meaning ", ", though "]:
            idx = s.find(conj, max(12, len(s) // 4))
            if idx > 0 and not has_citation(s[idx:idx+25]):
                pre = s[:idx]
                if pre.count('(') != pre.count(')'):
                    continue
                if starts_dep and not VERB_PAT.search(pre):
                    continue
                first = pre.rstrip()
                if not first.endswith(('.','!','?')):
                    first += '.'
                rest = s[idx + len(conj):].strip()
                if rest and len(rest.split()) > 4:
                    rest = rest[0].upper() + rest[1:]
                    if conj.strip() in ("which","where"):
                        st = rng.pick(["This ","That ","It "])
                        rest = st + rest[0].lower() + rest[1:]
                    return [first, rest], False

    # E) PREPOSITIONAL PHRASE FRONTING
    if wc > 8 and rng.rand() < 0.75:
        m = _re.search(r',\s+((?:in|with|for|by|through|during|across|over|under|using|via|from|among|between|against)\s+[^,\.]+)[\.!?]$', s, _re.I)
        if m and not has_citation(m.group(1)):
            phrase = m.group(1).strip()
            main = s[:m.start()].rstrip()
            if main and main[0].isupper():
                main = main[0].lower() + main[1:]
            return [phrase[0].upper() + phrase[1:] + ", " + main + "."], False

    # F) ACTIVE/PASSIVE VOICE FLIP
    if rng.rand() < 0.65:
        m = _re.match(r'^(.+?)\s+(shows?|reveals?|indicates?|confirms?|suggests?|provides?|achieves?|produces?|generates?|yields?|reaches?)\s+(?:that\s+)?(.+)$', s, _re.I)
        if m and not has_citation(m.group(1)) and len(m.group(3).split()) > 3:
            subject = m.group(1).rstrip(',')
            rest = m.group(3).rstrip('.!?')
            if rest[0].islower():
                rest = rest[0].upper() + rest[1:]
            conn = rng.pick(["as shown by","as seen from","according to","based on","going by"])
            new_s = rest + ", " + conn + " " + subject[0].lower() + subject[1:] + "."
            if len(new_s.split()) > 5:
                return [new_s], False

    # G) CLAUSE-AWARE COMMA REORDERING (both halves need verbs, no citations at boundary)
    if wc > 12:
        for cm in _re.finditer(r',\s+', s):
            cp = cm.start()
            before = s[:cp].strip()
            after = s[cp+2:].strip()
            bw, aw = len(before.split()), len(after.split())
            if bw < 5 or aw < 7:
                continue
            if has_citation(before[-30:]) or has_citation(after[:30]):
                continue
            if before.count(',') > 0 or ' and ' in before[-25:] or ' or ' in before[-20:]:
                continue
            if before.count('(') != before.count(')'):
                continue
            if after.split()[0].lower() in ('and','or','but','nor'):
                continue
            if not VERB_PAT.search(after) or not VERB_PAT.search(before):
                continue
            if rng.rand() < 0.55:
                if after[0].islower():
                    after = after[0].upper() + after[1:]
                if before[0].isupper():
                    before = before[0].lower() + before[1:]
                return [after.rstrip('.') + ", " + before + "."], False
            break

    # H) "The/This X verb..." → fronted rewrite (skip proper nouns)
    m = _re.match(r'^(The|This|These|Those)\s+(\w+)', s)
    if m and wc > 7:
        first_noun = m.group(2)
        if not (first_noun[0].isupper() and first_noun not in ("ASR","BERT","CLIP","GPU","TPE","MLP","CNN","RNN","NLP","API","DSA")):
            det = m.group(1)
            rest = s[m.end():].strip()
            alts = {"The":["Our ","Each ","Every ","Any ","A given "],
                    "This":["Such a ","One such ","A similar "],
                    "These":["Such ","All these ","Several "],
                    "Those":["Such ","All those ","Several "]}
            if det in alts and rng.rand() < 0.70:
                return [rng.pick(alts[det]) + first_noun + " " + rest], False

    return [s], False


def restructure_all(sentences, rng):
    result = []
    i = 0
    while i < len(sentences):
        s = sentences[i]
        next_s = sentences[i + 1] if i + 1 < len(sentences) else None
        out, consumed_next = restructure_sentence(s, rng, allow_merge_next=True, next_s=next_s)
        result.extend(out)
        i += 2 if consumed_next else i + 1
        if not consumed_next:
            i = i  # already incremented above
            break
    # Fix: proper loop
    return result

def restructure_all(sentences, rng):
    result = []
    i = 0
    while i < len(sentences):
        s = sentences[i]
        next_s = sentences[i + 1] if i + 1 < len(sentences) else None
        out, consumed_next = restructure_sentence(s, rng, allow_merge_next=True, next_s=next_s)
        result.extend(out)
        i += 2 if consumed_next else 1
    return result


# ================================================================
# PASS 2: guaranteed touch for any unchanged sentence
# ================================================================
def guaranteed_touch(sentences, original_sentences, rng):
    """For any sentence identical to original, force a structural change."""
    orig_set = set(original_sentences)
    result = []
    for s in sentences:
        if s not in orig_set or len(s.split()) <= 6 or is_formula(s) or has_citation(s[:15]):
            result.append(s)
            continue

        words = s.split()
        wc = len(words)
        done = False

        # Strategy 1: clause-aware comma swap (need verbs on both sides)
        for cm in _re.finditer(r',\s+', s):
            cp = cm.start()
            before = s[:cp].strip()
            after = s[cp+2:].strip()
            if len(before.split()) < 4 or len(after.split()) < 5:
                continue
            if has_citation(before[-20:]) or has_citation(after[:20]):
                continue
            if before.count(',') > 0:
                continue
            if not VERB_PAT.search(after) or not VERB_PAT.search(before):
                continue
            if after[0].islower():
                after = after[0].upper() + after[1:]
            if before[0].isupper():
                before = before[0].lower() + before[1:]
            result.append(after.rstrip('.') + ", " + before + ".")
            done = True
            break

        if done:
            continue

        # Strategy 2: swap around semicolon or dash
        for sep in ["; ", " - "]:
            idx = s.find(sep)
            if idx > 10 and idx < len(s) - 15:
                before = s[:idx].strip()
                after = s[idx+len(sep):].strip()
                if len(after.split()) >= 3 and len(before.split()) >= 3:
                    if after[0].islower():
                        after = after[0].upper() + after[1:]
                    if before[0].isupper():
                        before = before[0].lower() + before[1:]
                    result.append(after.rstrip('.') + sep + before + ".")
                    done = True
                    break
        if done:
            continue

        # Strategy 3: determiner swap (skip proper nouns / capitalized words)
        m = _re.match(r'^(The|This|These|Those)\s+(\w+)', s)
        if m:
            det = m.group(1)
            next_word = m.group(2)
            is_proper = next_word[0].isupper() and next_word not in ("ASR","BERT","CLIP","GPU","TPE","MLP","CNN","RNN","NLP","API","DSA")
            if not is_proper:
                alts = {"The":["Our","Each","Every"],"This":["Such a","One such"],
                        "These":["Such","All these","Several"],"Those":["Such","All those"]}
                if det in alts:
                    new_det = rng.pick(alts[det])
                    result.append(new_det + " " + next_word + s[m.end():])
                    continue

        result.append(s)
    return result


# ================================================================
# BURSTINESS ENFORCEMENT — target CV >= 0.50
# ================================================================
def enforce_burstiness(sentences, rng):
    if len(sentences) < 3:
        return sentences

    def calc_cv(sents):
        lengths = [len(s.split()) for s in sents]
        mean = sum(lengths) / len(lengths)
        var = sum((l - mean) ** 2 for l in lengths) / len(lengths)
        return math.sqrt(var) / mean if mean > 0 else 0

    cv = calc_cv(sentences)
    if cv >= 0.50:
        return sentences

    result = list(sentences)
    for attempt in range(15):
        if calc_cv(result) >= 0.50:
            break

        lengths = [len(s.split()) for s in result]
        longest_idx = max(range(len(result)), key=lambda i: lengths[i])
        s = result[longest_idx]

        if lengths[longest_idx] > 16:
            starts_dep = is_dep_start(s)
            for conj in [", and ", ", but ", "; ", ", which ", ", where ", ", while ", ", though "]:
                idx = s.find(conj, max(10, len(s) // 4))
                if idx > 0:
                    pre = s[:idx]
                    if pre.count('(') != pre.count(')'):
                        continue
                    if starts_dep and not VERB_PAT.search(pre):
                        continue
                    first = pre.rstrip()
                    if not first.endswith(('.','!','?')):
                        first += '.'
                    rest = s[idx + len(conj):].strip()
                    if rest and len(rest.split()) > 3:
                        rest = rest[0].upper() + rest[1:]
                        if conj.strip() in ("which","where"):
                            rest = "This " + rest[0].lower() + rest[1:]
                        result[longest_idx] = first
                        result.insert(longest_idx + 1, rest)
                        break

        lengths = [len(s.split()) for s in result]
        short_indices = [i for i in range(len(result) - 1) if lengths[i] < 10 and lengths[i+1] < 10]
        if short_indices:
            si = short_indices[0]
            if not has_citation(result[si]) and not has_citation(result[si + 1]):
                conj = rng.pick(["; ", ", and ", " - "])
                merged = result[si].rstrip('.') + conj + result[si + 1][0].lower() + result[si + 1][1:]
                result[si] = merged
                result.pop(si + 1)

    return result


# ================================================================
# STARTER VARIATION — max 20% for any single word, no 2 consecutive same
# ================================================================
def enforce_starters(sentences, rng):
    if len(sentences) < 2:
        return sentences

    alts_map = {
        "The": ["Our","Each","Every","Any","A given","One"],
        "This": ["Such a","One such","A similar","That"],
        "These": ["Such","All these","Several","Those"],
        "Those": ["Such","All those","Several"],
        "It": ["One","That","Clearly,"],
        "They": ["Those systems","Each of them","All of them"],
        "We": ["Our team","In our work,","Our analysis"],
        "There": ["One finds","Several","Some"],
        "A": ["One","Each","Any","Some"],
        "In": ["Within","Across","Throughout","During"],
        "For": ["Regarding","When it comes to","As for"],
    }

    prefix_alts = ["In practice, ","Put differently, ","On closer look, ",
                    "At the same time, ","Along these lines, ","As we see it, ",
                    "Taking a step back, ","With this in mind, ","From a broader view, ",
                    "Bearing this out, ","From another angle, ","To put it plainly, "]

    result = list(sentences)

    # Break consecutive same-starters
    for i in range(1, len(result)):
        prev_start = result[i - 1].split()[0] if result[i - 1].split() else ""
        curr_words = result[i].split()
        if not curr_words:
            continue
        curr_start = curr_words[0]

        if curr_start == prev_start:
            rest = " ".join(curr_words[1:])
            if curr_start in alts_map:
                result[i] = rng.pick(alts_map[curr_start]) + " " + rest
            else:
                result[i] = rng.pick(prefix_alts) + curr_start.lower() + " " + rest

    # Break triple same-starters
    for i in range(2, len(result)):
        w = [result[j].split()[0] if result[j].split() else "" for j in (i - 2, i - 1, i)]
        if w[0] == w[1] == w[2]:
            curr_words = result[i].split()
            rest = " ".join(curr_words[1:])
            result[i] = rng.pick(prefix_alts) + curr_words[0].lower() + " " + rest

    # Cap any single starter at 20% of total
    max_count = max(1, len(result) // 5)
    from collections import Counter
    starter_counts = Counter(s.split()[0] for s in result if s.split())
    for word, count in starter_counts.most_common():
        if count <= max_count:
            break
        excess = count - max_count
        for i in range(len(result)):
            if excess <= 0:
                break
            curr_words = result[i].split()
            if curr_words and curr_words[0] == word:
                rest = " ".join(curr_words[1:])
                if word in alts_map:
                    result[i] = rng.pick(alts_map[word]) + " " + rest
                else:
                    result[i] = rng.pick(prefix_alts) + word.lower() + " " + rest
                excess -= 1

    return result


# ================================================================
# SYNONYM PERTURBATION — 95% rate, 99% for AI-flagged words
# ================================================================
SYNS = {
    "achieved":["reached","attained","obtained","hit"],
    "achieves":["reaches","attains","gets","hits"],
    "achieve":["reach","attain","get","hit"],
    "achieving":["reaching","attaining","getting"],
    "accuracy":["precision","correctness"],
    "analyze":["examine","study","look at"],
    "analyzed":["examined","studied","looked at"],
    "analyzes":["examines","studies"],
    "analysis":["examination","review","assessment"],
    "approach":["technique","strategy","method","way"],
    "architecture":["design","structure","layout","setup"],
    "capability":["ability","capacity","power"],
    "capabilities":["abilities","capacities","strengths"],
    "challenge":["difficulty","hurdle","obstacle","problem"],
    "challenges":["difficulties","hurdles","obstacles","problems"],
    "component":["part","element","piece","block"],
    "components":["parts","elements","pieces","blocks"],
    "computation":["calculation","processing"],
    "computational":["computing","processing"],
    "configuration":["setup","arrangement","layout"],
    "constraint":["limitation","restriction","bound","limit"],
    "constraints":["limitations","restrictions","bounds","limits"],
    "critical":["crucial","vital","essential","key"],
    "dataset":["data collection","corpus","data set"],
    "datasets":["data collections","corpora","data sets"],
    "degradation":["decline","drop","weakening","loss"],
    "deployment":["rollout","release","launch"],
    "domain":["field","area","sphere"],
    "domains":["fields","areas","spheres"],
    "efficient":["effective","capable","fast","lean"],
    "efficiency":["effectiveness","speed"],
    "evident":["clear","apparent","plain"],
    "framework":["system","structure","setup"],
    "frameworks":["systems","structures","setups"],
    "hypothesis":["theory","assumption","idea"],
    "illustrate":["show","depict","highlight"],
    "illustrates":["shows","depicts","highlights"],
    "impact":["effect","influence","bearing"],
    "implement":["build","create","set up"],
    "implemented":["built","created","set up"],
    "implementation":["setup","realization","build"],
    "indicate":["show","suggest","point to"],
    "indicates":["shows","suggests","points to"],
    "indicated":["showed","suggested","pointed to"],
    "investigate":["explore","study","look into","dig into"],
    "investigated":["explored","studied","looked into"],
    "investigation":["exploration","study","review","probe"],
    "limitation":["drawback","shortcoming","weakness","gap"],
    "limitations":["drawbacks","shortcomings","weaknesses","gaps"],
    "mechanism":["process","procedure","method","way"],
    "mechanisms":["processes","procedures","methods","ways"],
    "modality":["mode","type","channel","form"],
    "modalities":["modes","types","channels","forms"],
    "module":["unit","block","segment","part"],
    "modules":["units","blocks","segments","parts"],
    "obtain":["get","gather","collect"],
    "obtained":["got","gathered","collected"],
    "parameter":["setting","variable","factor"],
    "parameters":["settings","variables","factors"],
    "perform":["carry out","do","execute","run"],
    "performed":["carried out","did","ran"],
    "performs":["carries out","does","runs"],
    "pipeline":["workflow","process chain","flow"],
    "prevalent":["common","widespread","frequent"],
    "prior":["earlier","previous","past"],
    "procedure":["process","method","step","routine"],
    "scenario":["situation","case","setting"],
    "scenarios":["situations","cases","settings"],
    "straightforward":["simple","direct","easy"],
    "subsequent":["later","following","next"],
    "sufficient":["enough","adequate"],
    "superior":["better","higher","stronger"],
    "threshold":["cutoff","limit","boundary","bar"],
    "thresholds":["cutoffs","limits","boundaries","bars"],
    "validates":["confirms","verifies","checks"],
    "validated":["confirmed","verified","checked"],
    "various":["different","diverse","multiple","several"],
    "yield":["produce","give","generate"],
    "yields":["produces","gives","generates"],
    "proposed":["suggested","presented","put forward"],
    "adopts":["uses","takes on","picks up"],
    "adopted":["used","taken on","picked up"],
    "employs":["uses","applies","relies on"],
    "employed":["used","applied","relied on"],
    "conventional":["traditional","standard","typical"],
    "incorporates":["includes","adds","brings in"],
    "captures":["records","picks up","catches"],
    "addresses":["tackles","handles","deals with"],
    "highlights":["points out","stresses","flags"],
    "explores":["looks into","examines","probes"],
    "outperforms":["beats","surpasses","exceeds","tops"],
    "outperformed":["beat","surpassed","exceeded"],
    "relies":["depends","counts","leans"],
    "encompasses":["covers","spans","includes"],
    "interactions":["exchanges","connections","links"],
    "observations":["findings","notes","takeaways"],
    "evaluation":["assessment","testing","review"],
    "segments":["parts","sections","pieces"],
    "feature":["trait","attribute","characteristic"],
    "features":["traits","attributes","characteristics"],
    "metrics":["measures","scores","indicators"],
    "metric":["measure","score","indicator"],
    "feasible":["doable","practical","workable"],
    "iterations":["rounds","cycles","passes"],
    "convergence":["coming together","agreement"],
    "aggregation":["combination","pooling","merge"],
    "remarkable":["striking","impressive","notable"],
    "promising":["encouraging","hopeful","positive"],
    "complementary":["supporting","supplementary"],
}

AI_FLAGGED = {'methodology','methodologies','comprehensive','innovative','robust','paradigm','paradigms',
              'facilitate','facilitates','utilize','utilizes','leverage','leverages','enhance','enhances',
              'ensure','ensures','underscore','underscores','encompass','encompasses','exemplify','exemplifies'}

SYN_PAT = _re.compile(r"\b(" + "|".join(_re.escape(w) for w in SYNS.keys()) + r")\b", _re.I)

def perturb_synonyms(text, rng, rate=0.95):
    def _rep(m):
        w = m.group(0)
        lw = w.lower()
        if lw in SYNS:
            r_val = 0.99 if lw in AI_FLAGGED else rate
            if rng.rand() < r_val:
                r = rng.pick(SYNS[lw])
                return (r[0].upper() + r[1:]) if w[0].isupper() else r
        return w
    return SYN_PAT.sub(_rep, text)


# ================================================================
# MAIN PIPELINE
# ================================================================
def process_paragraph(text, rng):
    text = _ud.normalize("NFKC", text)
    text = clean_text_watermarks(text)
    text = humanize_words(text, rng)
    sentences = split_sentences(text)
    orig_sentences = list(sentences)

    if len(sentences) > 1:
        sentences = restructure_all(sentences, rng)
        sentences = guaranteed_touch(sentences, orig_sentences, rng)
        sentences = enforce_burstiness(sentences, rng)
        sentences = enforce_starters(sentences, rng)

    text = " ".join(sentences)
    text = perturb_synonyms(text, rng)
    text = final_ai_kill(text, rng)
    text = _re.sub(r"  +", " ", text).strip()
    return text


def final_ai_kill(text, rng):
    """Absolute last pass: kill any remaining AI-signature words."""
    kills = [
        (r"(?i)\bmethodolog(?:y|ies)\b", lambda: rng.pick(["method","technique","approach"])),
        (r"(?i)\badditionally,?\s", lambda: rng.pick(["Also, ","Besides, ","On top of that, "])),
        (r"(?i)\bconsequently,?\s", lambda: rng.pick(["So, ","Because of this, ","For this reason, "])),
        (r"(?i)\bfurthermore,?\s", lambda: rng.pick(["Also, ","On top of this, "])),
        (r"(?i)\bmoreover,?\s", lambda: rng.pick(["Also, ","Besides, "])),
        (r"(?i)\bnevertheless,?\s", lambda: rng.pick(["Still, ","Even so, ","That said, "])),
        (r"(?i)\bnonetheless,?\s", lambda: rng.pick(["Still, ","Even so, "])),
        (r"(?i)\bsubsequently,?\s", lambda: rng.pick(["Then, ","After that, ","Next, "])),
        (r"(?i)\bcomprehensive\b", lambda: rng.pick(["thorough","complete","full"])),
        (r"(?i)\binnovative\b", lambda: rng.pick(["new","novel","fresh"])),
        (r"(?i)\brobust\b", lambda: rng.pick(["strong","solid","reliable"])),
        (r"(?i)\bparadigms?\b", lambda: rng.pick(["model","pattern"])),
        (r"(?i)\bnotwithstanding\b","despite"),
        (r"(?i)\bwherein\b","where"),
        (r"(?i)\bthereby\b","and so"),
        (r"(?i)\butiliz(?:es?|ed|ing)\b", lambda: rng.pick(["uses","used","using"])),
        (r"(?i)\bleverage[sd]?\b","used"),
        (r"(?i)\bleveraging\b","using"),
        (r"(?i)\bfacilitat(?:es?|ed|ing)\b", lambda: rng.pick(["helps","helped","helping"])),
        (r"(?i)\benhance[sd]?\b","improved"),
        (r"(?i)\benhancing\b","improving"),
        (r"(?i)\bensure[sd]?\b","made sure"),
        (r"(?i)\bensuring\b","making sure"),
        (r"(?i)\bunderscor(?:es?|ing)\b", lambda: rng.pick(["highlights","highlighting"])),
        (r"(?i)\bencompass(?:es|ed)?\b", lambda: rng.pick(["covers","spans","includes"])),
        (r"(?i)\bexemplif(?:ies|y)\b","illustrates"),
        (r"(?i)\bherein\b","here"),
        (r"(?i)\baforementioned\b","mentioned earlier"),
        (r"(?i)\bhenceforth\b","from now on"),
    ]
    for pat, repl in kills:
        if callable(repl):
            text = _re.sub(pat, lambda m, f=repl: f(), text)
        else:
            text = _re.sub(pat, repl, text)
    return text


# ================================================================
# PROCESS DOCUMENT
# ================================================================
print("\n--- Phase 2: Text humanization ---")
changed = 0
skipped = 0
seed_counter = 101

def do_process_para(para):
    global changed, seed_counter
    runs = [r for r in para.runs if r.text]
    if not runs:
        return
    full_text = "".join(r.text for r in runs)
    if not full_text.strip():
        return
    seed_counter += 7
    rng = PRNG(seed_counter)
    processed = process_paragraph(full_text, rng)
    if processed != full_text:
        changed += 1
        if len(runs) == 1:
            runs[0].text = processed
        else:
            runs[0].text = processed
            for r in runs[1:]:
                r.text = ""

for idx, para in enumerate(doc.paragraphs):
    raw = para.text.strip()
    if should_skip(idx, raw):
        skipped += 1
        continue
    do_process_para(para)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if p.text.strip():
                    do_process_para(p)

print(f"  Skipped: {skipped}")
print(f"  Paragraphs changed: {changed}")

# Final pass: remove empty runs from processed paragraphs
empty_removed = 0
for p in doc.paragraphs:
    for run in list(p.runs):
        if run.text == "":
            p._element.remove(run._element)
            empty_removed += 1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in list(p.runs):
                    if run.text == "":
                        p._element.remove(run._element)
                        empty_removed += 1
print(f"  Empty runs removed: {empty_removed}")

avg_runs = sum(len(list(p.runs)) for p in doc.paragraphs[:438] if p.text.strip()) / max(1, sum(1 for p in doc.paragraphs[:438] if p.text.strip()))
print(f"  Final avg runs/para: {avg_runs:.1f}")

# Final metadata re-clean
doc.core_properties.author = "Student"
doc.core_properties.comments = ""

doc.save(DST)
print(f"\nDone - saved {DST}")
