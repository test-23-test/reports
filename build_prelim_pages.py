"""
build_prelim_pages.py

Generate a standalone preliminary-pages DOCX:
  Cover Page (with CU Online logo + decorative ornament)
  Declaration
  Acknowledgement
  Abstract / Executive Summary

Each page has a header (logo) and footer (page number + university info).

Output: Preliminary_Pages_YONO_SBI.docx
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm

HERE = Path(__file__).resolve().parent
FIG_DIR = HERE / "figures"
LOGO = FIG_DIR / "template_image1.png"         # CU Online logo
ORNAMENT = FIG_DIR / "template_image2.png"     # decorative flourish
OUT = HERE / "Preliminary_Pages_YONO_SBI.docx"

TITLE = "Digital Banking in Today\u2019s World:\nA Case Study of YONO SBI"
TITLE_UPPER = "DIGITAL BANKING IN TODAY\u2019S WORLD:\nA CASE STUDY OF YONO SBI"
DEGREE = "Master of Business Administration (MBA)"
YEAR = "2024\u20132025"
UNIVERSITY = "Chandigarh University"
DEPT = "Department of Management Studies"
TAGLINE = "CU Online \u2013 Discover. Learn. Empower."

# Colour constants
NAVY = RGBColor(0x1F, 0x4E, 0x79)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)


# ── helpers ──────────────────────────────────────────────────────────────────

def _field_code(run, code: str) -> None:
    for el in (
        _fld("begin"),
        _instr(code),
        _fld("end"),
    ):
        run._r.append(el)


def _fld(typ: str):
    e = OxmlElement("w:fldChar")
    e.set(qn("w:fldCharType"), typ)
    return e


def _instr(code: str):
    e = OxmlElement("w:instrText")
    e.set(qn("xml:space"), "preserve")
    e.text = f" {code} "
    return e


def _shade(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_border_bottom(para, color: str = "1F4E79", sz: int = 6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(4)
    pf.space_before = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.widow_control = True

    for name, sz, sb, sa in (("Heading 1", 16, 10, 4),
                               ("Heading 2", 14, 8, 3),
                               ("Heading 3", 12, 6, 2)):
        s = doc.styles[name]
        s.font.name = "Times New Roman"
        s.font.size = Pt(sz)
        s.font.bold = True
        s.font.color.rgb = NAVY
        s.paragraph_format.space_before = Pt(sb)
        s.paragraph_format.space_after = Pt(sa)
        s.paragraph_format.line_spacing = 1.0
        s.paragraph_format.keep_with_next = True


def setup_margins(doc: Document) -> None:
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        sec.header_distance = Cm(0.8)
        sec.footer_distance = Cm(0.8)


def setup_header_footer(doc: Document) -> None:
    """Add logo header and page-number + university footer to every section."""
    for sec in doc.sections:
        sec.different_first_page_header_footer = False

        # ── Header: blue line separator ──
        hdr = sec.header
        hdr.is_linked_to_previous = False
        for p in list(hdr.paragraphs):
            p.clear()
        hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.space_after = Pt(0)
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.line_spacing = Pt(2)
        _add_border_bottom(hp, color="1F4E79", sz=12)

        # ── Footer: university name left | page number centre | year right ──
        ftr = sec.footer
        ftr.is_linked_to_previous = False
        for p in list(ftr.paragraphs):
            p.clear()
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.paragraph_format.space_before = Pt(2)

        tbl = ftr.add_table(rows=1, cols=3, width=Inches(6.5))
        tbl.alignment = 1  # centre
        for cell in tbl.rows[0].cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)

        # Left cell – university
        c0 = tbl.rows[0].cells[0]
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = c0.paragraphs[0].add_run(UNIVERSITY)
        r0.font.size = Pt(8)
        r0.font.color.rgb = NAVY

        # Centre cell – page number
        c1 = tbl.rows[0].cells[1]
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = c1.paragraphs[0].add_run()
        r1.font.size = Pt(9)
        _field_code(r1, "PAGE")

        # Right cell – department
        c2 = tbl.rows[0].cells[2]
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = c2.paragraphs[0].add_run(DEPT)
        r2.font.size = Pt(8)
        r2.font.color.rgb = NAVY

        # Remove table borders
        tbl_el = tbl._tbl
        tblPr = tbl_el.tblPr if tbl_el.tblPr is not None else OxmlElement("w:tblPr")
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "none")
            e.set(qn("w:sz"), "0")
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), "auto")
            borders.append(e)
        tblPr.append(borders)


def spacer(doc: Document, pts: int = 6) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(pts)
    pf.line_spacing = Pt(2)


def para(doc: Document, text: str, *, bold=False, italic=False,
         center=False, size=None, color=None, spacing_after=None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    if spacing_after is not None:
        p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(it)


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(2)
    p.add_run().add_break(WD_BREAK.PAGE)


# ── Section builders ─────────────────────────────────────────────────────────

def build_cover(doc: Document) -> None:
    spacer(doc, 20)

    # Logo (large, centred)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        p.add_run().add_picture(str(LOGO), width=Inches(4.0))

    spacer(doc, 10)

    # Ornament
    if ORNAMENT.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        p.add_run().add_picture(str(ORNAMENT), width=Inches(1.4))

    para(doc, "A PROJECT REPORT ON", center=True, bold=True, size=14,
         color=NAVY, spacing_after=4)

    para(doc, TITLE_UPPER, center=True, bold=True, size=18, color=RED,
         spacing_after=6)

    spacer(doc, 4)
    para(doc, "Submitted in Partial Fulfilment of the Requirements",
         center=True, italic=True, size=12, spacing_after=2)
    para(doc, "for the Award of the Degree of", center=True, italic=True,
         size=12, spacing_after=2)
    para(doc, DEGREE, center=True, bold=True, size=14, color=NAVY,
         spacing_after=4)

    spacer(doc, 8)

    para(doc, "Submitted By:", center=True, bold=True, size=12,
         spacing_after=2)
    para(doc, "[Student Name]", center=True, size=13, bold=True,
         spacing_after=2)
    para(doc, "Enrollment No.: [XXXXXXXX]", center=True, size=12,
         spacing_after=2)

    spacer(doc, 6)

    para(doc, "Under the Guidance of:", center=True, bold=True, size=12,
         spacing_after=2)
    para(doc, "[Guide Name], [Designation]", center=True, size=13,
         bold=True, spacing_after=2)
    para(doc, DEPT, center=True, size=12, spacing_after=2)
    para(doc, UNIVERSITY, center=True, size=12, bold=True, color=NAVY,
         spacing_after=2)

    spacer(doc, 10)

    # Ornament again
    if ORNAMENT.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(str(ORNAMENT), width=Inches(1.4))

    para(doc, f"Academic Year: {YEAR}", center=True, bold=True, size=14,
         color=NAVY)

    page_break(doc)


def build_declaration(doc: Document) -> None:
    h = doc.add_heading("Declaration", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc,
         "I, [Student Name], hereby solemnly declare that the project "
         f"report titled \u201C{TITLE.replace(chr(10), ': ')}\u201D, "
         "submitted in partial fulfilment of the requirements for the "
         "award of the degree of Master of Business Administration "
         "(MBA), is my original work.")

    para(doc, "I further declare that:", bold=True)
    bullets(doc, [
        f"This project has been carried out by me during the academic "
        f"year {YEAR} under the supervision of [Guide Name] "
        "([Designation]).",
        "The work has not been submitted previously to any other "
        "university, institution, or examination body for the award "
        "of any degree, diploma, or certification.",
        "All sources of information used in this report have been "
        "duly acknowledged and referenced in accordance with "
        "academic ethics and plagiarism norms.",
        "The data presented in this report is authentic to the best "
        "of my knowledge and has not been fabricated or manipulated.",
        "I understand that if any part of this declaration is found "
        "to be false, my project may be rejected and disciplinary "
        "action may be taken as per institutional rules.",
    ])

    spacer(doc, 12)
    para(doc, "Place: ____________________", spacing_after=2)
    para(doc, "Date:  ____________________", spacing_after=2)

    spacer(doc, 12)
    para(doc, "Student Signature: ____________________", spacing_after=2)
    para(doc, "Student Name: [Student Name]", spacing_after=2)
    para(doc, "Enrollment No.: [XXXXXXXX]", spacing_after=2)

    page_break(doc)


def build_acknowledgement(doc: Document) -> None:
    h = doc.add_heading("Acknowledgement", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc,
         "The completion of this project would not have been possible "
         "without the support, guidance, and encouragement of several "
         "individuals and institutions, to whom I wish to express my "
         "sincere gratitude.")
    para(doc,
         "I am deeply grateful to my project guide, [Guide Name], "
         "[Designation], Department of Management Studies, "
         f"{UNIVERSITY}, for providing continuous direction, "
         "constructive feedback, and academic rigour throughout the "
         "course of this study. Their patience and insight were "
         "instrumental in shaping this report.")
    para(doc,
         "I would also like to thank the Dean, the Head of Department, "
         "and the faculty members of the Department of Management "
         f"Studies at {UNIVERSITY} for providing the academic resources "
         "and institutional support that made this project possible.")
    para(doc,
         "I am thankful to the branch managers and staff of State Bank "
         "of India who spared their time for informal discussions, and "
         "to the YONO SBI users who took part in the survey for this "
         "study. Their candid responses form the empirical core of "
         "this work.")
    para(doc,
         "Finally, I wish to thank my family and peers for their "
         "unwavering support and encouragement during the course of "
         "this MBA project.")

    spacer(doc, 12)
    para(doc, "[Student Name]", bold=True, spacing_after=2)
    para(doc, "Enrollment No.: [XXXXXXXX]", spacing_after=2)

    page_break(doc)


def build_abstract(doc: Document) -> None:
    h = doc.add_heading("Abstract / Executive Summary", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc,
         "This study examines the evolution, performance, and strategic "
         "significance of digital banking in India, using YONO SBI as "
         "the primary case. The research addresses five specific "
         "objectives: analysing the growth of digital banking since "
         "2016, evaluating the benefits and challenges of platforms "
         "such as YONO, examining their impact on financial management "
         "practice, identifying emerging trends, and deriving "
         "career-relevant implications for MBA students.")
    para(doc,
         "A descriptive research design with exploratory elements was "
         "adopted. Primary data was collected through a structured "
         "30-item questionnaire administered to 150 YONO SBI users "
         "across urban, semi-urban, and rural India, supplemented by "
         "semi-structured interviews with branch managers and one "
         "fintech professional. Secondary data was drawn from SBI and "
         "RBI publications, NPCI statistics, and peer-reviewed journals. "
         "Data was analysed in Microsoft Excel and IBM SPSS 26 using "
         "descriptive statistics, Pearson correlation, multiple "
         "regression, chi-square tests, and Cronbach\u2019s alpha.")
    para(doc,
         "The study finds that perceived usefulness and ease of use "
         "are the strongest predictors of user satisfaction with YONO, "
         "followed by security perception and accessibility. Younger, "
         "urban respondents report higher usage frequency, while rural "
         "and older users cite onboarding friction and language support "
         "as key barriers. The launch of YONO 2.0 in December 2025 "
         "addresses many of these issues. The report concludes with "
         "actionable recommendations for SBI, regulators, and MBA "
         "students in financial services.")

    spacer(doc, 8)

    # Keywords
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run("Keywords: ")
    r1.bold = True
    p.add_run(
        "Digital Banking, YONO SBI, Technology Acceptance Model, "
        "Financial Inclusion, UPI, Mobile Banking, India, "
        "Perceived Usefulness, Perceived Ease of Use, "
        "Customer Satisfaction."
    )


# ── Main ─────────────────────────────────────────────────────────────────────

def main() -> None:
    print(f"Building: {OUT}")
    doc = Document()
    setup_styles(doc)
    setup_margins(doc)

    build_cover(doc)
    build_declaration(doc)
    build_acknowledgement(doc)
    build_abstract(doc)

    # Apply headers/footers after all content (affects all sections)
    setup_header_footer(doc)

    doc.save(OUT)
    print(f"[OK] Saved: {OUT}")
    import os
    print(f"     Size : {os.path.getsize(OUT) / 1024:.1f} KB")


if __name__ == "__main__":
    main()
