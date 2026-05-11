import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from difflib import SequenceMatcher

orig = Document('CU_finalprojectreport-1.docx')
rewr = Document('CU_finalprojectreport_rewritten.docx')

rewrite_indices = [
    85,88,89,90,91,92,93,94,95,97,98,99,101,102,106,107,108,111,114,
    116,117,118,119,120,121,122,123,124,127,128,129,131,132,133,134,
    135,136,137,138,140,141,143,144,145,146,147,148,149,150,151,153,
    154,156,157,159,160,163,164,166,167,169,170,172,173,174,175,176,
    181,182,183,184,185,188,190,191,194,196,197,199,202,203,204,205,
    206,207,208,209,210,211,213,214,215,217,222,223,225,227,229,231,
    236,239,246,247,248,249,250,252,
    275,276,277,279,280,281,284,288,289,290,292,294,296,298,299,302,308,309,314,315,
    318,320,339,345,347,353,355,360,362,363,364,365,366,
    385,387,388,390,391,392,393,394,396,397,398,399,400,402,403,404,405,406,408,409,410,411,414,415,
    418,420,421,423,424,426,427,429,430,434,435,436,437,439,443,445,446,447,448,449,450,451,452,453,454,455,457,
    460,462,463,465,466,468,469,472,473,475,476,480,481,482,483,
    486,488,490
]

total_sim = 0
count = 0
high60 = 0

for idx in rewrite_indices:
    o = orig.paragraphs[idx].text.strip()
    r = rewr.paragraphs[idx].text.strip()
    if not o or not r:
        continue
    ratio = SequenceMatcher(None, o.lower().split(), r.lower().split()).ratio()
    total_sim += ratio
    count += 1
    if ratio > 0.60:
        high60 += 1

avg_sim = total_sim / count if count else 0
print(f"Paragraphs analyzed: {count}")
print(f"Avg word-level similarity to original: {avg_sim*100:.1f}%")
print(f"Avg content change: {(1-avg_sim)*100:.1f}%")
print(f"Paragraphs >60% similar: {high60} / {count}")
print()

# Burstiness analysis
all_lens = []
for idx in rewrite_indices:
    r = rewr.paragraphs[idx].text.strip()
    if not r or len(r) < 30:
        continue
    sents = [s.strip() for s in re.split(r'[.!?]+', r) if s.strip()]
    lens = [len(s.split()) for s in sents]
    all_lens.extend(lens)

if all_lens:
    avg_len = sum(all_lens) / len(all_lens)
    variance = sum((l - avg_len)**2 for l in all_lens) / len(all_lens)
    std_dev = variance ** 0.5
    print("=== BURSTINESS (sentence length variation) ===")
    print(f"Total sentences: {len(all_lens)}")
    print(f"Avg words/sentence: {avg_len:.1f}")
    print(f"Std deviation: {std_dev:.1f}")
    print(f"Range: {min(all_lens)} to {max(all_lens)} words")
    print(f"Burstiness ratio (std/avg): {std_dev/avg_len:.2f}")
    print()

# Human-like markers
full_text = " ".join(rewr.paragraphs[idx].text for idx in rewrite_indices if rewr.paragraphs[idx].text.strip())
contraction_pat = re.compile(r"\b\w+'[tsrd]\b|\bdon't\b|\bcan't\b|\bwon't\b|\bdoesn't\b|\bwasn't\b|\bisn't\b", re.I)
contractions = len(contraction_pat.findall(full_text))
questions = full_text.count("?")
dashes = full_text.count("\u2014")
print("=== HUMAN-LIKE MARKERS ===")
print(f"Contractions: {contractions}")
print(f"Rhetorical questions: {questions}")
print(f"Em dashes: {dashes}")
print()

# Structure
print(f"Structure: {len(rewr.paragraphs)} paras, {len(rewr.tables)} tables")
o_imgs = len([r for r in orig.part.rels.values() if "image" in r.reltype])
r_imgs = len([r for r in rewr.part.rels.values() if "image" in r.reltype])
print(f"Images: {r_imgs} (orig: {o_imgs})")

for idx, label in [(6,"Declaration"),(19,"Acknowledgement"),(35,"ExecSummary"),(492,"References"),(518,"Appendix")]:
    ok = orig.paragraphs[idx].text == rewr.paragraphs[idx].text
    print(f"{label}: {'INTACT' if ok else 'CHANGED!'}")
