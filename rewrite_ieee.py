import sys, io, copy, random, re as _re, unicodedata as _ud
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

random.seed(42)

# =========================================================
# LOAD & CLONE
# =========================================================
SRC = 'paper_ieee.docx'
DST = 'paper_ieee_rewritten.docx'

doc = Document(SRC)
print(f"Loaded {SRC}: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# =========================================================
# PROTECTED ZONES  (paragraph indices to SKIP rewriting)
# =========================================================
# 0-3  : title, authors, department, supervisor
# 5    : keywords line
# Algorithm pseudocode lines, math formulas, figure captions,
# table headers, and references are detected dynamically below.

SKIP_INDICES = {0, 1, 2, 3, 5}

# References section starts at index 438
REFERENCES_START = 438

def is_section_header(txt):
    return bool(_re.match(
        r'^(I{1,3}V?|VI{0,3}|IX|X|XI{0,3})\.\s', txt
    )) or bool(_re.match(r'^\d+\.\d+', txt))

def is_algorithm_line(txt):
    return bool(_re.match(r'^\d{1,2}[\.\)]\s', txt)) or \
           txt.startswith("Input:") or txt.startswith("Output:")

def is_formula(txt):
    if len(txt) > 200:
        return False

    special = ['\u2211', '\u220f', '\u222b', '\u2208', '\u2200',
               '\u2203', '\u2264', '\u2265', '\u2260', '\u221e',
               '\u2202', '\u2207', '\u00d7', '\u2297', '\u2299',
               '\u221d']
    special_count = sum(1 for c in txt if c in special)
    if special_count >= 2:
        return True

    words = txt.split()
    if len(words) < 8 and _re.search(r'[=≤≥∈∀∃∑∏∫∂∇⊗⊙≈∝]', txt):
        return True
    if txt.count('=') >= 2 and len(words) < 12:
        return True
    if _re.search(r'\b[a-z]_\{', txt):
        return True
    return False

def is_figure_or_table_caption(txt):
    return bool(_re.match(r'^(Fig\.\s*\d|TABLE\s+[IVXLC]|Table\s+[IVXLC])', txt))

def is_bullet_contribution(txt):
    return txt.startswith('\u2022') or txt.startswith('- ') or txt.startswith('* ')

def should_skip(idx, txt):
    if idx in SKIP_INDICES:
        return True
    if idx >= REFERENCES_START:
        return True
    if not txt or len(txt.strip()) < 15:
        return True
    if is_section_header(txt):
        return True
    if is_algorithm_line(txt):
        return True
    if is_formula(txt):
        return True
    if is_figure_or_table_caption(txt):
        return True
    return False


# =========================================================
# PARAGRAPH-LEVEL RESTRUCTURING
# Aggressive sentence rearrangement for technical prose
# =========================================================

def _pick(lst):
    return lst[random.randint(0, len(lst) - 1)]

def restructure_paragraph(text):
    if not text or len(text.strip()) < 30:
        return text

    # Skip restructuring for citation-heavy text (preserve reference integrity)
    citation_count = len(_re.findall(r'\[\d+\]', text))
    if citation_count > 3:
        return text

    sentences = _re.split(r'(?<=[.!?])\s+', text.strip())
    if len(sentences) < 2:
        return text

    result = []
    for i, sent in enumerate(sentences):
        words = sent.split()
        wc = len(words)
        has_citations = bool(_re.search(r'\[\d+\]', sent))

        # Swap clause halves at commas for ~15% of medium sentences WITHOUT citations
        if 15 < wc < 35 and random.random() < 0.15 and not has_citations:
            comma_idx = sent.find(', ', max(8, len(sent)//3))
            if comma_idx > 0 and comma_idx < len(sent) - 15:
                first_half = sent[:comma_idx].strip()
                second_half = sent[comma_idx+2:].strip()
                sw = second_half.split()
                if second_half and len(sw) > 5 and not _re.match(r'^(e\.g\.|i\.e\.|viz\.|[a-z])', second_half):
                    swapped = second_half[0].upper() + second_half[1:] + ", " + first_half[0].lower() + first_half[1:]
                    if not swapped.endswith(('.','!','?')):
                        swapped += '.'
                    sent = swapped

        # Occasionally swap two adjacent sentences (only if neither has citations)
        if i > 0 and random.random() < 0.06 and result and not has_citations:
            prev_has_cit = bool(_re.search(r'\[\d+\]', result[-1]))
            if not prev_has_cit:
                prev = result.pop()
                result.append(sent)
                result.append(prev)
                continue

        # Front-load "however/although" clauses ~10% of the time
        however_match = _re.search(r',?\s+(however|although|though|yet),?\s+', sent, _re.IGNORECASE)
        if however_match and random.random() < 0.10 and not has_citations:
            pos = however_match.start()
            before = sent[:pos].strip().rstrip(',')
            after = sent[however_match.end():].strip()
            conj = however_match.group(1).capitalize()
            if after and before and len(after.split()) > 4:
                sent = f"{conj} {after.rstrip('.')}, {before[0].lower()}{before[1:]}."

        result.append(sent)

    return ' '.join(result)


# =========================================================
# HUMANIZER (all 19 steps from enhanced rewrite_report.py)
# =========================================================

def humanize_text(text):
    if not text or not text.strip():
        return text

    # 1) TRANSITION WORDS
    tr_rules = [
        (r"\bFurthermore,?\s", lambda: _pick(["Also, ", "Plus, ", "And ", "Besides, "])),
        (r"\bMoreover,?\s", lambda: _pick(["Also, ", "What's more, ", "And ", "Plus, "])),
        (r"\bAdditionally,?\s", lambda: _pick(["Also, ", "Plus, ", "And ", "Besides, "])),
        (r"\bConsequently,?\s", lambda: _pick(["So, ", "As a result, ", "That's why "])),
        (r"\bNevertheless,?\s", lambda: _pick(["Still, ", "Even so, ", "That said, ", "But "])),
        (r"\bNonetheless,?\s", lambda: _pick(["Still, ", "Even so, ", "Yet, "])),
        (r"\bSubsequently,?\s", lambda: _pick(["Then, ", "After that, ", "Later, "])),
        (r"\bConversely,?\s", lambda: _pick(["On the flip side, ", "Then again, ", "But "])),
        (r"\bUltimately,?\s", lambda: _pick(["In the end, ", "When all is said and done, "])),
        (r"(?i)\bIn\s+addition,?\s", lambda: _pick(["Also, ", "Plus, ", "And "])),
        (r"(?i)\bAs\s+a\s+result,?\s", lambda: _pick(["So, ", "Because of this, "])),
        (r"(?i)\bIn\s+particular,?\s", lambda: _pick(["Especially, ", "Mainly, "])),
        (r"\bSpecifically,?\s", lambda: _pick(["In particular, ", "Namely, ", "That is, "])),
        (r"\bAccordingly,?\s", lambda: _pick(["So, ", "For that reason, "])),
        (r"\bHence,?\s", lambda: _pick(["So, ", "That's why "])),
        (r"\bThus,?\s", lambda: _pick(["So, ", "This way, "])),
    ]
    for pattern, replacer in tr_rules:
        text = _re.sub(pattern, lambda m: replacer(), text)

    # 2) FORMAL VOCABULARY -> SIMPLER WORDS
    voc_rules = [
        (r"(?i)\butilizes\b", "uses"), (r"(?i)\butilize\b", "use"),
        (r"(?i)\butilized\b", "used"), (r"(?i)\butilizing\b", "using"),
        (r"(?i)\bfacilitates\b", "helps"), (r"(?i)\bfacilitate\b", "help"),
        (r"(?i)\bfacilitated\b", "helped"), (r"(?i)\bfacilitating\b", "helping"),
        (r"(?i)\bdemonstrates\b", "shows"), (r"(?i)\bdemonstrate\b", "show"),
        (r"(?i)\bdemonstrated\b", "showed"), (r"(?i)\bdemonstrating\b", "showing"),
        (r"(?i)\bcomprehensive\b", lambda: _pick(["thorough", "complete", "full"])),
        (r"(?i)\bsignificantly\b", lambda: _pick(["greatly", "noticeably", "a lot"])),
        (r"(?i)\bsignificant\b", lambda: _pick(["big", "major", "notable"])),
        (r"(?i)\bsubstantially\b", lambda: _pick(["greatly", "largely", "considerably"])),
        (r"(?i)\bsubstantial\b", lambda: _pick(["big", "large", "considerable"])),
        (r"(?i)\bapproximately\b", lambda: _pick(["about", "around", "roughly"])),
        (r"(?i)\bnumerous\b", lambda: _pick(["many", "several", "quite a few"])),
        (r"(?i)\boptimal\b", lambda: _pick(["best", "ideal"])),
        (r"(?i)\benhances\b", "improves"), (r"(?i)\benhance\b", "improve"),
        (r"(?i)\benhanced\b", "improved"), (r"(?i)\benhancing\b", "improving"),
        (r"(?i)\bleverages\b", "uses"), (r"(?i)\bleverage\b", "use"),
        (r"(?i)\bleveraging\b", "using"), (r"(?i)\bleveraged\b", "used"),
        (r"(?i)\bensures\b", "makes sure"), (r"(?i)\bensure\b", "make sure"),
        (r"(?i)\bensuring\b", "making sure"),
        (r"(?i)\bpivotal\b", lambda: _pick(["key", "central", "important"])),
        (r"(?i)\bmultifaceted\b", lambda: _pick(["complex", "varied"])),
        (r"(?i)\bstreamlines\b", "simplifies"), (r"(?i)\bstreamline\b", "simplify"),
        (r"(?i)\bstreamlining\b", "simplifying"),
        (r"(?i)\bparamount\b", lambda: _pick(["top priority", "most important"])),
        (r"(?i)\brobust\b", lambda: _pick(["strong", "solid", "reliable"])),
        (r"(?i)\bseamlessly\b", "smoothly"), (r"(?i)\bseamless\b", "smooth"),
        (r"(?i)\binnovative\b", lambda: _pick(["new", "creative", "fresh"])),
        (r"(?i)\bholistic\b", lambda: _pick(["overall", "whole", "complete"])),
        (r"(?i)\bparadigms\b", "models"), (r"(?i)\bparadigm\b", "model"),
        (r"(?i)\bmethodologies\b", "methods"),
        (r"(?i)\bunderscores\b", "highlights"), (r"(?i)\bunderscore\b", "highlight"),
        (r"(?i)\bunderscoring\b", "highlighting"),
        (r"(?i)\bendeavors\b", "efforts"), (r"(?i)\bendeavor\b", "effort"),
        (r"(?i)\baugments\b", "boosts"), (r"(?i)\baugment\b", "boost"),
        (r"(?i)\baugmenting\b", "boosting"),
        (r"(?i)\bmitigates\b", "reduces"), (r"(?i)\bmitigate\b", "reduce"),
        (r"(?i)\bmitigating\b", "reducing"), (r"(?i)\bmitigation\b", "reduction"),
        (r"(?i)\bascertains\b", "finds out"), (r"(?i)\bascertain\b", "find out"),
        (r"(?i)\bascertained\b", "found out"),
        (r"(?i)\bcommences\b", "starts"), (r"(?i)\bcommence\b", "start"),
        (r"(?i)\bcommenced\b", "started"), (r"(?i)\bcommencing\b", "starting"),
        (r"(?i)\bterminates\b", "ends"), (r"(?i)\bterminate\b", "end"),
        (r"(?i)\bterminated\b", "ended"),
        (r"(?i)\bexpedites\b", "speeds up"), (r"(?i)\bexpedite\b", "speed up"),
        (r"(?i)\bexpediting\b", "speeding up"),
        (r"(?i)\brequisite\b", "needed"),
        (r"(?i)\bpredominantly\b", lambda: _pick(["mostly", "mainly"])),
        (r"(?i)\bpredominant\b", lambda: _pick(["main", "primary"])),
        (r"(?i)\bcorroborates\b", "backs up"), (r"(?i)\bcorroborate\b", "back up"),
        (r"(?i)\bcorroborating\b", "backing up"),
        (r"(?i)\belucidated\b", "explained"), (r"(?i)\belucidate\b", "explain"),
        (r"(?i)\belucidating\b", "explaining"),
        (r"(?i)\bameliorate\b", "improve"), (r"(?i)\bameliorates\b", "improves"),
        (r"(?i)\bamelioration\b", "improvement"),
        (r"(?i)\bdelineates\b", "outlines"), (r"(?i)\bdelineate\b", "outline"),
        (r"(?i)\bdelineating\b", "outlining"),
        (r"(?i)\bpropensity\b", "tendency"),
        (r"(?i)\bdiscernible\b", "noticeable"),
        (r"(?i)\bpertaining to\b", "about"),
        (r"(?i)\bwith respect to\b", "about"),
        (r"(?i)\bwith regard to\b", "about"),
        (r"(?i)\bin the context of\b", "in"),
        (r"(?i)\bconducive\b", "helpful"),
        (r"(?i)\bexemplifies\b", "shows"), (r"(?i)\bexemplify\b", "show"),
        (r"(?i)\bexemplified\b", "showed"),
        (r"(?i)\bnotwithstanding\b", "despite"),
        (r"(?i)\bhereinafter\b", "from here on"),
        (r"(?i)\binsofar as\b", "as far as"),
        (r"(?i)\bwherein\b", "where"),
        (r"(?i)\bthereby\b", "by doing this"),
        (r"(?i)\bwhereas\b", "while"),
        (r"(?i)\bprofound\b", lambda: _pick(["deep", "major", "big"])),
        (r"(?i)\bprofoundly\b", lambda: _pick(["deeply", "seriously"])),
        (r"(?i)\bplethora\b", lambda: _pick(["bunch", "ton", "lot"])),
        (r"(?i)\bmyriad\b", lambda: _pick(["tons of", "many", "all sorts of"])),
        (r"(?i)\bakin to\b", lambda: _pick(["like", "similar to", "close to"])),
        (r"(?i)\bfor instance\b", lambda: _pick(["for example", "like", "say"])),
        (r"(?i)\bin contrast\b", lambda: _pick(["on the other hand", "but", "then again"])),
        (r"(?i)\bin summary\b", lambda: _pick(["to sum up", "all in all", "in short"])),
        (r"(?i)\bin conclusion\b", lambda: _pick(["to wrap up", "all in all"])),
        (r"(?i)\bshowcases\b", "shows"), (r"(?i)\bshowcase\b", "show"),
        (r"(?i)\bshowcased\b", "showed"), (r"(?i)\bshowcasing\b", "showing"),
        (r"(?i)\bencompasses\b", "covers"), (r"(?i)\bencompass\b", "cover"),
        (r"(?i)\bencompassing\b", "covering"),
        (r"(?i)\bnarrative\b", lambda: _pick(["story", "account", "picture"])),
        (r"(?i)\bunderpins\b", lambda: _pick(["supports", "backs", "drives"])),
        (r"(?i)\bunderpinning\b", lambda: _pick(["supporting", "backing", "driving"])),
        (r"(?i)\brealm\b", lambda: _pick(["area", "field", "space"])),
        (r"(?i)\blandscape\b", lambda: _pick(["scene", "picture", "space"])),
        (r"(?i)\becosystem\b", lambda: _pick(["space", "world", "setup"])),
        (r"(?i)\btangible\b", lambda: _pick(["real", "concrete", "actual"])),
        (r"(?i)\bnuanced\b", lambda: _pick(["subtle", "detailed", "fine-grained"])),
        (r"(?i)\boverarching\b", lambda: _pick(["main", "big-picture", "central"])),
        (r"(?i)\bcrucial\b", lambda: _pick(["key", "vital", "really important"])),
        (r"(?i)\bcrucially\b", lambda: _pick(["importantly", "vitally"])),
        (r"(?i)\bindispensable\b", lambda: _pick(["essential", "must-have"])),
        (r"(?i)\bintricate\b", lambda: _pick(["complex", "detailed", "involved"])),
        (r"(?i)\bcommendable\b", lambda: _pick(["impressive", "solid", "good"])),
        (r"(?i)\bmeticulous\b", lambda: _pick(["careful", "thorough", "precise"])),
        (r"(?i)\bmeticulously\b", lambda: _pick(["carefully", "thoroughly"])),
        (r"(?i)\bdelve(?:s)?\b", lambda: _pick(["dig", "look", "go"])),
        (r"(?i)\bdelving\b", lambda: _pick(["digging", "looking"])),
        (r"(?i)\bfostering\b", lambda: _pick(["building", "encouraging"])),
        (r"(?i)\bfosters\b", lambda: _pick(["builds", "encourages"])),
        (r"(?i)\bfoster\b", lambda: _pick(["build", "encourage"])),
        (r"(?i)\bharnessing\b", lambda: _pick(["using", "tapping into"])),
        (r"(?i)\bharness(?:es)?\b", lambda: _pick(["use", "tap into"])),
        (r"(?i)\bpivotal\b", lambda: _pick(["key", "central", "critical"])),
        (r"(?i)\bunderpinned\b", lambda: _pick(["supported", "backed", "driven"])),
    ]
    for pattern, repl in voc_rules:
        if callable(repl):
            text = _re.sub(pattern, lambda m, fn=repl: fn(), text)
        else:
            text = _re.sub(pattern, repl, text)

    # 3) AI FILLER PHRASES
    ph_rules = [
        (r"(?i)\bIt is important to note that\s*", ""),
        (r"(?i)\bIt should be noted that\s*", ""),
        (r"(?i)\bIt is worth mentioning that\s*", ""),
        (r"(?i)\bIt is worth noting that\s*", ""),
        (r"(?i)\bIt goes without saying that?\s*", ""),
        (r"(?i)\bNeedless to say,?\s*", ""),
        (r"(?i)\bin order to\b", "to"),
        (r"(?i)\bdue to the fact that\b", "because"),
        (r"(?i)\ba wide range of\b", "many"),
        (r"(?i)\ba plethora of\b", "many"),
        (r"(?i)\bplays a (?:crucial|vital|key|important|significant) role\b", "matters a lot"),
        (r"(?i)\bin today'?s rapidly evolving\b", "in today's"),
        (r"(?i)\bin the ever-changing landscape of\b", "in"),
        (r"(?i)\bat the present time\b", "now"),
        (r"(?i)\bat this point in time\b", "now"),
        (r"(?i)\bin the realm of\b", "in"),
        (r"(?i)\bfor the purpose of\b", "for"),
        (r"(?i)\bin the event that\b", "if"),
        (r"(?i)\bprior to\b", "before"),
        (r"(?i)\bsubsequent to\b", "after"),
        (r"(?i)\bin conjunction with\b", "with"),
        (r"(?i)\bon the basis of\b", "based on"),
        (r"(?i)\bin the absence of\b", "without"),
        (r"(?i)\bhas the potential to\b", "can"),
        (r"(?i)\bis capable of\b", "can"),
        (r"(?i)\ba considerable amount of\b", "a lot of"),
        (r"(?i)\bthe vast majority of\b", "most"),
        (r"(?i)\bin close proximity to\b", "near"),
        (r"(?i)\bat this juncture\b", "now"),
    ]
    for pattern, repl in ph_rules:
        text = _re.sub(pattern, repl, text)

    # 4) CONTRACTIONS
    c_rules = [
        (r"\bdoes not\b", "doesn't"), (r"\bDoes not\b", "Doesn't"),
        (r"\bdo not\b", "don't"), (r"\bDo not\b", "Don't"),
        (r"\bdid not\b", "didn't"), (r"\bDid not\b", "Didn't"),
        (r"\bis not\b", "isn't"), (r"\bIs not\b", "Isn't"),
        (r"\bare not\b", "aren't"), (r"\bAre not\b", "Aren't"),
        (r"\bwas not\b", "wasn't"), (r"\bWas not\b", "Wasn't"),
        (r"\bwere not\b", "weren't"), (r"\bWere not\b", "Weren't"),
        (r"\bcannot\b", "can't"), (r"\bCannot\b", "Can't"),
        (r"\bcould not\b", "couldn't"), (r"\bCould not\b", "Couldn't"),
        (r"\bwould not\b", "wouldn't"), (r"\bWould not\b", "Wouldn't"),
        (r"\bwill not\b", "won't"), (r"\bWill not\b", "Won't"),
        (r"\bshould not\b", "shouldn't"), (r"\bShould not\b", "Shouldn't"),
        (r"\bit is\b", "it's"), (r"\bIt is\b", "It's"),
        (r"\bthat is\b", "that's"), (r"\bThat is\b", "That's"),
        (r"\bthere is\b", "there's"), (r"\bThere is\b", "There's"),
        (r"\bwhat is\b", "what's"), (r"\bWhat is\b", "What's"),
        (r"\bthey are\b", "they're"), (r"\bThey are\b", "They're"),
        (r"\bwe are\b", "we're"), (r"\bWe are\b", "We're"),
        (r"\byou are\b", "you're"), (r"\bYou are\b", "You're"),
        (r"\bI am\b", "I'm"), (r"\bI have\b", "I've"),
        (r"\bI would\b", "I'd"), (r"\bI will\b", "I'll"),
        (r"\bwe have\b", "we've"), (r"\bWe have\b", "We've"),
        (r"\bthey have\b", "they've"), (r"\bThey have\b", "They've"),
        (r"\bwho is\b", "who's"), (r"\bWho is\b", "Who's"),
        (r"\bwhere is\b", "where's"), (r"\bWhere is\b", "Where's"),
        (r"\bhe is\b", "he's"), (r"\bHe is\b", "He's"),
        (r"\bshe is\b", "she's"), (r"\bShe is\b", "She's"),
        (r"\bhave not\b", "haven't"), (r"\bHave not\b", "Haven't"),
        (r"\bhas not\b", "hasn't"), (r"\bHas not\b", "Hasn't"),
        (r"\blet us\b", "let's"), (r"\bLet us\b", "Let's"),
    ]
    for pattern, repl in c_rules:
        text = _re.sub(pattern, repl, text)

    def _safe_lower(cap_letter, text_after_dot):
        """Lowercase unless the word is an acronym or a person's name."""
        word_match = _re.match(r'([A-Za-z]+)', cap_letter + text_after_dot)
        if not word_match:
            return cap_letter.lower()
        full_word = word_match.group(1)
        # All-caps = acronym (BERT, CLIP, OCR, etc.) -> keep
        if len(full_word) > 1 and full_word.isupper():
            return cap_letter
        # Followed by " et al." or " and [A-Z]" = author name -> keep
        rest = text_after_dot[len(full_word)-1:]  # skip rest of first word
        if _re.match(r'\s+et\s+al\.', rest) or _re.match(r'\s+and\s+[A-Z][a-z]', rest):
            return cap_letter
        # Otherwise lowercase
        return cap_letter.lower()

    # 5) SENTENCE OPENER VARIATION — aggressive
    def _vary_opener(m):
        if random.random() < 0.22:
            prefix = _pick(["And ", "But ", "So ", "Now, ", "Still, ", "Yet ",
                            "Sure, ", "True, ", "Mind you, ", "Of course, ",
                            "Notably, ", "In fact, "])
            rest = text[m.end():]
            lowered = _safe_lower(m.group(2), m.group(2) + rest)
            return m.group(1) + prefix + lowered
        return m.group(0)
    text = _re.sub(r"(\. )([A-Z])", _vary_opener, text)

    # 6) SPLIT LONG SENTENCES AT CONJUNCTIONS — more aggressive
    clause_starters = {"the", "this", "these", "that", "it", "they", "we", "our",
                       "a", "an", "each", "every", "most", "some", "such", "its",
                       "their", "when", "if", "as", "since", "while"}
    parts = _re.split(r"(?<=[.!?])\s+", text)
    rebuilt = []
    for s in parts:
        wc = len(s.split())
        if wc > 22 and random.random() < 0.55:
            split_done = False
            for conj in [" but ", " while ", " although ", " because ", " and ", " which ", " where ", " since "]:
                idx = s.find(conj, max(10, len(s) // 4))
                if idx > 0:
                    first = s[:idx].rstrip()
                    second = s[idx + len(conj):].strip()
                    second_words = second.split()
                    first_words = first.split()
                    if (second and len(second_words) > 4 and len(first_words) > 4
                            and second_words[0].lower() in clause_starters):
                        if not first.endswith((".", "!", "?")):
                            first += "."
                        second = second[0].upper() + second[1:]
                        rebuilt.append(first)
                        rebuilt.append(second)
                        split_done = True
                    break
            if not split_done:
                rebuilt.append(s)
        else:
            rebuilt.append(s)
    text = " ".join(rebuilt)

    # 7) INFORMAL DISCOURSE MARKERS — heavy injection
    def _inject_marker(m):
        if random.random() < 0.20:
            marker = _pick([
                "Honestly, ", "Frankly, ", "To be fair, ",
                "In a way, ", "Actually, ", "Really, ",
                "Look, ", "The thing is, ", "Truth be told, ",
                "Granted, ", "Admittedly, ", "Interestingly, ",
                "No doubt, ", "Fair enough, ", "Point being, ",
                "Bottom line, ", "For what it's worth, ",
                "As it happens, ", "If anything, ",
            ])
            rest = text[m.end():]
            lowered = _safe_lower(m.group(2), m.group(2) + rest)
            return m.group(1) + marker + lowered
        return m.group(0)
    text = _re.sub(r"(\. )([A-Z])", _inject_marker, text)

    # 8) HEDGING LANGUAGE — much heavier
    def _hedge(m):
        if random.random() < 0.18:
            h = _pick(["quite ", "fairly ", "rather ", "pretty ", "somewhat ",
                       "arguably ", "relatively ", "reasonably "])
            return h + m.group(0)
        return m.group(0)
    text = _re.sub(
        r"\b(important|clear|strong|difficult|common|effective|useful|evident|high|low|similar|likely|complex|critical|notable|relevant|robust|specific|broad|deep|large|small|rapid|consistent|distinct|obvious|practical)\b",
        _hedge, text
    )

    # 9) PERSONAL VOICE TOUCHES
    pv_rules = [
        (r"(?i)\bIt can be argued that\b", lambda: _pick(["You could say", "One might say"])),
        (r"(?i)\bOne can observe that\b", lambda: _pick(["You can see that", "It's clear that"])),
        (r"(?i)\bIt is evident that\b", lambda: _pick(["Clearly,", "Obviously,"])),
        (r"(?i)\bIt is clear that\b", lambda: _pick(["Clearly,", "Obviously,"])),
        (r"(?i)\bIt is apparent that\b", lambda: _pick(["Clearly,", "Obviously,"])),
        (r"(?i)\bIt can be seen that\b", lambda: _pick(["You can see", "We can see"])),
        (r"(?i)\bIt is noteworthy that\b", lambda: _pick(["Worth noting,", "Interestingly,"])),
    ]
    for pattern, replacer in pv_rules:
        text = _re.sub(pattern, lambda m, fn=replacer: fn(), text)

    # 10) ADVERB VARIATION
    adv_rules = [
        (r"(?i)\bconsiderably\b", lambda: _pick(["a lot", "quite a bit", "noticeably"])),
        (r"(?i)\bremarkably\b", lambda: _pick(["surprisingly", "really"])),
        (r"(?i)\bparticularly\b", lambda: _pick(["especially", "mainly"])),
        (r"(?i)\bprimarily\b", lambda: _pick(["mainly", "mostly", "chiefly"])),
        (r"(?i)\bfundamentally\b", lambda: _pick(["at its core", "basically", "at heart"])),
        (r"(?i)\bincreasingly\b", lambda: _pick(["more and more", "progressively"])),
        (r"(?i)\bcritically\b", lambda: _pick(["crucially", "vitally"])),
        (r"(?i)\bsystematically\b", lambda: _pick(["step by step", "methodically"])),
        (r"(?i)\beffectively\b", lambda: _pick(["well", "in practice", "successfully"])),
        (r"(?i)\bconsequently\b", lambda: _pick(["so", "as a result"])),
        (r"(?i)\baccordingly\b", lambda: _pick(["so", "for that reason"])),
    ]
    for pattern, replacer in adv_rules:
        text = _re.sub(pattern, lambda m, fn=replacer: fn(), text)

    # 11) PASSIVE -> ACTIVE VOICE HINTS
    pv2_rules = [
        (r"(?i)\bwas conducted\b", lambda: _pick(["took place", "happened", "was done"])),
        (r"(?i)\bwere conducted\b", lambda: _pick(["took place", "happened"])),
        (r"(?i)\bwas observed\b", lambda: _pick(["showed up", "came through", "was noticed"])),
        (r"(?i)\bwas identified\b", lambda: _pick(["stood out", "came up", "was spotted"])),
        (r"(?i)\bwere identified\b", lambda: _pick(["stood out", "came up"])),
        (r"(?i)\bis characterized by\b", lambda: _pick(["stands out for", "is known for"])),
        (r"(?i)\bare characterized by\b", lambda: _pick(["stand out for", "are known for"])),
        (r"(?i)\bwas determined\b", lambda: _pick(["turned out", "came out as"])),
        (r"(?i)\bwas established\b", lambda: _pick(["came about", "was set up"])),
        (r"(?i)\bwas implemented\b", lambda: _pick(["went live", "was rolled out", "was put in place"])),
        (r"(?i)\bwere implemented\b", lambda: _pick(["went live", "were rolled out"])),
        (r"(?i)\bwas utilized\b", "was used"),
        (r"(?i)\bwas achieved\b", lambda: _pick(["came through", "worked out"])),
        (r"(?i)\bwas obtained\b", lambda: _pick(["came in", "was gathered"])),
    ]
    for pattern, replacer in pv2_rules:
        text = _re.sub(pattern, lambda m, fn=replacer: fn(), text)

    # 12) REPETITIVE SENTENCE-START BREAKING
    sentences = _re.split(r"(?<=[.!?])\s+", text)
    if len(sentences) > 2:
        for i in range(1, len(sentences)):
            prev_start = sentences[i - 1].split()[0] if sentences[i - 1].split() else ""
            curr_words = sentences[i].split()
            if not curr_words:
                continue
            curr_start = curr_words[0]
            if curr_start == prev_start and curr_start in ("The", "This", "These", "That", "It", "They", "There"):
                if random.random() < 0.6:
                    alt_starters = {
                        "The": ["What stands out about the", "Looking at the", "As for the", "When it comes to the"],
                        "This": ["What this", "Here, this", "In practice, this", "On that note, this"],
                        "These": ["All these", "Taken together, these", "Each of these"],
                        "That": ["And that", "Which is why that", "Notably, that"],
                        "It": ["What it", "In short, it", "As it turns out, it", "Basically, it"],
                        "They": ["In practice, they", "What they", "As it happens, they"],
                        "There": ["As it stands, there", "Right now, there", "In fact, there"],
                    }
                    alts = alt_starters.get(curr_start, [])
                    if alts:
                        replacement = _pick(alts)
                        rest = " ".join(curr_words[1:])
                        sentences[i] = replacement + " " + rest
        text = " ".join(sentences)

    # 13) PARENTHETICAL ASIDES — heavier
    asides = [
        "(and this matters)", "(which is a big deal)", "(not surprisingly)",
        "(to no one's surprise)", "(as you'd expect)", "(worth keeping in mind)",
        "(and this is key)", "(for better or worse)", "(understandably)",
        "(at least in theory)", "(in hindsight)", "(surprisingly enough)",
        "(to put it mildly)", "(believe it or not)", "(rightly so)",
        "(no kidding)", "(go figure)", "(makes sense)", "(fair point)",
    ]
    def _inject_aside(m):
        if random.random() < 0.08:
            aside = _pick(asides)
            return m.group(0).rstrip() + " " + aside + " "
        return m.group(0)
    text = _re.sub(r"[,;]\s", _inject_aside, text)

    # 14) RHETORICAL QUESTION CONVERSION
    rq_patterns = [
        (r"(?i)\bIt is (clear|obvious|evident|apparent) that (.+?)\.", lambda m:
            _pick(["Isn't it " + m.group(1) + " that " + m.group(2) + "?",
                   "Can anyone deny that " + m.group(2) + "?"])),
        (r"(?i)\bEveryone (knows|agrees|understands) that (.+?)\.", lambda m:
            "Does everyone really " + m.group(1).rstrip("s") + " that " + m.group(2) + "?"),
    ]
    for pattern, replacer in rq_patterns:
        if random.random() < 0.5:
            text = _re.sub(pattern, replacer, text, count=1)

    # 15) CLAUSE REORDERING
    def _reorder_because(m):
        if random.random() < 0.15:
            reason = m.group(1).strip().rstrip(",")
            main = m.group(2).strip()
            if main and reason:
                return main[0].upper() + main[1:] + ", because " + reason[0].lower() + reason[1:]
        return m.group(0)
    text = _re.sub(r"\b[Bb]ecause ([^,]+), ([^.!?]+[.!?])", _reorder_because, text)

    # 16) SENTENCE MERGING
    merged_sentences = _re.split(r"(?<=[.!?])\s+", text)
    if len(merged_sentences) > 2:
        rebuilt2 = []
        i = 0
        while i < len(merged_sentences):
            s = merged_sentences[i]
            wc = len(s.split())
            if (wc < 8 and i + 1 < len(merged_sentences)
                    and len(merged_sentences[i + 1].split()) < 8
                    and random.random() < 0.25):
                joiner = _pick(["; ", " - ", " and "])
                next_s = merged_sentences[i + 1]
                if next_s:
                    combined = s.rstrip(".!?") + joiner + next_s[0].lower() + next_s[1:]
                    rebuilt2.append(combined)
                    i += 2
                    continue
            rebuilt2.append(s)
            i += 1
        text = " ".join(rebuilt2)

    # 17) AI N-GRAM PATTERN BREAKING
    ngram_rules = [
        (r"(?i)\bIt is worth noting\b", lambda: _pick(["Keep in mind", "Notice", "One thing to note"])),
        (r"(?i)\bThis is particularly true\b", lambda: _pick(["This really shows", "You see this especially"])),
        (r"(?i)\bThis is especially the case\b", lambda: _pick(["This really holds", "You notice this most"])),
        (r"(?i)\bA growing body of\b", lambda: _pick(["More and more", "A lot of recent"])),
        (r"(?i)\bIn light of\b", lambda: _pick(["Given", "Considering", "Because of"])),
        (r"(?i)\bOn the other hand\b", lambda: _pick(["Then again", "But", "That said"])),
        (r"(?i)\bIn this regard\b", lambda: _pick(["On that front", "Here", "In this area"])),
        (r"(?i)\bFrom this perspective\b", lambda: _pick(["Seen this way", "Looked at like that"])),
        (r"(?i)\bIt is imperative that\b", lambda: _pick(["It's critical that", "We really need to"])),
        (r"(?i)\bThis suggests that\b", lambda: _pick(["That points to", "Which hints that", "So it seems"])),
        (r"(?i)\bThis indicates that\b", lambda: _pick(["That shows", "Which tells us", "So it looks like"])),
        (r"(?i)\bThis implies that\b", lambda: _pick(["That means", "Which suggests", "So basically"])),
        (r"(?i)\bIt is essential to\b", lambda: _pick(["You really need to", "It's critical to", "The key thing is to"])),
        (r"(?i)\bplay(?:s)? an important role\b", lambda: _pick(["matter a lot", "count for a lot", "carry real weight"])),
        (r"(?i)\bserves as a\b", lambda: _pick(["works as a", "acts as a", "doubles as a"])),
        (r"(?i)\ba key factor\b", lambda: _pick(["a big driver", "a real factor", "something that matters"])),
        (r"(?i)\ba critical component\b", lambda: _pick(["a core piece", "something essential", "a must-have part"])),
        (r"(?i)\bin today's (?:digital|modern|contemporary) (?:age|era|world)\b",
         lambda: _pick(["nowadays", "today", "these days", "in the current climate"])),
        (r"(?i)\bthe landscape of\b", lambda: _pick(["the world of", "the state of", "how things stand in"])),
        (r"(?i)\bpaves the way for\b", lambda: _pick(["opens the door to", "sets the stage for", "makes room for"])),
        (r"(?i)\bshed(?:s)? light on\b", lambda: _pick(["clarifies", "helps explain", "brings clarity to"])),
        (r"(?i)\btake(?:s)? into account\b", lambda: _pick(["factor in", "consider", "keep in mind"])),
        (r"(?i)\bthe findings suggest\b", lambda: _pick(["the results point to", "what came out is", "the data shows"])),
        (r"(?i)\bthe results indicate\b", lambda: _pick(["the numbers show", "what we found is", "the data tells us"])),
        (r"(?i)\bas a whole\b", lambda: _pick(["overall", "all in all", "taken together"])),
        (r"(?i)\bgiven the fact that\b", lambda: _pick(["since", "because", "considering"])),
    ]
    for pattern, replacer in ngram_rules:
        text = _re.sub(pattern, lambda m, fn=replacer: fn(), text)

    # 18) EMPHATIC EXPRESSIONS — boosted
    def _emphaticize(m):
        if random.random() < 0.10:
            emphatics = [
                "What really matters here is that ",
                "The bottom line is, ",
                "Put simply, ",
                "Here's the thing: ",
                "The real takeaway is that ",
                "Simply put, ",
                "Long story short, ",
                "In plain terms, ",
                "To be blunt, ",
            ]
            rest = text[m.end():]
            lowered = _safe_lower(m.group(2), m.group(2) + rest)
            return m.group(1) + _pick(emphatics) + lowered
        return m.group(0)
    text = _re.sub(r"(\. )([A-Z])", _emphaticize, text)

    # 19) DOUBLE-SPACE AFTER PERIOD VARIATION
    def _vary_spacing(m):
        if random.random() < 0.08:
            return ".  " + m.group(1)
        return ". " + m.group(1)
    text = _re.sub(r"\. ([A-Z])", _vary_spacing, text)

    # ── 20) ACADEMIC STIFFNESS BREAKERS ──
    acad_rules = [
        (r"(?i)\bthe aforementioned\b", lambda: _pick(["the above", "the earlier", "that"])),
        (r"(?i)\baforementioned\b", lambda: _pick(["above", "earlier", "said"])),
        (r"(?i)\bthe present study\b", lambda: _pick(["this work", "our work", "this paper"])),
        (r"(?i)\bthe current study\b", lambda: _pick(["this work", "our work"])),
        (r"(?i)\bwe posit that\b", lambda: _pick(["we think", "our view is that", "we believe"])),
        (r"(?i)\bit has been shown that\b", lambda: _pick(["studies show", "research shows", "evidence says"])),
        (r"(?i)\bthe empirical evidence\b", lambda: _pick(["the data", "the numbers", "what we found"])),
        (r"(?i)\bempirically\b", lambda: _pick(["experimentally", "through experiments", "with real data"])),
        (r"(?i)\bsalient\b", lambda: _pick(["key", "standout", "striking"])),
        (r"(?i)\bextant\b", lambda: _pick(["existing", "current"])),
        (r"(?i)\bnon-trivial\b", lambda: _pick(["real", "meaningful", "noteworthy"])),
        (r"(?i)\bpreclude(?:s)?\b", lambda: _pick(["prevent", "rule out", "block"])),
        (r"(?i)\bwarrants?\b", lambda: _pick(["calls for", "deserves", "needs"])),
        (r"(?i)\bcoalesce\b", "come together"),
        (r"(?i)\bdichotomy\b", lambda: _pick(["split", "divide", "tension"])),
        (r"(?i)\bcommensurate\b", lambda: _pick(["matching", "proportional", "in line"])),
        (r"(?i)\bjuxtapos(?:ed|es|ition|ing)\b", lambda: _pick(["side by side", "compared", "contrast"])),
        (r"(?i)\bexhibit(?:s|ed)?\b", lambda: _pick(["show", "display", "have"])),
        (r"(?i)\bmanifest(?:s|ed)?\b", lambda: _pick(["show up", "appear", "come through"])),
        (r"(?i)\bpertain(?:s|ing)?\b", lambda: _pick(["relate", "apply", "connect"])),
        (r"(?i)\belicit(?:s|ed)?\b", lambda: _pick(["draw out", "bring out", "trigger"])),
        (r"(?i)\bposit(?:s|ed)?\b", lambda: _pick(["suggest", "argue", "claim"])),
        (r"(?i)\bpurport(?:s|ed)?\b", lambda: _pick(["claim", "supposedly", "allege"])),
        (r"(?i)\bamenable\b", lambda: _pick(["open", "suited", "receptive"])),
        (r"(?i)\btherein\b", lambda: _pick(["in that", "there", "in it"])),
    ]
    for pattern, repl in acad_rules:
        if callable(repl):
            text = _re.sub(pattern, lambda m, fn=repl: fn(), text)
        else:
            text = _re.sub(pattern, repl, text)

    # ── 21) SENTENCE FRAGMENT CREATION ──
    # ~6% of sentences get a short fragment prepended
    def _add_fragment(m):
        if random.random() < 0.06:
            frag = _pick([
                "Worth noting.",
                "A key point.",
                "Not trivial.",
                "Big deal.",
                "Makes sense.",
                "No surprise there.",
                "And here's why.",
                "One more thing.",
                "Quick aside.",
                "Important distinction.",
            ])
            return m.group(1) + frag + " "
        return m.group(0)
    text = _re.sub(r"(\. )(?=[A-Z])", _add_fragment, text)

    # ── 22) WORD ORDER MICRO-VARIATION ──
    # Move trailing prepositional phrases to front ~8% of sentences
    def _front_load_prep(sent):
        m = _re.search(r'^(.+?)((?:,\s*)?(?:in|on|at|by|for|with|from|through|during|after|before|across)\s+[^.!?]{8,40})[.!?]$', sent, _re.IGNORECASE)
        if m and random.random() < 0.08:
            main = m.group(1).rstrip(', ')
            prep = m.group(2).strip().lstrip(', ')
            return prep[0].upper() + prep[1:] + ", " + main[0].lower() + main[1:] + "."
        return sent

    sents = _re.split(r'(?<=[.!?])\s+', text)
    sents = [_front_load_prep(s) for s in sents]
    text = " ".join(sents)

    # ── 23) TONAL SOFTENERS ──
    # Replace assertive phrases with softer alternatives
    soft_rules = [
        (r"(?i)\bclearly shows\b", lambda: _pick(["seems to show", "appears to show", "points toward"])),
        (r"(?i)\bclearly demonstrates\b", lambda: _pick(["seems to show", "appears to confirm"])),
        (r"(?i)\bproves that\b", lambda: _pick(["suggests that", "points to the idea that"])),
        (r"(?i)\bundeniably\b", lambda: _pick(["pretty clearly", "to a large extent"])),
        (r"(?i)\bwithout question\b", lambda: _pick(["most likely", "almost certainly"])),
        (r"(?i)\bwithout doubt\b", lambda: _pick(["very likely", "almost surely"])),
        (r"(?i)\bincontrovertibly\b", lambda: _pick(["by most accounts", "convincingly"])),
        (r"(?i)\bunequivocally\b", lambda: _pick(["pretty clearly", "for the most part"])),
        (r"(?i)\bthe study reveals\b", lambda: _pick(["the data hints", "our results suggest", "the work shows"])),
        (r"(?i)\bthe analysis reveals\b", lambda: _pick(["the analysis suggests", "what we see is"])),
        (r"(?i)\bthe data reveals\b", lambda: _pick(["the data suggests", "the numbers hint"])),
    ]
    for pattern, repl in soft_rules:
        text = _re.sub(pattern, lambda m, fn=repl: fn(), text)

    # ── 24) BURSTINESS INJECTOR ──
    # Occasionally add very short sentence after a long one
    burst_sents = _re.split(r'(?<=[.!?])\s+', text)
    if len(burst_sents) > 3:
        burst_rebuilt = []
        for i, s in enumerate(burst_sents):
            burst_rebuilt.append(s)
            if len(s.split()) > 20 and random.random() < 0.12:
                mini = _pick([
                    "That matters.", "This is key.", "Big difference.",
                    "Not ideal.", "Good sign.", "Worth noting.",
                    "No small thing.", "Fair enough.", "Hard to ignore.",
                    "Think about that.", "Pretty telling.",
                ])
                burst_rebuilt.append(mini)
        text = " ".join(burst_rebuilt)

    return text


# =========================================================
# AI TEXT CLEANER
# =========================================================

def ai_text_clean(text, preserve_boundary_spaces=False):
    leading = ""
    trailing = ""
    if preserve_boundary_spaces:
        lm = _re.match(r'^(\s+)', text)
        tm = _re.search(r'(\s+)$', text)
        if lm:
            leading = lm.group(1)
        if tm:
            trailing = tm.group(1)

    text = _ud.normalize("NFKC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _re.sub(r"[\u200B\u200C\u200D\u2060\uFEFF\u180E]", "", text)
    text = _re.sub(r"[\u00A0\u2007\u202F\u2009\u200A]", " ", text)
    text = _re.sub(r"[\u0000-\u0008\u000B\u000C\u000E-\u001F\u007F-\u009F]", "", text)
    text = _re.sub(r"[\u2013\u2014\u2212]", "-", text)
    text = _re.sub(r'[\u201C\u201D\u201E\u00AB\u00BB]', '"', text)
    text = _re.sub(r"[\u2018\u2019\u201A\u201B\u2039\u203A]", "'", text)
    text = text.replace("\u2026", "...")
    text = _re.sub(r"[\u2022\u25E6\u00B7\u2219]", "-", text)
    text = _re.sub(
        r"[\u00A9\u00AE\u2000-\u3300]|[\U0001F000-\U0001F9FF]",
        "", text
    )
    text = _re.sub(r"\s+([.,!?;:])", r"\1", text)
    text = _re.sub(r"[ \t]*\n[ \t]*", "\n", text)
    text = _re.sub(r"[ \t]{2,}", " ", text)
    text = _re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    text = text.strip()

    if preserve_boundary_spaces:
        if leading and not text.startswith(" "):
            text = " " + text
        if trailing and not text.endswith(" "):
            text = text + " "

    return text


# =========================================================
# APPLY TO DOCUMENT
# =========================================================

restructured = 0
humanized = 0
cleaned = 0
skipped = 0

def process_paragraph_wholistic(para, do_restructure=True):
    """Process paragraph as a single text unit with DOUBLE-PASS humanization."""
    global restructured, humanized, cleaned
    runs = para.runs
    if not runs:
        return

    full_text = "".join(r.text for r in runs)
    if not full_text.strip() or len(full_text.strip()) < 15:
        return

    processed = full_text
    if do_restructure:
        processed = restructure_paragraph(processed)

    # PASS 1: full humanization
    processed = humanize_text(processed)
    # PASS 2: re-seed and run again for deeper transformation
    processed = humanize_text(processed)

    if processed != full_text:
        humanized += 1

        if len(runs) == 1:
            runs[0].text = ai_text_clean(processed)
            cleaned += 1
            if do_restructure:
                restructured += 1
        else:
            runs[0].text = ai_text_clean(processed)
            for r in runs[1:]:
                r.text = ""
            cleaned += 1
            if do_restructure:
                restructured += 1
    else:
        for run in runs:
            if run.text:
                prev = run.text
                run.text = ai_text_clean(run.text, preserve_boundary_spaces=True)
                if run.text != prev:
                    cleaned += 1


for idx, para in enumerate(doc.paragraphs):
    raw = para.text.strip()
    if should_skip(idx, raw):
        skipped += 1
        # Still clean Unicode artifacts on skipped paragraphs (except refs)
        if idx < REFERENCES_START:
            for run in para.runs:
                if run.text:
                    prev = run.text
                    run.text = ai_text_clean(run.text, preserve_boundary_spaces=True)
                    if run.text != prev:
                        cleaned += 1
        continue

    process_paragraph_wholistic(para, do_restructure=True)

# Also process tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                process_paragraph_wholistic(p, do_restructure=False)

print(f"\nParagraphs skipped (headers/formulas/refs): {skipped}")
print(f"Paragraphs restructured+humanized: {restructured}")
print(f"Runs humanized: {humanized}")
print(f"Runs AI-cleaned: {cleaned}")

doc.save(DST)
print(f"\nDone - saved {DST}")
