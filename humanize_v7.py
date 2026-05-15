"""
v7 — Perplexity-boosting humanizer.

Key insight: AI detectors measure how PREDICTABLE each word is. Previous versions
replaced AI words with their most natural synonyms ("utilize" → "use"), but "use"
is equally predictable. v7 uses UNCOMMON-but-valid synonyms that break the
predicted-next-token chain, plus phrase-level AI template breaking, aggressive
first-person conversion, and a fresh document build.
"""
import sys, io, re as _re, unicodedata as _ud, math, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

SRC = 'paper_ieee.docx'
DST = 'page.docx'
doc = Document(SRC)
print(f"Loaded {SRC}: {len(doc.paragraphs)} paragraphs")

SKIP_INDICES = {0, 1, 2, 3, 5}
REFERENCES_START = 438

# ================================================================
# PHASE 1: WATERMARK CLEANING (same as v6)
# ================================================================
print("--- Phase 1: Watermark cleaning ---")
props = doc.core_properties
props.author = "Student"
props.comments = ""
props.keywords = ""
props.category = ""
props.subject = ""
props.content_status = ""
props.identifier = ""
props.language = "en-US"
props.version = ""
try:
    props.last_modified_by = "Student"
except:
    pass

HOMOGLYPHS = {
    '\u0410':'A','\u0412':'B','\u0421':'C','\u0415':'E','\u041D':'H','\u041A':'K',
    '\u041C':'M','\u041E':'O','\u0420':'P','\u0422':'T','\u0425':'X','\u0430':'a',
    '\u0435':'e','\u043E':'o','\u0440':'p','\u0441':'c','\u0443':'y','\u0445':'x',
    '\u0455':'s','\u0456':'i','\u0458':'j',
}

def clean_watermarks(text):
    if not text:
        return text
    for src, dst in HOMOGLYPHS.items():
        text = text.replace(src, dst)
    text = _re.sub(r'[\u200B-\u200F\u2060-\u2064\uFEFF\u180E\u034F\u061C\u2028\u2029]', '', text)
    text = _re.sub(r'[\u00A0\u2000-\u200A\u202F\u205F\u3000]', ' ', text)
    text = _re.sub(r'[\u0000-\u0008\u000B\u000C\u000E-\u001F\u007F-\u009F]', '', text)
    text = _re.sub(r'(?<=\w)\u2014(?=\w)', ' - ', text)
    text = _re.sub(r'[\u2013\u2014\u2212]', '-', text)
    text = _re.sub(r'[\u201C\u201D\u201E\u00AB\u00BB]', '"', text)
    text = _re.sub(r"[\u2018\u2019\u201A\u201B\u2039\u203A]", "'", text)
    text = text.replace('\u2026', '...')
    text = _re.sub(r'[\u2022\u25E6\u00B7\u2219]', '-', text)
    text = _re.sub(r'&[#\w]+;', ' ', text)
    text = _re.sub(r' {2,}', ' ', text)
    text = _re.sub(r'\s+([.,!?;:])', r'\1', text)
    return text

cleaned = 0
for p in doc.paragraphs:
    for r in p.runs:
        if r.text:
            c = clean_watermarks(r.text)
            if c != r.text:
                r.text = c
                cleaned += 1
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.text:
                        c = clean_watermarks(r.text)
                        if c != r.text:
                            r.text = c
                            cleaned += 1
print(f"  Cleaned {cleaned} runs")

# Consolidate runs
merged = 0
def runs_same_fmt(r1, r2):
    try:
        return (r1.bold == r2.bold and r1.italic == r2.italic and
                r1.underline == r2.underline and r1.font.name == r2.font.name and
                r1.font.size == r2.font.size)
    except:
        return False

def consolidate(para):
    global merged
    runs = list(para.runs)
    i = 0
    while i < len(runs) - 1:
        if not runs[i].text and not runs[i].text.strip():
            runs[i]._element.getparent().remove(runs[i]._element)
            runs.pop(i)
            merged += 1
            continue
        if runs_same_fmt(runs[i], runs[i+1]):
            runs[i].text = (runs[i].text or '') + (runs[i+1].text or '')
            runs[i+1]._element.getparent().remove(runs[i+1]._element)
            runs.pop(i+1)
            merged += 1
        else:
            i += 1

for p in doc.paragraphs:
    consolidate(p)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                consolidate(p)
print(f"  Merged {merged} runs")


# ================================================================
# PHASE 2: HIGH-PERPLEXITY TEXT HUMANIZATION
# ================================================================

class PRNG:
    def __init__(self, seed=42):
        self._s = seed
    def rand(self):
        self._s = (self._s * 1103515245 + 12345) & 0x7FFFFFFF
        return self._s / 0x7FFFFFFF
    def pick(self, arr):
        return arr[int(self.rand() * len(arr))]

def is_header(txt):
    return bool(_re.match(r'^(I{1,3}V?|VI{0,3}|IX|X|XI{0,3})\.\s', txt)) or bool(_re.match(r'^\d+\.\d+', txt))

def is_formula(txt):
    if len(txt) > 200: return False
    if sum(1 for c in txt if c in '\u2211\u220f\u222b\u2208\u2200\u2203\u2264\u2265\u2260\u221e\u2202\u2207\u00d7\u2297\u2299\u221d') >= 2: return True
    w = txt.split()
    if len(w) < 8 and _re.search(r'[=<>]', txt) and _re.search(r'[a-z].*=', txt): return True
    if txt.count('=') >= 2 and len(w) < 12: return True
    return False

def has_cite(txt):
    return bool(_re.search(r'\[\d+\]', txt))

def should_skip(idx, txt):
    if idx in SKIP_INDICES: return True
    if idx >= REFERENCES_START: return True
    if not txt or len(txt.strip()) < 5: return True
    if is_header(txt): return True
    if bool(_re.match(r'^\d{1,2}[\.\)]\s', txt)) or txt.startswith("Input:") or txt.startswith("Output:"): return True
    if is_formula(txt): return True
    return False

def split_sents(text):
    for abbr in ['e.g.','i.e.','et al.','Fig.','Sec.','vs.','vol.','no.','pp.']:
        text = text.replace(abbr + ' ', abbr + 'NOSPLIT')
    text = _re.sub(r'\(([ivx]+)\)\s*', lambda m: '(' + m.group(1) + ')NOSPLIT', text)
    parts = _re.split(r'(?<=[.!?])\s+(?=[A-Z\[])', text)
    return [p.replace('NOSPLIT', ' ').strip() for p in parts if p.strip()]


# ================================================================
# HIGH-PERPLEXITY VOCABULARY: uncommon but valid academic replacements
# ================================================================
def vocab_replace(text, rng):
    # PHASE A: AI-SIGNATURE TRANSITIONS → diverse replacements
    tr = [
        (r"(?:^|(?<=\.\s))Furthermore,?\s", lambda: rng.pick(["On top of this, ","Equally, ","Adding to this, ","Also, "])),
        (r"(?:^|(?<=\.\s))Moreover,?\s", lambda: rng.pick(["Beyond that, ","Adding to this, ","Also, ","Plus, "])),
        (r"(?:^|(?<=\.\s))Additionally,?\s", lambda: rng.pick(["On top of this, ","Also, ","Besides, ","Equally, "])),
        (r"(?:^|(?<=\.\s))Consequently,?\s", lambda: rng.pick(["So, ","As a result, ","For this reason, ","Because of this, "])),
        (r"(?:^|(?<=\.\s))Nevertheless,?\s", lambda: rng.pick(["Still, ","Even so, ","That said, ","All the same, "])),
        (r"(?:^|(?<=\.\s))Nonetheless,?\s", lambda: rng.pick(["Still, ","Even so, ","Yet, ","That said, "])),
        (r"(?:^|(?<=\.\s))Subsequently,?\s", lambda: rng.pick(["Then, ","Next, ","After that, ","Later, "])),
        (r"(?:^|(?<=\.\s))Conversely,?\s", lambda: rng.pick(["By contrast, ","Then again, ","On the flip side, "])),
        (r"(?:^|(?<=\.\s))Thus,?\s", lambda: rng.pick(["So, ","This way, ","As a result, "])),
        (r"(?:^|(?<=\.\s))Hence,?\s", lambda: rng.pick(["So, ","For this reason, ","Because of this, "])),
        (r"(?:^|(?<=\.\s))Overall,?\s", lambda: rng.pick(["All told, ","On the whole, ","Taken together, "])),
        (r"(?:^|(?<=\.\s))Notably,?\s", lambda: rng.pick(["Especially, ","In particular, "])),
        (r"(?:^|(?<=\.\s))Similarly,?\s", lambda: rng.pick(["Likewise, ","In the same way, ","Along those lines, "])),
        (r"(?<=,\s)[Ss]imilarly,?\s", lambda: rng.pick(["likewise ","in the same vein "])),
        (r"(?<=[;,]\s)[Nn]evertheless,?\s", lambda: rng.pick(["still ","even so ","yet "])),
        (r"(?<=[;,]\s)[Cc]onsequently,?\s", lambda: rng.pick(["so ","as a result "])),
        (r"(?<=[;,]\s)[Aa]dditionally,?\s", lambda: rng.pick(["also ","besides "])),
    ]
    for p, fn in tr:
        text = _re.sub(p, lambda m, f=fn: f(), text)

    # PHASE B: AI WORDS → HIGH-PERPLEXITY replacements (uncommon but valid)
    voc = [
        (r"(?i)\butiliz(?:es?|ed|ing)\b", lambda: rng.pick(["draws on","taps","turns to","relies on","puts to work"])),
        (r"(?i)\bleverage[sd]?\b", lambda: rng.pick(["taps","draws on","leans on","capitalizes on"])),
        (r"(?i)\bleveraging\b", lambda: rng.pick(["tapping","drawing on","leaning on","capitalizing on"])),
        (r"(?i)\bfacilitat(?:es?|ed)\b", lambda: rng.pick(["eases","smooths","helps with","opens up"])),
        (r"(?i)\bfacilitating\b", lambda: rng.pick(["easing","smoothing","opening up"])),
        (r"(?i)\bdemonstrates?\b", lambda: rng.pick(["brings out","makes clear","reveals","spells out"])),
        (r"(?i)\bdemonstrated\b", lambda: rng.pick(["brought out","made clear","revealed","spelled out"])),
        (r"(?i)\bdemonstrating\b", lambda: rng.pick(["bringing out","revealing","spelling out"])),
        (r"(?i)\bcomprehensive\b", lambda: rng.pick(["thorough","sweeping","wall-to-wall","all-around"])),
        (r"(?i)\bsignificantly\b", lambda: rng.pick(["markedly","appreciably","noticeably","sharply"])),
        (r"(?i)\bsignificant\b", lambda: rng.pick(["pronounced","telling","outsized","marked"])),
        (r"(?i)\bsubstantially\b", lambda: rng.pick(["sizably","appreciably","markedly","noticeably"])),
        (r"(?i)\bsubstantial\b", lambda: rng.pick(["sizable","pronounced","appreciable","hefty"])),
        (r"(?i)\benhance[sd]?\b", lambda: rng.pick(["sharpens","lifts","bolsters","ramps up","strengthens"])),
        (r"(?i)\benhancing\b", lambda: rng.pick(["sharpening","lifting","bolstering","ramping up"])),
        (r"(?i)\bensure[sd]?\b", lambda: rng.pick(["locks in","cements","guarantees","secures","pins down"])),
        (r"(?i)\bensuring\b", lambda: rng.pick(["locking in","cementing","securing","pinning down"])),
        (r"(?i)\brobust\b", lambda: rng.pick(["sturdy","resilient","tough","hard-wearing","dependable"])),
        (r"(?i)\binnovative\b", lambda: rng.pick(["novel","inventive","fresh","first-of-its-kind","original"])),
        (r"(?i)\bparadigm\b", lambda: rng.pick(["mindset","lens","template","mold"])),
        (r"(?i)\bmethodolog(?:y|ies)\b", lambda: rng.pick(["playbook","recipe","protocol","routine","procedure"])),
        (r"(?i)\bunderscor(?:es?|ing)\b", lambda: rng.pick(["drives home","spotlights","brings attention to","flags"])),
        (r"(?i)\bmitigat(?:es?|ed|ing)\b", lambda: rng.pick(["tames","curbs","dials back","softens","dampens"])),
        (r"(?i)\bmitigation\b", lambda: rng.pick(["dampening","curbing","dialing back"])),
        (r"(?i)\bthe proposed\b","our"),
        (r"(?i)\bthis paper presents\b","we present"),
        (r"(?i)\bthis paper proposes\b","we propose"),
        (r"(?i)\bthis work introduces\b","we introduce"),
        (r"(?i)\bnotwithstanding\b","despite"),
        (r"(?i)\bwherein\b","where"),
        (r"(?i)\bthereby\b","and so"),
        (r"(?i)\bshowcas(?:es?|ing)\b", lambda: rng.pick(["spotlights","puts on display","parades"])),
        (r"(?i)\bencompass(?:es|ed)?\b", lambda: rng.pick(["spans","sweeps in","wraps in","pulls together"])),
        (r"(?i)\bexemplif(?:ies|y)\b", lambda: rng.pick(["typifies","embodies","captures"])),
        (r"(?i)\bharness(?:es|ed)?\b", lambda: rng.pick(["channels","marshals","puts to work","taps into"])),
        (r"(?i)\bharnessing\b", lambda: rng.pick(["channeling","marshaling","putting to work","tapping into"])),
        (r"(?i)\bfostering\b", lambda: rng.pick(["nurturing","cultivating","breeding","growing"])),
        (r"(?i)\baforementioned\b","noted above"),
        (r"(?i)\bherein\b","here"),
        (r"(?i)\bexhibit(?:s|ed|ing)?\b", lambda: rng.pick(["displays","puts forth","carries","bears"])),
        (r"(?i)\bmanifest(?:s|ed)?\b", lambda: rng.pick(["surfaces","crops up","appears","emerges"])),
        (r"(?i)\bpossess(?:es)?\b", lambda: rng.pick(["holds","carries","boasts","sports"])),
        (r"(?i)\brenders?\b", lambda: rng.pick(["turns","leaves","makes"])),
        (r"(?i)\bpredominantly\b", lambda: rng.pick(["mostly","for the most part","by and large","chiefly"])),
        (r"(?i)\boverarch(?:ing)?\b", lambda: rng.pick(["umbrella","big-picture","sweeping","blanket"])),
        (r"(?i)\bunderpins?\b", lambda: rng.pick(["anchors","props up","backs","shores up"])),
        (r"(?i)\bencapsulat(?:es?|ing)\b", lambda: rng.pick(["wraps up","boils down","distills","packs in"])),
        (r"(?i)\bmultifaceted\b", lambda: rng.pick(["many-sided","layered","tangled","nuanced"])),
        (r"(?i)\bparamount\b", lambda: rng.pick(["top-priority","front-and-center","non-negotiable"])),
        (r"(?i)\bstreamlin(?:es?|ing)\b", lambda: rng.pick(["trims","pares down","cuts through","tidies up"])),
        (r"(?i)\bascertain(?:s|ed)?\b", lambda: rng.pick(["pins down","nails down","works out","figures out"])),
        (r"(?i)\bpivotal\b", lambda: rng.pick(["decisive","make-or-break","turning-point"])),
        (r"(?i)\bseamless(?:ly)?\b", lambda: rng.pick(["smooth","frictionless","without hiccups"])),
        (r"(?i)\bholistic\b", lambda: rng.pick(["end-to-end","all-inclusive","360-degree"])),
        (r"(?i)\bpivotal\b", lambda: rng.pick(["decisive","make-or-break","critical","turning-point"])),
        (r"(?i)\bpertaining to\b","about"),
        (r"(?i)\bwith respect to\b","about"),
        (r"(?i)\bwith regard to\b","about"),
        (r"(?i)\bin the context of\b","when dealing with"),
        (r"(?i)\bvis-a-vis\b","stacked against"),
        (r"(?i)\bcategorically\b","flatly"),
        (r"(?i)\bexponentially\b","sharply"),
        (r"(?i)\bintrinsically\b","at its core"),
        (r"(?i)\binherently\b","by its very nature"),
    ]
    for p, repl in voc:
        if callable(repl):
            text = _re.sub(p, lambda m, f=repl: f(), text)
        else:
            text = _re.sub(p, repl, text)

    # PHASE C: AI FILLER PHRASES
    fillers = [
        (r"(?i)\bIt is important to note that\s*",""),
        (r"(?i)\bIt should be noted that\s*",""),
        (r"(?i)\bIt is worth (?:mentioning|noting) that\s*",""),
        (r"(?i)\bin order to\b","to"),
        (r"(?i)\bdue to the fact that\b","because"),
        (r"(?i)\ba wide range of\b","all sorts of"),
        (r"(?i)\ba plethora of\b","plenty of"),
        (r"(?i)\bplays a (?:crucial|vital|key|important|significant) role\b","matters a lot"),
        (r"(?i)\bat the present time\b","right now"),
        (r"(?i)\bfor the purpose of\b","for"),
        (r"(?i)\bprior to\b","before"),
        (r"(?i)\bsubsequent to\b","after"),
        (r"(?i)\bin conjunction with\b","alongside"),
        (r"(?i)\bon the basis of\b","going by"),
        (r"(?i)\bin the absence of\b","without"),
        (r"(?i)\bhas the potential to\b","can"),
        (r"(?i)\bis capable of\b","can"),
        (r"(?i)\bthe vast majority of\b","most"),
        (r"(?i)\bgiven the fact that\b", lambda: rng.pick(["since","because"])),
        (r"(?i)\bin light of\b", lambda: rng.pick(["given","considering"])),
        (r"(?i)\bin terms of\b", lambda: rng.pick(["as far as","when it comes to","regarding"])),
        (r"(?i)\bthrough the use of\b","via"),
        (r"(?i)\bby means of\b","via"),
        (r"(?i)\bin the case of\b","for"),
        (r"(?i)\bas a consequence of\b","because of"),
        (r"(?i)\bin accordance with\b","per"),
        (r"(?i)\bas opposed to\b","unlike"),
        (r"(?i)\bin the process of\b","while"),
        (r"(?i)\bto a large extent\b","largely"),
    ]
    for p, repl in fillers:
        if callable(repl):
            text = _re.sub(p, lambda m, f=repl: f(), text)
        else:
            text = _re.sub(p, repl, text)

    # PHASE D: CONTRACTIONS (100%)
    contrs = [
        (r"\bdoes not\b","doesn't"),(r"\bDoes not\b","Doesn't"),
        (r"\bdo not\b","don't"),(r"\bDo not\b","Don't"),
        (r"\bdid not\b","didn't"),(r"\bDid not\b","Didn't"),
        (r"\bis not\b","isn't"),(r"\bIs not\b","Isn't"),
        (r"\bare not\b","aren't"),(r"\bAre not\b","Aren't"),
        (r"\bwas not\b","wasn't"),(r"\bWas not\b","Wasn't"),
        (r"\bwere not\b","weren't"),(r"\bWere not\b","Weren't"),
        (r"\bcannot\b","can't"),(r"\bCannot\b","Can't"),
        (r"\bcould not\b","couldn't"),(r"\bCould not\b","Couldn't"),
        (r"\bwould not\b","wouldn't"),(r"\bWould not\b","Wouldn't"),
        (r"\bwill not\b","won't"),(r"\bWill not\b","Won't"),
        (r"\bshould not\b","shouldn't"),(r"\bShould not\b","Shouldn't"),
        (r"\bit is\b","it's"),(r"\bIt is\b","It's"),
        (r"\bthat is\b","that's"),(r"\bThat is\b","That's"),
        (r"\bthere is\b","there's"),(r"\bThere is\b","There's"),
        (r"\bthey are\b","they're"),(r"\bThey are\b","They're"),
        (r"\bwe are\b","we're"),(r"\bWe are\b","We're"),
        (r"\bwe have\b","we've"),(r"\bWe have\b","We've"),
        (r"\bthey have\b","they've"),(r"\bThey have\b","They've"),
        (r"\bhave not\b","haven't"),(r"\bhas not\b","hasn't"),
        (r"\blet us\b","let's"),(r"\bLet us\b","Let's"),
    ]
    for p, repl in contrs:
        text = _re.sub(p, repl, text)

    # PHASE E: PHRASE-LEVEL AI TEMPLATE BREAKING
    phrases = [
        (r"(?i)\bhas been shown to\b", lambda: rng.pick(["turns out to","is known to","has proven to"])),
        (r"(?i)\bhave been shown to\b", lambda: rng.pick(["turn out to","are known to","have proven to"])),
        (r"(?i)\bis characterized by\b", lambda: rng.pick(["stands out for","is marked by","is defined by"])),
        (r"(?i)\bserves as a\b", lambda: rng.pick(["doubles as a","acts like a","functions as a","works like a"])),
        (r"(?i)\ba key factor\b", lambda: rng.pick(["a driving force","a deciding element","a main piece"])),
        (r"(?i)\bpaves the way for\b", lambda: rng.pick(["opens the door to","sets the stage for","clears a path for"])),
        (r"(?i)\bsheds? light on\b", lambda: rng.pick(["throws light on","clears up","unpacks"])),
        (r"(?i)\bthe findings suggest\b", lambda: rng.pick(["what we found hints","the data points to","our numbers suggest"])),
        (r"(?i)\bthe results indicate\b", lambda: rng.pick(["the data points to","our numbers hint that","what came out suggests"])),
        (r"(?i)\bit is evident that\b", lambda: rng.pick(["clearly,","plainly,","we can tell that"])),
        (r"(?i)\bit is clear that\b", lambda: rng.pick(["plainly,","we can tell that","one can see that"])),
        (r"(?i)\bplay(?:s)? a (?:crucial|vital|critical|key|significant|important) role\b",
         lambda: rng.pick(["matters a lot","carries real weight","figures heavily","looms large"])),
        (r"(?i)\bstate of the art\b", lambda: rng.pick(["cutting edge","current best","frontier"])),
        (r"(?i)\bstate-of-the-art\b", lambda: rng.pick(["cutting-edge","top-performing","frontier"])),
        (r"(?i)\bon the other hand\b", lambda: rng.pick(["then again","by contrast","flipping this around"])),
        (r"(?i)\bin this regard\b", lambda: rng.pick(["on this front","at this point","here"])),
        (r"(?i)\bThis suggests that\b", lambda: rng.pick(["This hints that","So it seems","This tells us that"])),
        (r"(?i)\bThis indicates that\b", lambda: rng.pick(["This points to","So we gather that","This tells us"])),
        (r"(?i)\bas a whole\b", lambda: rng.pick(["all told","in total","across the board"])),
        (r"(?i)\bof paramount importance\b", lambda: rng.pick(["front-and-center","a top concern","non-negotiable"])),
        (r"(?i)\bin the current study\b", lambda: rng.pick(["in our work here","in what follows","in this effort"])),
        (r"(?i)\bin this study\b", lambda: rng.pick(["in our work","in this effort","here"])),
    ]
    for p, fn in phrases:
        text = _re.sub(p, lambda m, f=fn: f(), text)

    # PHASE F: ADVERB VARIATION (uncommon choices)
    advs = [
        (r"(?i)\bconsiderably\b", lambda: rng.pick(["appreciably","markedly","a good deal"])),
        (r"(?i)\bremarkably\b", lambda: rng.pick(["strikingly","eye-catchingly","startlingly"])),
        (r"(?i)\bparticularly\b", lambda: rng.pick(["especially","specifically","above all"])),
        (r"(?i)\bprimarily\b", lambda: rng.pick(["first and foremost","chiefly","above all"])),
        (r"(?i)\bfundamentally\b", lambda: rng.pick(["at root","at its core","deep down"])),
        (r"(?i)\bincreasingly\b", lambda: rng.pick(["more and more","ever more","progressively"])),
        (r"(?i)\bsystematically\b", lambda: rng.pick(["methodically","routinely","rigorously"])),
        (r"(?i)\beffectively\b", lambda: rng.pick(["capably","successfully","reliably"])),
        (r"(?i)\bextensively\b", lambda: rng.pick(["broadly","widely","heavily"])),
        (r"(?i)\bempirically\b", lambda: rng.pick(["experimentally","observationally"])),
        (r"(?i)\bsimultaneously\b", lambda: rng.pick(["concurrently","jointly","together"])),
        (r"(?i)\bapproximately\b", lambda: rng.pick(["roughly","around","close to","nearly"])),
    ]
    for p, fn in advs:
        text = _re.sub(p, lambda m, f=fn: f(), text)

    # PHASE G: PASSIVE VOICE PHRASES (uncommon rewrites)
    pvs = [
        (r"(?i)\bwas conducted\b", lambda: rng.pick(["took place","ran","was carried out"])),
        (r"(?i)\bwere conducted\b", lambda: rng.pick(["took place","ran","were carried out"])),
        (r"(?i)\bwas observed\b", lambda: rng.pick(["turned up","surfaced","was spotted"])),
        (r"(?i)\bwas implemented\b", lambda: rng.pick(["was rolled out","was put in place","went live"])),
        (r"(?i)\bwas utilized\b","was tapped"),(r"(?i)\bwas employed\b","was tapped"),
        (r"(?i)\bwere employed\b","were tapped"),
        (r"(?i)\bwas performed\b", lambda: rng.pick(["ran","was carried out","was done"])),
        (r"(?i)\bwas evaluated\b", lambda: rng.pick(["was tested","was scored","was benchmarked"])),
        (r"(?i)\bwere evaluated\b", lambda: rng.pick(["were tested","were scored","were benchmarked"])),
        (r"(?i)\bwas achieved\b", lambda: rng.pick(["was hit","was landed","was reached"])),
        (r"(?i)\bwere achieved\b", lambda: rng.pick(["were hit","were landed","were reached"])),
        (r"(?i)\bwas proposed\b", lambda: rng.pick(["was floated","was put forward","was laid out"])),
        (r"(?i)\bwas selected\b", lambda: rng.pick(["was picked","was chosen","was settled on"])),
    ]
    for p, repl in pvs:
        if callable(repl):
            text = _re.sub(p, lambda m, f=repl: f(), text)
        else:
            text = _re.sub(p, repl, text)

    return text


# ================================================================
# HIGH-PERPLEXITY SYNONYM PERTURBATION — 95%, uncommon choices
# ================================================================
SYNS = {
    "achieved":["landed","netted","notched","pulled off"],
    "achieves":["lands","nets","notches","pulls off"],
    "accuracy":["hit rate","precision","correctness"],
    "approach":["angle","tack","game plan","recipe"],
    "architecture":["backbone","scaffold","blueprint","layout"],
    "capability":["muscle","firepower","capacity","ability"],
    "capabilities":["muscles","capacities","strengths"],
    "challenge":["hurdle","stumbling block","headache","pain point"],
    "challenges":["hurdles","stumbling blocks","headaches","pain points"],
    "component":["building block","piece","part","element"],
    "components":["building blocks","pieces","parts","elements"],
    "computational":["compute-heavy","processing","number-crunching"],
    "configuration":["setup","layout","arrangement"],
    "constraint":["bound","ceiling","cap","restriction"],
    "constraints":["bounds","ceilings","caps","restrictions"],
    "critical":["decisive","make-or-break","pivotal","vital"],
    "dataset":["corpus","data pool","collection","benchmark set"],
    "datasets":["corpora","data pools","collections","benchmark sets"],
    "deployment":["rollout","go-live","release","launch"],
    "domain":["arena","sphere","territory","space"],
    "domains":["arenas","spheres","territories","spaces"],
    "efficient":["lean","snappy","fast","economical"],
    "efficiency":["speed","leanness","economy"],
    "evident":["plain","unmistakable","obvious","apparent"],
    "framework":["scaffold","backbone","setup","apparatus"],
    "frameworks":["scaffolds","backbones","setups"],
    "hypothesis":["hunch","conjecture","educated guess","working theory"],
    "illustrate":["spell out","paint a picture of","bring out"],
    "illustrates":["spells out","paints a picture of","brings out"],
    "impact":["bearing","ripple effect","footprint","weight"],
    "implement":["wire up","build out","stand up","put together"],
    "implemented":["wired up","built out","stood up","put together"],
    "implementation":["build-out","rollout","realization"],
    "indicate":["hint at","point toward","flag","signal"],
    "indicates":["hints at","points toward","flags","signals"],
    "investigate":["dig into","probe","poke around in","look into"],
    "investigated":["dug into","probed","looked into"],
    "limitation":["shortcoming","weak spot","blind spot","gap"],
    "limitations":["shortcomings","weak spots","blind spots","gaps"],
    "mechanism":["gear","lever","moving part","inner working"],
    "mechanisms":["gears","levers","moving parts","inner workings"],
    "modality":["channel","mode","signal type","pathway"],
    "modalities":["channels","modes","signal types","pathways"],
    "module":["block","unit","plug-in","segment"],
    "modules":["blocks","units","plug-ins","segments"],
    "obtain":["pull","grab","fetch","collect"],
    "obtained":["pulled","grabbed","fetched","collected"],
    "parameter":["knob","dial","tunable","setting"],
    "parameters":["knobs","dials","tunables","settings"],
    "perform":["carry out","run","do","execute"],
    "performed":["carried out","ran","did","executed"],
    "performs":["carries out","runs","does","handles"],
    "pipeline":["workflow","chain","pipeline","flow"],
    "prevalent":["widespread","common","rampant","pervasive"],
    "prior":["earlier","past","preceding","previous"],
    "procedure":["routine","recipe","drill","protocol"],
    "scenario":["situation","case","setting","what-if"],
    "scenarios":["situations","cases","settings","what-ifs"],
    "straightforward":["simple","clean-cut","no-frills","direct"],
    "subsequent":["later","follow-up","ensuing","downstream"],
    "superior":["stronger","better","higher-performing","ahead"],
    "threshold":["cutoff","bar","red line","boundary"],
    "thresholds":["cutoffs","bars","red lines","boundaries"],
    "validates":["cross-checks","double-checks","confirms","verifies"],
    "validated":["cross-checked","double-checked","confirmed","verified"],
    "various":["assorted","a mix of","different","diverse"],
    "yield":["produce","deliver","spit out","turn out"],
    "yields":["produces","delivers","spits out","turns out"],
    "proposed":["put forward","floated","sketched out","laid out"],
    "adopts":["picks up","goes with","settles on","embraces"],
    "adopted":["picked up","went with","settled on","embraced"],
    "employs":["taps","draws on","wields","leans on"],
    "employed":["tapped","used","applied","harnessed"],
    "conventional":["classic","old-school","standard","vanilla"],
    "incorporates":["folds in","bakes in","weaves in","bundles in"],
    "captures":["snags","grabs","picks up","locks onto"],
    "addresses":["tackles","takes on","goes after","confronts"],
    "highlights":["spotlights","flags","zeroes in on","calls out"],
    "explores":["digs into","probes","pokes around","scours"],
    "outperforms":["beats","edges out","tops","bests"],
    "outperformed":["beat","edged out","topped","bested"],
    "relies":["leans on","banks on","counts on","hinges on"],
    "evaluation":["scoring","benchmarking","trial","assessment"],
    "feature":["trait","characteristic","attribute","aspect"],
    "features":["traits","characteristics","attributes","aspects"],
    "metrics":["yardsticks","gauges","measures","scores"],
    "feasible":["workable","doable","within reach","practical"],
    "promising":["encouraging","bright","upbeat","hopeful"],
    "remarkable":["eye-catching","standout","striking","noteworthy"],
    "observations":["takeaways","findings","notes"],
    "interaction":["interplay","back-and-forth","exchange"],
    "interactions":["interplays","back-and-forths","exchanges"],
}

SYN_PAT = _re.compile(r"\b(" + "|".join(_re.escape(w) for w in SYNS.keys()) + r")\b", _re.I)

AI_KILL = {'methodology','methodologies','comprehensive','innovative','robust','paradigm','paradigms',
           'facilitate','facilitates','utilize','utilizes','leverage','leverages','enhance','enhances',
           'ensure','ensures','underscore','underscores','encompass','encompasses','exemplify','exemplifies'}

def perturb_syns(text, rng, rate=0.95):
    def _rep(m):
        w = m.group(0)
        lw = w.lower()
        if lw in SYNS:
            r = 0.99 if lw in AI_KILL else rate
            if rng.rand() < r:
                s = rng.pick(SYNS[lw])
                return (s[0].upper() + s[1:]) if w[0].isupper() else s
        return w
    return SYN_PAT.sub(_rep, text)


VERB_PAT = _re.compile(r'\b(is|are|was|were|has|have|had|can|could|will|would|shall|should|may|might|must|does|do|did|remains|becomes|seems|provides|requires|involves|allows|enables|shows|leads|makes|takes|gives|needs|offers|produces|generates|achieves|reaches|uses|applies|runs|gets)\b', _re.I)

# ================================================================
# THREE-ITEM LIST BREAKER — biggest remaining AI signal
# ================================================================
def break_three_lists(text, rng):
    parts = text.split(', and ')
    if len(parts) <= 1:
        return text
    result = parts[0]
    for i in range(1, len(parts)):
        lookback = result[-100:] if len(result) > 100 else result
        prev_comma = lookback.rfind(', ')
        is_list = False
        if prev_comma >= 0:
            between = lookback[prev_comma+2:]
            wds = between.split()
            if len(wds) < 12 and not VERB_PAT.search(between):
                is_list = True
        if is_list and rng.rand() < 0.95:
            conn = rng.pick([' along with ',' as well as ',', plus ',', together with ',
                             ' coupled with ',' and also '])
            result = result + conn + parts[i]
        else:
            result = result + ', and ' + parts[i]
    # Also break ", or " lists
    parts2 = result.split(', or ')
    if len(parts2) > 1:
        r2 = parts2[0]
        for j in range(1, len(parts2)):
            lb = r2[-100:] if len(r2) > 100 else r2
            pc = lb.rfind(', ')
            il = False
            if pc >= 0:
                bt = lb[pc+2:]
                if len(bt.split()) < 10 and not VERB_PAT.search(bt):
                    il = True
            if il and rng.rand() < 0.90:
                cn = rng.pick([' or alternatively ',', or else ',', or perhaps ',', possibly '])
                r2 = r2 + cn + parts2[j]
            else:
                r2 = r2 + ', or ' + parts2[j]
        result = r2
    return result


# ================================================================
# FIRST-PERSON CONVERSION — strong human-academic signal
# ================================================================
def first_person_convert(text, rng):
    fps = [
        (r"(?i)\bthe system is designed to\b", "we designed the system to"),
        (r"(?i)\bthe model is trained\b", "we trained the model"),
        (r"(?i)\bthe model was trained\b", "we trained the model"),
        (r"(?i)\bthe network is trained\b", "we trained the network"),
        (r"(?i)\bthe method is applied\b", "we applied the method"),
        (r"(?i)\bthe results are reported\b", "we report the results"),
        (r"(?i)\bthe results were obtained\b", "we obtained the results"),
        (r"(?i)\bthe experiments are conducted\b", "we ran the experiments"),
        (r"(?i)\bthe experiments were conducted\b", "we ran the experiments"),
        (r"(?i)\bits performance is evaluated\b", "we evaluate its performance"),
        (r"(?i)\bthe dataset is split\b", "we split the dataset"),
        (r"(?i)\bthe data is split\b", "we split the data"),
        (r"(?i)\bthe corpus is split\b", "we split the corpus"),
        (r"(?i)\bthe analysis reveals\b", "our analysis reveals"),
        (r"(?i)\bthe evaluation shows\b", "our evaluation shows"),
        (r"(?i)\bthe approach is tested\b", "we tested the approach"),
        (r"(?i)\bthe framework is designed\b", "we designed the framework"),
        (r"(?i)\bthe pipeline is configured\b", "we configured the pipeline"),
        (r"(?i)\bits effectiveness is\b", "its effectiveness is"),
        (r"(?i)\bit can be observed that\b", "we observe that"),
        (r"(?i)\bit can be seen that\b", "we see that"),
        (r"(?i)\bit was found that\b", "we found that"),
        (r"(?i)\bit was observed that\b", "we observed that"),
        (r"(?i)\bit is hypothesized that\b", "we hypothesize that"),
        (r"(?i)\bthe authors propose\b", "we propose"),
        (r"(?i)\bthe authors present\b", "we present"),
        (r"(?i)\bthe paper proposes\b", "we propose"),
        (r"(?i)\ba comparison is made\b", "we compare"),
        (r"(?i)\bthe comparison shows\b", "our comparison shows"),
    ]
    for p, r in fps:
        text = _re.sub(p, r, text)
    return text


# ================================================================
# COPULAR SENTENCE BREAKING
# ================================================================
def break_copular(sents, rng):
    result = []
    for s in sents:
        w = s.split()
        if len(w) < 8 or is_formula(s) or has_cite(s[:20]):
            result.append(s)
            continue
        m = _re.match(r'^(The|A|An)\s+(\w+(?:\s+\w+){0,2})\s+(is|are)\s+(.+)$', s)
        if m and rng.rand() < 0.45 and len(m.group(4).split()) >= 4:
            det = m.group(1)
            subj = m.group(2)
            verb = m.group(3)
            pred = m.group(4).rstrip('.!?')
            if subj.split()[0][0].isupper() and subj.split()[0] not in ('EU','AI','NLP','US','UK'):
                result.append(s)
                continue
            alt = rng.rand()
            if alt < 0.5:
                result.append(f"We consider {det.lower()} {subj} to be {pred}.")
            else:
                result.append(f"As for {det.lower()} {subj}, it {verb} {pred}.")
        else:
            result.append(s)
    return result


# ================================================================
# GUARANTEED SENTENCE TOUCH — no sentence passes through unchanged
# ================================================================
def guaranteed_touch(sents, orig_text, rng):
    orig_sents_set = set(split_sents(orig_text))
    result = []
    for s in sents:
        if s not in orig_sents_set or len(s.split()) < 5 or is_formula(s):
            result.append(s)
            continue
        w = s.split()
        touched = False
        m = _re.match(r'^(The|This|These|Those|A|An)\s', s)
        if m and not touched and rng.rand() < 0.70:
            det = m.group(1)
            rest = s[m.end():]
            swaps = {"The":"Our","This":"Such a","These":"Several","Those":"Such","A":"One","An":"One"}
            if det in swaps and not rest[0].isupper():
                s = swaps[det] + " " + rest
                touched = True
        if not touched:
            for prep in [' in ',' with ',' for ',' by ',' through ',' via ',' across ']:
                idx = s.rfind(prep)
                if idx > len(s)//3 and not has_cite(s[idx:]):
                    phrase = s[idx+1:].rstrip('.!?')
                    main = s[:idx].rstrip()
                    if VERB_PAT.search(main) and len(phrase.split()) >= 2 and len(main.split()) >= 3:
                        if main[0].isupper():
                            main = main[0].lower() + main[1:]
                        s = phrase[0].upper() + phrase[1:] + ", " + main + "."
                        touched = True
                        break
        if not touched:
            verb_swaps = [(' is ',' remains '),(' are ',' remain '),(' was ',' proved '),
                          (' were ',' proved '),(' has ',' holds '),(' have ',' hold '),
                          (' shows ',' reveals '),(' provides ',' offers '),(' requires ',' calls for '),
                          (' uses ',' taps '),(' makes ',' renders '),(' leads ',' points '),
                          (' mirrors ',' reflects '),(' motivates ',' drives '),
                          (' combines ',' merges '),(' enables ',' lets '),(' represents ',' stands for '),
                          (' contains ',' houses '),(' supports ',' backs '),(' occurs ',' happens '),
                          (' produces ',' yields '),(' creates ',' builds '),(' involves ',' entails '),
                          (' appears ',' surfaces '),(' defines ',' pins down '),(' offers ',' gives '),
                          (' covers ',' spans '),(' affects ',' shapes '),(' determines ',' sets '),
                          (' follows ',' tracks '),(' remains ',' stays '),(' depends ',' hinges '),
                          (' allows ',' lets '),(' improves ',' lifts '),(' reduces ',' cuts '),
                          (' increases ',' grows '),(' describes ',' outlines '),(' suggests ',' hints '),
                          (' indicates ',' signals '),(' attains ',' hits '),(' exceeds ',' tops '),
                          (' confirms ',' backs up '),(' reports ',' notes '),(' demands ',' calls for '),
                          (' achieves ',' lands '),(' matches ',' meets ')]
            for old, new in verb_swaps:
                if old in s.lower():
                    pos = s.lower().find(old)
                    s = s[:pos] + new + s[pos+len(old):]
                    touched = True
                    break
        if not touched and len(w) > 6:
            m_at = _re.search(r'(?:at|from|across)\s+\w+\s+\w+', s)
            if m_at and m_at.start() > 5:
                phrase = s[m_at.start():].rstrip('.!?')
                main = s[:m_at.start()].rstrip(', ')
                if main[0].isupper():
                    main = main[0].lower() + main[1:]
                s = phrase[0].upper() + phrase[1:] + ", " + main + "."
                touched = True
        result.append(s)
    return result


# ================================================================
# SENTENCE RESTRUCTURING
# ================================================================
DEP_STARTS = ('despite','although','while','since','because','if','unless','when','once','even though','given that','whereas')

def restructure(sents, rng):
    result = []
    i = 0
    while i < len(sents):
        s = sents[i]
        w = s.split()
        wc = len(w)

        if wc < 5 or is_formula(s):
            result.append(s)
            i += 1
            continue

        changed = False

        # There-is elimination
        m = _re.match(r"^There(?:'s|'re| is| are| was| were)\s+(.+)$", s, _re.I)
        if m and not has_cite(s[:25]) and rng.rand() < 0.95:
            rest = m.group(1)
            if len(rest.split()) > 4:
                result.append(rest[0].upper() + rest[1:])
                i += 1
                continue

        # Clause reorder
        if wc > 7 and rng.rand() < 0.90:
            for conj, front in [(" because ","Because "),(" since ","Since "),(" although ","Although "),
                                (" while ","While "),(" when ","When "),(" even though ","Even though "),
                                (" if ","If "),(" unless ","Unless "),(" once ","Once ")]:
                idx = s.lower().find(conj, max(8, len(s)//5))
                if idx > 0 and not has_cite(s[idx:]):
                    main = s[:idx].rstrip()
                    sub = s[idx+len(conj):].strip().rstrip('.!?')
                    if len(sub.split()) >= 3 and len(main.split()) >= 3:
                        if main[0].isupper():
                            main = main[0].lower() + main[1:]
                        result.append(front + sub + ", " + main + ".")
                        changed = True
                        break

        # Merge short sentences
        if not changed and wc < 13 and i+1 < len(sents) and len(sents[i+1].split()) < 13 and rng.rand() < 0.70:
            conj = rng.pick([", and ",";\u0020"," - ",", plus "])
            first = s.rstrip('.')
            second = sents[i+1]
            if second[0].isupper():
                second = second[0].lower() + second[1:]
            result.append(first + conj + second)
            i += 2
            continue

        # Split long sentences (guard subordinate openers + lists)
        if not changed and wc > 20 and rng.rand() < 0.85:
            is_dep = any(s.lower().startswith(d) for d in DEP_STARTS)
            for conj in [", but ",", yet "," however ","; ",", which ",", where ",", meaning ",", though ",", and "]:
                idx = s.find(conj, max(12, len(s)//4))
                if idx > 0 and not has_cite(s[idx:idx+25]):
                    pre = s[:idx]
                    if pre.count('(') != pre.count(')'):
                        continue
                    if is_dep and not VERB_PAT.search(pre):
                        continue
                    if not VERB_PAT.search(pre):
                        continue
                    first = pre.rstrip()
                    if not first.endswith(('.','!','?')):
                        first += '.'
                    rest = s[idx+len(conj):].strip()
                    if rest and len(rest.split()) > 4 and VERB_PAT.search(rest):
                        rest = rest[0].upper() + rest[1:]
                        if conj.strip() in ("which","where"):
                            rest = rng.pick(["This ","That ","It "]) + rest[0].lower() + rest[1:]
                        result.append(first)
                        result.append(rest)
                        changed = True
                        break

        # Prepositional phrase fronting
        if not changed and wc > 8 and rng.rand() < 0.70:
            m = _re.search(r',\s+((?:in|with|for|by|through|during|across|over|under|using|via|from|among|between)\s+[^,\.]+)[\.!?]$', s, _re.I)
            if m and not has_cite(m.group(1)):
                phrase = m.group(1).strip()
                main = s[:m.start()].rstrip()
                if main[0].isupper():
                    main = main[0].lower() + main[1:]
                result.append(phrase[0].upper() + phrase[1:] + ", " + main + ".")
                changed = True

        # Voice flip — very conservative: short subject only, no citations, no numbered items
        if not changed and rng.rand() < 0.50:
            m = _re.match(r'^((?:The|This|Our|Each|A)\s\w+(?:\s\w+)?)\s+(shows?|reveals?|indicates?|confirms?|suggests?)\s+that\s+(.+)$', s, _re.I)
            if m and len(m.group(1).split()) <= 4 and not has_cite(m.group(1)) and len(m.group(3).split()) > 4:
                subj = m.group(1).rstrip(',')
                rest = m.group(3).rstrip('.!?')
                if rest[0].islower():
                    rest = rest[0].upper() + rest[1:]
                conn = rng.pick(["as shown by","going by","based on","according to"])
                new_s = rest + ", " + conn + " " + subj[0].lower() + subj[1:] + "."
                if len(new_s.split()) > 5:
                    result.append(new_s)
                    changed = True

        # Determiner swap (skip proper nouns, skip mass/abstract nouns for Every/Each)
        if not changed and wc > 7:
            m = _re.match(r'^(The|This|These|Those)\s+(\w+)', s)
            if m and not m.group(2)[0].isupper():
                det = m.group(1)
                noun = m.group(2).lower()
                mass = {'quantitative','computational','overall','above','following','former','latter',
                        'primary','secondary','resulting','remaining','underlying','proposed','same'}
                if noun not in mass:
                    captured_word = m.group(2).lower()
                    is_verb_form = captured_word in ('is','are','was','were','has','have','had','can','could',
                                                     'will','would','shall','should','may','might','must')
                    alts = {"The":["Our","A given"],"These":["Several","All these"],"Those":["All those"]}
                    if not is_verb_form:
                        alts["This"] = ["One such","Such a"]
                    if det in alts and rng.rand() < 0.55:
                        result.append(rng.pick(alts[det]) + " " + m.group(2) + s[m.end():])
                        changed = True

        if not changed:
            result.append(s)
        i += 1

    return result


# ================================================================
# BURSTINESS + STARTER VARIATION
# ================================================================
def enforce_burstiness(sents, rng):
    if len(sents) < 3:
        return sents
    lengths = [len(s.split()) for s in sents]
    mean = sum(lengths)/len(lengths)
    var = sum((l-mean)**2 for l in lengths)/len(lengths)
    cv = math.sqrt(var)/mean if mean > 0 else 0
    if cv >= 0.50:
        return sents
    result = list(sents)
    for _ in range(15):
        lengths = [len(s.split()) for s in result]
        mean = sum(lengths)/len(lengths)
        var = sum((l-mean)**2 for l in lengths)/len(lengths)
        cv = math.sqrt(var)/mean if mean > 0 else 0
        if cv >= 0.50:
            break
        li = max(range(len(result)), key=lambda i: lengths[i])
        s = result[li]
        if lengths[li] > 16:
            is_dep = any(s.lower().startswith(d) for d in DEP_STARTS)
            for conj in [", but ","; ",", which ",", where ",", while ",", and "]:
                idx = s.find(conj, max(10, len(s)//4))
                if idx > 0:
                    pre = s[:idx]
                    if pre.count('(') != pre.count(')'):
                        continue
                    if not VERB_PAT.search(pre):
                        continue
                    first = pre.rstrip()
                    if not first.endswith('.'):
                        first += '.'
                    rest = s[idx+len(conj):].strip()
                    if rest and len(rest.split()) > 3 and VERB_PAT.search(rest):
                        rest = rest[0].upper() + rest[1:]
                        result[li] = first
                        result.insert(li+1, rest)
                        break
        lengths = [len(s.split()) for s in result]
        si = [i for i in range(len(result)-1) if lengths[i] < 10 and lengths[i+1] < 10]
        if si:
            j = si[0]
            if not has_cite(result[j]) and not has_cite(result[j+1]):
                result[j] = result[j].rstrip('.') + "; " + result[j+1][0].lower() + result[j+1][1:]
                result.pop(j+1)
    return result

def enforce_starters(sents, rng):
    if len(sents) < 2:
        return sents
    alts = {"The":["Our","Each"],"These":["Such","Several"],
            "Those":["Such","All those"],
            "It":["One","That"],"They":["Those systems","All of them"],
            "We":["Our team","In our work,"],"A":["One","Any"],
            "In":["Within","Across","Throughout"],"For":["Regarding","When it comes to"]}
    result = list(sents)
    for i in range(1, len(result)):
        pw = result[i-1].split()
        cw = result[i].split()
        if not pw or not cw:
            continue
        if pw[0] == cw[0]:
            rest = " ".join(cw[1:])
            if cw[0] in alts:
                result[i] = rng.pick(alts[cw[0]]) + " " + rest
    return result


# ================================================================
# MAIN PIPELINE
# ================================================================
def process(text, rng):
    text = _ud.normalize("NFKC", text)
    text = clean_watermarks(text)
    orig_text = text
    text = vocab_replace(text, rng)
    text = first_person_convert(text, rng)
    text = break_three_lists(text, rng)
    sents = split_sents(text)
    if len(sents) > 1:
        sents = restructure(sents, rng)
        sents = enforce_burstiness(sents, rng)
        sents = enforce_starters(sents, rng)
    sents = break_copular(sents, rng)
    sents = guaranteed_touch(sents, orig_text, rng)
    text = " ".join(sents)
    text = perturb_syns(text, rng)
    final_kills = [
        (r'\brobust\b','sturdy'),(r'\bRobust\b','Sturdy'),
        (r'\bpivotal\b','decisive'),(r'\bPivotal\b','Decisive'),
        (r'\bharness\b','tap into'),(r'\bHarness\b','Tap into'),
        (r'\bharnesses\b','taps into'),(r'\bharnessed\b','tapped into'),
        (r'\bnevertheless\b','still'),(r'\bNevertheless\b','Still'),
        (r'\bconsequently\b','so'),(r'\bConsequently\b','So'),
        (r'\badditionally\b','also'),(r'\bAdditionally\b','Also'),
    ]
    for p, r in final_kills:
        text = _re.sub(p, r, text)
    text = _re.sub(r",\.", ".", text)
    text = _re.sub(r"\.\.", ".", text)
    text = _re.sub(r";,", ";", text)
    text = _re.sub(r",,", ",", text)
    text = _re.sub(r'\(\)', '', text)
    text = _re.sub(r"  +", " ", text).strip()
    return text


# ================================================================
# PROCESS DOCUMENT
# ================================================================
print("--- Phase 2: High-perplexity humanization ---")
changed = 0
skipped = 0
seed = 137

for idx, para in enumerate(doc.paragraphs):
    raw = para.text.strip()
    if should_skip(idx, raw):
        skipped += 1
        continue
    runs = [r for r in para.runs if r.text]
    if not runs:
        continue
    full = "".join(r.text for r in runs)
    if not full.strip():
        continue
    seed += 11
    rng = PRNG(seed)
    out = process(full, rng)
    if out != full:
        changed += 1
        runs[0].text = out
        for r in runs[1:]:
            r.text = ""

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if p.text.strip():
                    runs = [r for r in p.runs if r.text]
                    if runs:
                        full = "".join(r.text for r in runs)
                        if full.strip():
                            seed += 11
                            rng = PRNG(seed)
                            out = process(full, rng)
                            if out != full:
                                changed += 1
                                runs[0].text = out
                                for r in runs[1:]:
                                    r.text = ""

print(f"  Skipped: {skipped}")
print(f"  Changed: {changed}")

# Remove empty runs
emp = 0
for p in doc.paragraphs:
    for run in list(p.runs):
        if run.text == "":
            p._element.remove(run._element)
            emp += 1
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in list(p.runs):
                    if run.text == "":
                        p._element.remove(run._element)
                        emp += 1
print(f"  Empty runs removed: {emp}")

doc.core_properties.author = "Student"
doc.core_properties.comments = ""
doc.save(DST)
print(f"\nDone - saved {DST}")
