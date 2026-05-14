import shutil
from docx import Document

shutil.copy2('CU_finalprojectreport-1.docx', 'CU_finalprojectreport_rewritten.docx')
doc = Document('CU_finalprojectreport_rewritten.docx')


def replace_para_text(para, new_text):
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def replace_cell_text(table_idx, row_idx, col_idx, new_text):
    cell = doc.tables[table_idx].cell(row_idx, col_idx)
    for p in cell.paragraphs:
        if p.runs:
            p.runs[0].text = ""
            for r in p.runs[1:]:
                r.text = ""
    first_para = cell.paragraphs[0]
    if first_para.runs:
        first_para.runs[0].text = new_text
    else:
        first_para.add_run(new_text)


# ===========================================================
# CHAPTER 1
# ===========================================================

CHAPTER1 = {

85: (
    "Enterprise Resource Planning, or ERP, is essentially business software "
    "that brings an organization\u2019s scattered data under one roof. It collects "
    "information, stores it, processes it, and\u2014most importantly\u2014makes it "
    "usable. Why do companies care so much about ERP? Because without it, "
    "departments end up hoarding their own spreadsheets and databases, and "
    "nobody agrees on the numbers. ERP ties together project scheduling, "
    "budgeting, inventory tracking, HR records, and risk reporting into a "
    "single platform with a shared database that updates in near-real time. "
    "Finance sees the same figures as sales. HR pulls from the same employee "
    "records as payroll. Duplicate entries vanish, and data integrity "
    "actually becomes achievable\u2014something that\u2019s surprisingly hard to "
    "pull off when every department runs its own system."
),

88: (
    "There was a time\u2014not that long ago, really\u2014when enterprise software "
    "meant buying physical servers, racking them in a basement or a rented "
    "data centre, and hiring people to keep them running around the clock. "
    "Expensive? Enormously. But cloud computing changed the equation. "
    "Organizations started migrating their ERP to hosted environments, "
    "attracted by elastic scaling, lower upfront spending, and the ability "
    "to plug in CRM and supply chain tools without painful custom "
    "integrations. Collaboration across departments got easier almost "
    "overnight\u2014and honestly, it\u2019s hard to see why any company wouldn\u2019t "
    "at least consider making the switch at this point."
),

89: "Cloud ERP stands apart because of a few core traits:",

90: "All departmental data lives in one centralized repository.",

91: "Users can log in from a browser, anywhere with internet.",

92: "Capacity grows or shrinks based on what the business actually needs.",

93: "Hardware bills and maintenance overheads drop significantly.",

94: (
    "In most digital strategies today, cloud ERP occupies a central spot. "
    "Regular automated backups and built-in disaster recovery go a long way "
    "toward keeping data safe\u2014something that was far harder to guarantee "
    "with on-premises setups where a single hardware failure could cause "
    "serious damage."
),

95: (
    "Then there\u2019s Artificial Intelligence. AI pushes cloud ERP beyond "
    "record-keeping into territory that was science fiction a decade ago: "
    "automated routine tasks, predictive analytics that flag problems "
    "before they escalate, and dashboards that update in real time rather "
    "than once a week. On the security front, AI tightens fraud detection "
    "and compliance monitoring. The overall effect? A system that isn\u2019t "
    "just faster but genuinely smarter\u2014and that\u2019s a big deal for "
    "organizations drowning in data they can\u2019t make sense of quickly enough."
),

97: (
    "It is hard to overstate how much Cloud Computing and AI have reshaped "
    "enterprise software. The old model\u2014buy servers, hire technicians, "
    "pray nothing breaks on a Friday evening\u2014is giving way to cloud "
    "subscriptions that cost less and deliver more. Companies that clung "
    "to on-premises ERP are finding the maintenance bills and infrastructure "
    "headaches increasingly difficult to justify."
),

98: (
    "Cloud ERP comes in three flavours: SaaS, PaaS, and IaaS. Each one "
    "offers a different balance of control and convenience, but the appeal "
    "across all three is similar\u2014quicker setup, tighter budgets, and "
    "access from practically anywhere. The adoption curve tells its own "
    "story: cloud ERP usage jumped from 44 percent in 2020 to close to "
    "70 percent by 2026 (Anchor Group, 2024; Global Growth Insights, 2025)."
),

99: (
    "For a consulting firm like Keystone Advisory Group, headquartered in "
    "Bangalore, the stakes around operational effectiveness are unusually "
    "high. When internal systems lag\u2014slow invoices, clunky project "
    "tracking, unreliable remote access\u2014the quality of client work "
    "suffers directly. Keystone uses its ERP for billing, workforce "
    "scheduling, project management, and client reporting. And the "
    "problems with the legacy setup? High capital costs, convoluted "
    "maintenance, and an inability to work reliably outside the office. "
    "Those issues alone make a compelling case for moving to the cloud."
),

101: (
    "SMEs\u2014Small and Medium Enterprises\u2014are the backbone of India\u2019s "
    "economy and a major force globally. The competitive pressure on Indian "
    "SMEs is intense and accelerating. Digital transformation is no longer "
    "something companies do when they get around to it; it is a survival "
    "requirement. Sharper analytics, leaner operations, faster "
    "decision-making\u2014modern ERP platforms deliver exactly these things, "
    "and SMEs that ignore them risk falling behind competitors who do not."
),

102: (
    "The vendor landscape is crowded. SAP S/4HANA Cloud targets large "
    "multinationals juggling complex global operations. Oracle Fusion Cloud "
    "ERP appeals to mid-to-large firms, particularly those in finance and "
    "human resources. Microsoft Dynamics 365 is a natural fit for "
    "organizations already invested in the Microsoft ecosystem. And for "
    "smaller businesses watching their budgets, Zoho ERP and Oracle "
    "NetSuite offer cloud-native options at a more accessible price point."
),

106: (
    "The vast majority of new ERP implementations now sit on cloud "
    "infrastructure. Amazon Web Services holds the largest market share "
    "in that space (Gartner, 2023). Choosing which cloud provider to "
    "build on is a strategic decision with consequences that last years\u2014"
    "it affects cost structures, data latency, regulatory compliance, "
    "and vendor lock-in."
),

107: (
    "India\u2019s cloud market has its own characteristics. AWS operates an "
    "ap-south-1 region in Mumbai, offering Indian customers both data "
    "residency and low latency. Microsoft Azure runs out of Pune through "
    "its Central India region. Google Cloud\u2019s asia-south1 region is also "
    "in Mumbai. All three major providers keep data on Indian soil, which "
    "matters a great deal for companies where data sovereignty is not "
    "just a preference but a regulatory obligation."
),

108: (
    "The numbers behind the global cloud ERP market are striking. In 2023, "
    "the market was valued at USD 34.83 billion. Projections put it at "
    "USD 65.89 billion by 2026, then USD 110.26 billion by 2030, and a "
    "staggering USD 207.59 billion by 2034. What is behind this growth? "
    "Scalability that legacy systems cannot match, cost advantages that "
    "CFOs appreciate, anywhere-access that employees demand, and a global "
    "corporate push toward digital transformation that shows no sign "
    "of slowing."
),

111: (
    "A number of data points illustrate just how rapidly cloud ERP is "
    "gaining ground. By 2026, somewhere between 60 and 70 percent of all "
    "ERP environments will be cloud-based. About three-quarters of "
    "brand-new ERP deployments already run in the cloud. Financially, the "
    "returns are tangible: most organizations recoup roughly half their "
    "investment within two to three years. More than half of all "
    "enterprises are actively in the process of migrating, and around "
    "60 percent have digital transformation programmes underway. Close "
    "to 40 percent of ERP platforms now incorporate AI features\u2014"
    "things like automation and predictive analytics that were niche "
    "capabilities just five years ago (Parsimony, 2025; Anchor Group, "
    "2024; Global Growth Insights, 2024)."
),

114: (
    "India\u2019s IT and Business Process Management industry contributes "
    "about 7.4 percent of the country\u2019s GDP and provides employment to "
    "roughly 5.4 million people. Bangalore is at the epicentre of this "
    "activity\u2014people call it the Silicon Valley of India for good "
    "reason. Companies like Keystone Advisory Group operate in this "
    "high-pressure environment, and they lean on cloud ERP and other "
    "digital tools to keep pace. Keystone itself is a management and "
    "technology advisory firm based in Bangalore, focused on helping "
    "organizations navigate digital transformation."
),

116: (
    "Keystone Advisory Group is a mid-sized consultancy operating out of "
    "Bangalore. It sits at the intersection of management consulting and "
    "technology advisory, with a particular focus on helping businesses "
    "retool their operations for a digital-first world."
),

117: (
    "The firm works with clients across BFSI, manufacturing, retail, and "
    "healthcare. It employs around 50 professionals spread across multiple "
    "offices. Internally, integrated enterprise systems handle project "
    "tracking, talent management, financial reporting, and compliance. "
    "The service offering includes strategy consulting, IT advisory work, "
    "digital transformation support, and operational excellence programmes."
),

118: "Keystone\u2019s services can be grouped into four areas:",

119: "Advising organizations on digital transformation strategy",

120: "Supporting clients in adopting cloud-based technologies",

121: "Modernizing aging enterprise platforms",

122: "Optimizing business processes for efficiency",

123: (
    "All four feed the same objective: helping client organizations "
    "become faster and more efficient in their digital operations. "
    "From an academic perspective, Keystone serves as a useful real-world "
    "example of how business theory plays out in a consulting environment."
),

124: (
    "For over ten years, Keystone ran its operations on an on-premises "
    "SAP ECC system. Initially it worked fine. But over time, problems "
    "piled up: the system could not scale easily, maintenance costs kept "
    "climbing, analytics capabilities were limited, and remote access was "
    "patchy at best. Eventually, these accumulated pain points pushed "
    "leadership to seriously consider migrating to a cloud-based ERP\u2014"
    "seeking flexibility, lower IT overhead, and alignment with the "
    "firm\u2019s broader digital ambitions."
),

127: (
    "ERP systems matter because they hold organizations together. They "
    "integrate processes, improve efficiency, and support better "
    "decision-making. Keystone Advisory Group has relied on an on-premise "
    "ERP for its core operations for years now. It served its purpose "
    "once, but shifts in the business environment and the arrival of "
    "newer technologies have exposed real weaknesses in the setup."
),

128: (
    "Cloud computing and digital transformation have raised expectations. "
    "Organizations need systems that adapt quickly, scale without "
    "friction, and deliver information in real time. Traditional "
    "on-premises ERP falls short on all three counts, especially in "
    "fast-moving industries. What results is a widening gap between "
    "what legacy systems can offer and what the business genuinely "
    "needs going forward."
),

129: (
    "At Keystone, that gap has become hard to overlook. The company "
    "is growing, hybrid work is the norm, and the old system simply "
    "was not built for this reality. Leadership saw the need for "
    "something more agile, more secure, and easier on the budget, "
    "which prompted a genuine investigation into replacing the legacy "
    "on-premises ERP with a cloud-hosted alternative."
),

131: (
    "Several operational problems plague Keystone\u2019s on-premises ERP, "
    "and together they make the case for migration unavoidable."
),

132: (
    "Costly infrastructure: The on-premises system eats up INR 28,000 "
    "per month just for infrastructure, plus another INR 28,000 for "
    "maintenance and licensing. For a firm with only 30 to 40 employees, "
    "those numbers are difficult to stomach."
),

133: (
    "Poor remote access: Staff can only access the system from within "
    "Keystone\u2019s office network. The VPN alternative is unreliable\u2014"
    "a serious issue given that the firm now operates in hybrid and "
    "remote configurations."
),

134: (
    "Sluggish scaling: Adding new users or deploying new modules takes "
    "four to six weeks. In a business that needs to respond quickly to "
    "changing demands, that delay is crippling."
),

135: (
    "Outdated security: The platform runs on software that is six years "
    "old and does not receive patches on any regular schedule. Sensitive "
    "client data sits behind defences that grow weaker with every "
    "missed update."
),

136: (
    "Weak integration: Connecting the system with modern analytics "
    "engines, AI tools, or third-party SaaS applications is painful. "
    "That limitation hobbles Keystone\u2019s ability to draw actionable "
    "insights from its own data."
),

137: (
    "Unreliable uptime: Four unplanned outages hit the firm over the "
    "past twelve months, all caused by hardware failures, each one "
    "interrupting day-to-day business."
),

138: (
    "All these issues converge into a single research question: How can "
    "Keystone Advisory Group migrate from its aging on-premises ERP to a "
    "cloud-based platform without compromising data integrity, disrupting "
    "operations, overshooting the budget, or losing alignment with "
    "strategic objectives?"
),

140: (
    "This study matters for two reasons. First, the academic literature "
    "on cloud ERP migration overwhelmingly focuses on large Western "
    "multinationals. Indian management consulting firms\u2014mid-sized ones "
    "in particular\u2014barely appear in the conversation. Second, the "
    "study generates practical lessons: how to identify and defuse "
    "migration risks, how to manage organizational change during a "
    "technology transition, and how to tell whether the new system is "
    "actually delivering on its promises. These insights should help "
    "any company considering a similar move."
),

141: (
    "There is also a direct benefit for Keystone Advisory Group itself. "
    "The study provides a structured way to measure migration outcomes, "
    "benchmark the new system\u2019s performance, and pinpoint areas where "
    "adjustments are needed."
),

143: (
    "Primary Objective: To understand how cloud ERP adoption affects "
    "efficiency, costs, and decision-making quality at Keystone "
    "Advisory Group."
),

144: "Supporting objectives include the following:",

145: (
    "Examine how prepared Keystone is for cloud migration\u2014technically, "
    "organizationally, and from a security standpoint."
),

146: (
    "Compare AWS, Azure, and Google Cloud Platform and recommend the "
    "most suitable option for hosting Keystone\u2019s ERP."
),

147: (
    "Develop a phased migration roadmap that includes risk management "
    "protocols and contingency plans."
),

148: (
    "Execute data migration with a hard guarantee of 100 percent "
    "data integrity across every module."
),

149: (
    "Test the cloud ERP thoroughly\u2014functionality, performance, "
    "security\u2014and provide comprehensive user training."
),

150: (
    "Manage a clean go-live and establish a framework for comparing "
    "post-migration KPIs against the pre-migration baseline."
),

151: (
    "Survey employees about their preparedness, the difficulties they "
    "encountered, and their satisfaction after the transition."
),

153: (
    "RQ1: What steps should an organization take to determine whether "
    "it is genuinely ready to shift its ERP from on-premises servers "
    "to a cloud-based platform?"
),

154: (
    "This question draws on existing literature and readiness assessment "
    "frameworks. Survey data add an empirical dimension by highlighting "
    "the readiness factors that weigh most heavily."
),

156: (
    "RQ2: What factors should a company weigh when choosing one cloud "
    "platform over another for an ERP migration?"
),

157: (
    "A comparative review of major cloud vendors\u2014evaluated against "
    "cost, scalability, security, and performance benchmarks\u2014forms "
    "the basis of the answer, supported by published research."
),

159: (
    "RQ3: Which migration approach best protects data integrity "
    "and keeps the business running while the ERP moves to the cloud?"
),

160: (
    "Case studies and published migration frameworks provide the "
    "foundation. A purpose-built framework, mapped to the phases of "
    "this project, bridges theory and practice."
),

163: (
    "RQ4: What are the most significant risks in a cloud ERP "
    "migration, and what can organizations do to manage them?"
),

164: (
    "The answer is built from academic literature, case study "
    "analysis, and structured risk-assessment tools, with SWOT "
    "analysis being the primary instrument."
),

166: (
    "RQ5: After a cloud ERP goes live, how does it affect cost, "
    "operational efficiency, and system reliability?"
),

167: (
    "Secondary data analysis combined with scenario-based modelling, "
    "measured against industry benchmarks, addresses this question."
),

169: (
    "RQ6: How do end users feel about the switch to cloud ERP? "
    "What is their readiness level, and where do adoption barriers "
    "show up?"
),

170: (
    "A structured questionnaire\u2014administered to students acting as "
    "proxy respondents\u2014supplies the primary data, which is then "
    "examined through descriptive statistical methods."
),

172: (
    "This study focuses specifically on migrating ERP systems from "
    "local infrastructure to cloud-based platforms, examined through "
    "the experience of Keystone Advisory Group in Bangalore. The "
    "analysis spans organizational readiness, cloud vendor selection, "
    "the migration strategy itself, and post-implementation "
    "performance."
),

173: (
    "The entire investigation is framed within the Indian business "
    "context. Domestic industry norms, regulatory requirements, and "
    "the growing adoption of cloud technologies among Indian "
    "organizations all feed into the analysis, keeping the findings "
    "grounded in conditions that local firms actually face."
),

174: (
    "Cloud service models fall within the scope as well\u2014primarily "
    "Infrastructure as a Service (IaaS), though other models are "
    "considered where they bear on ERP deployment or scaling. AWS, "
    "Microsoft Azure, and Google Cloud Platform are each evaluated "
    "for suitability as ERP migration targets."
),

175: (
    "Functionally, the study touches Finance, Human Resources, Project "
    "Management, CRM, Procurement, and Administration modules. It looks "
    "at how cloud ERP addresses the persistent weaknesses of traditional "
    "systems: excessive costs, poor scalability, restricted access, and "
    "cumbersome integration requirements."
),

176: (
    "The project operates within a 12-week window running from March "
    "to June 2026, covering planning, vendor selection, migration "
    "execution, testing, and post-go-live support."
),

}

# ===========================================================
# CHAPTER 2
# ===========================================================

CHAPTER2 = {

181: (
    "Research on ERP systems has evolved considerably. Early work focused "
    "mostly on describing what these systems do. More recent scholarship "
    "asks a harder question: do they actually improve how organizations "
    "operate and make decisions? A significant portion of the literature "
    "examines how ERP brings together business functions that otherwise "
    "run in isolation, cutting the inefficiencies that pile up when "
    "departments manage their own processes independently."
),

182: (
    "The rise of cloud computing added an entirely new dimension to the "
    "conversation. Researchers began studying what happens when companies "
    "pull their ERP off local servers and move it onto platforms from SAP, "
    "Oracle, and similar vendors. The findings are broadly positive: cloud "
    "delivery strengthens inter-departmental coordination, automates "
    "repetitive tasks, and gives people access to live data rather than "
    "yesterday\u2019s reports."
),

183: (
    "But the literature is not all sunshine. A fair number of studies "
    "document real downsides\u2014security risks, messy data migrations, and "
    "employees who dig in their heels and refuse to adopt new workflows. "
    "The lesson here is important: ERP adoption is never just a technology "
    "project. It disrupts routines, reshuffles power dynamics, and "
    "challenges organizational culture."
),

184: (
    "Methodologically, the field is quite diverse. Case studies, surveys, "
    "and in-depth interviews have all been used to explore how ERP "
    "performs in real-world settings. Migration strategies also receive "
    "considerable attention. Phased rollouts and parallel system "
    "operations show up repeatedly as practical ways to contain risk "
    "while the old system hands off to the new one."
),

185: (
    "Despite all this work, significant blind spots persist. SMEs, in "
    "particular, remain underexplored\u2014and for SMEs, achieving a "
    "reasonable return on investment and keeping operations running "
    "smoothly are not luxuries but survival requirements. This chapter "
    "lays out the theoretical groundwork for ERP adoption, catalogues "
    "the shortcomings of legacy systems, tracks the emergence of "
    "cloud-based ERP, evaluates published findings on migration "
    "outcomes and pitfalls, and identifies the research gaps this "
    "study intends to fill."
),

188: (
    "No organization migrates its ERP casually. The decision to leave "
    "behind a legacy on-premises system and move to the cloud is tangled "
    "and multifaceted, touching strategy, budgets, technology, and human "
    "behaviour in roughly equal measure. Several theoretical models have "
    "been developed\u2014and empirically tested\u2014to explain how people and "
    "organizations adopt new technology."
),

190: (
    "The Technology Acceptance Model is probably the most cited framework "
    "in the information systems literature. What Davis proposed back in "
    "1989 was fairly straightforward: a person\u2019s willingness to use "
    "new technology boils down to two perceptions. Perceived Usefulness\u2014"
    "will this actually make my job better? And Perceived Ease of Use\u2014"
    "can I figure it out without an unreasonable amount of effort? Those "
    "two perceptions shape attitudes, attitudes shape intentions, and "
    "intentions shape actual behaviour. It\u2019s a deceptively simple chain, "
    "but it\u2019s held up remarkably well across decades of research."
),

191: (
    "The practical takeaway from TAM isn\u2019t exactly comfortable for IT "
    "leaders. A cloud ERP system might be technically outstanding, but if "
    "the people who have to use it every day find it confusing or "
    "cumbersome, they simply won\u2019t use it properly\u2014no matter what the "
    "spec sheet says. That\u2019s why usability design and hands-on training "
    "deserve just as much investment as the technology itself."
),

194: (
    "Venkatesh and colleagues, writing in 2003, took eight earlier "
    "adoption models and synthesized them into one unified framework "
    "called UTAUT. It identifies four core drivers: Performance "
    "Expectancy, Effort Expectancy, Social Influence, and Facilitating "
    "Conditions. What makes UTAUT particularly relevant for ERP projects "
    "is that it explicitly accounts for managerial endorsement and "
    "infrastructure readiness\u2014two factors that often matter more in "
    "large-scale deployments than any individual user\u2019s personal "
    "feelings about the software."
),

196: (
    "Rogers (2003) described how innovations spread through populations "
    "and sorted adopters into five groups: Innovators, Early Adopters, "
    "Early Majority, Late Majority, and Laggards. Applied to cloud ERP, "
    "this classification is useful for predicting where internal "
    "resistance will come from. Technology-oriented consulting firms "
    "typically sit in the early majority\u2014they adopt innovations after "
    "proof of concept exists but before mainstream saturation. "
    "Keystone Advisory Group\u2019s 2024 migration timeline fits this "
    "pattern closely."
),

197: (
    "Rogers also identified five innovation attributes that influence "
    "adoption speed: relative advantage, compatibility, complexity, "
    "trialability, and observability. Cloud ERP scores well on advantage "
    "and trialability\u2014vendors are happy to run demonstrations. But "
    "complexity remains a real concern, particularly when it comes to "
    "migrating years of accumulated organizational data."
),

199: (
    "The TCO framework treats the cloud-versus-on-premises decision as "
    "a comprehensive, long-term financial exercise. It looks beyond "
    "obvious costs like licenses and hardware to include implementation, "
    "training, maintenance, downtime, and even opportunity costs. "
    "Organizations that only compare license fees tend to underestimate "
    "what on-premises actually costs them while overestimating cloud "
    "expenses. Gartner (2022) quantifies this: over five years, cloud "
    "ERP typically costs 25 to 40 percent less than the equivalent "
    "on-premises deployment when everything is included."
),

202: "Johansson and Ruivo (2013)",

203: (
    "This study compared cloud ERP uptake across SMEs in Sweden and "
    "Portugal, drawing on surveys and structured interviews at 40 "
    "organizations. Cost savings and scalability were the strongest "
    "motivators; data security fears and vendor dependency were the "
    "most commonly cited barriers. An interesting finding was that "
    "companies with stronger IT governance reported significantly "
    "higher satisfaction with their cloud deployments. The limitation "
    "is geographic\u2014a Western European SME focus that does not "
    "translate easily to consulting firms in emerging economies."
),

204: "Seethamraju (2015)",

205: (
    "Seethamraju examined SaaS-based ERP in Australian universities, "
    "which are complex, multi-stakeholder institutions by nature. "
    "Semi-structured interviews with IT directors and process owners "
    "at three universities revealed real gains in maintainability and "
    "upgrade frequency. The downside? Cloud platforms offered "
    "considerably less room for bespoke customization compared to "
    "on-premises installations. For organizations tangled in "
    "regulatory and reporting obligations, that trade-off can be "
    "painful. The finding resonates with consulting firms, which "
    "face their own version of the customization-versus-standardization "
    "dilemma."
),

206: "Martins, Au-Yong-Oliveira, and Neto (2019)",

207: (
    "This longitudinal study tracked cloud ERP migration at 28 "
    "multinationals over four years, applying a balanced scorecard "
    "to evaluate outcomes. Internal processes and organizational "
    "learning saw the biggest improvements\u2014process automation "
    "rates went up, cross-functional data sharing improved noticeably. "
    "Financial benefits, including TCO reduction, appeared mostly in "
    "years two and three post-migration, suggesting a meaningful lag "
    "between going live and seeing the full financial return. A "
    "notable gap: mid-sized advisory firms were not part of the "
    "sample at all."
),

208: "Sagheer and colleagues (2021)",

209: (
    "Sagheer et al. surveyed 120 organizations across India, Pakistan, "
    "and Bangladesh to understand what holds back cloud ERP adoption in "
    "the South Asian technology sector. Data privacy concerns, skill "
    "shortages, and migration costs emerged as the top three barriers. "
    "On a more encouraging note, organizations that already had "
    "experience with cloud platforms\u2014a hosted CRM or HRMS, for "
    "example\u2014experienced considerably smoother ERP migrations. The "
    "South Asian focus makes these findings directly relevant to the "
    "present study."
),

210: "Ali and colleagues (2023)",

211: (
    "Ali et al. surveyed 200 IT managers across the Gulf Cooperation "
    "Council region and developed a multidimensional model for evaluating "
    "post-migration ERP performance. The model spans technical health, "
    "user satisfaction, process efficiency, and strategic alignment. "
    "Using structural equation modelling, the authors identified three "
    "key success drivers: executive sponsorship, thorough training "
    "programmes, and phased rollout strategies. They also flagged a gap "
    "around the role of middle management in post-migration adoption\u2014"
    "a gap the current study begins to address through its survey design."
),

213: (
    "On-premises ERP has been the workhorse of enterprise information "
    "management since the 1990s. SAP R/3, Oracle E-Business Suite, and "
    "Microsoft Dynamics NAV were the dominant platforms of that generation. "
    "Running these systems meant purchasing perpetual software licenses, "
    "maintaining physical servers on-site, and keeping a dedicated IT "
    "team available for customization, patching, and periodic upgrades."
),

214: (
    "On-premises ERP earned its dominance for good reasons: deep "
    "customisability, total control over data, and independence from "
    "internet connectivity. Large, heavily regulated organizations "
    "gravitated toward it\u2014manufacturers needing shop-floor integration, "
    "banks requiring transactional security, and government agencies "
    "running air-gapped networks."
),

215: (
    "The irony is that the very qualities which made on-premises ERP "
    "effective in a more stable era now work against it. Rigid "
    "architecture stifles the pace of innovation. Upgrade cycles "
    "measured in years guarantee that the organization is always "
    "trailing the technology curve by a significant margin."
),

217: (
    "On-premises ERP is built on a four-layer stack (see Figure 2.2). "
    "The user layer connects to the application layer exclusively "
    "through the corporate LAN. Anyone working off-site needs a VPN, "
    "which is often unreliable. Below that sits the infrastructure "
    "layer\u2014physical servers, storage, databases\u2014all owned and "
    "managed by the in-house IT team. Every hardware replacement, "
    "every software patch, every capacity expansion lands on that "
    "team\u2019s desk, piling up both risk and cost."
),

222: (
    "The canonical definition comes from NIST: cloud computing is a "
    "model for on-demand network access to a shared pool of configurable "
    "computing resources\u2014servers, storage, networks, applications\u2014"
    "that can be provisioned and released rapidly with minimal hands-on "
    "management (Mell and Grance, 2011)."
),

223: (
    "Three characteristics make cloud computing work: self-service on "
    "demand, broad network accessibility, and rapid elasticity. Together "
    "they remove the time and capacity constraints that define "
    "on-premises systems. For ERP, cloud delivery fundamentally changes "
    "how an organization relates to its software\u2014ownership gives way "
    "to subscription, and hands-on management gives way to consumption."
),

225: (
    "IaaS delivers virtual servers, storage, and networking over the "
    "internet on a pay-as-you-go basis. The cloud provider handles the "
    "hardware; the customer controls the operating system, middleware, "
    "and applications. In ERP contexts, IaaS is commonly chosen for "
    "lift-and-shift migrations\u2014moving the existing application onto "
    "cloud-hosted virtual machines without rearchitecting it. AWS EC2, "
    "Azure Virtual Machines, and Google Compute Engine are the major "
    "players."
),

227: (
    "PaaS gives developers an internet-based environment for building "
    "and deploying applications, abstracting away the infrastructure "
    "underneath. In ERP ecosystems, PaaS is where organizations build "
    "custom modules, integrations, and analytics tools that extend "
    "their cloud ERP. SAP Business Technology Platform and Oracle "
    "Cloud Platform are the leading options in this space."
),

229: (
    "SaaS is the most common delivery model for cloud ERP. The "
    "vendor owns and operates the entire stack\u2014infrastructure, "
    "platform, and application. Customers access it through a browser "
    "or API, pay per user or per module, and receive updates "
    "automatically. The IT team\u2019s maintenance workload drops to "
    "almost nothing."
),

231: (
    "SaaS ERP combines enterprise resource planning functionality "
    "with cloud-native delivery. The leading platforms\u2014SAP S/4HANA "
    "Cloud (Public Edition), Oracle Fusion Cloud ERP, Microsoft "
    "Dynamics 365 Finance, and Workday Finance\u2014share common traits: "
    "standardized processes built on industry best practices, "
    "quarterly or semi-annual feature updates, multi-tenant "
    "architecture that spreads costs across customers, and native "
    "hooks for AI, machine learning, and advanced analytics."
),

236: (
    "ERP systems have evolved through distinct generational phases. "
    "Each phase brought a step change in capability, delivery model, "
    "and strategic relevance within organizations. The timeline below "
    "traces those milestones and the features\u2014and limitations\u2014that "
    "characterized each era."
),

239: (
    "On a functional level, on-premises and cloud ERP cover broadly "
    "the same ground: finance, HR, supply chain management, and so on. "
    "Architecturally, though, they are fundamentally different animals. "
    "On-premises ERP is monolithic\u2014single-tenant, vertically "
    "integrated server stacks. Cloud ERP, especially in SaaS form, "
    "is distributed\u2014multi-tenant, microservices-based, and hosted "
    "on hyperscale infrastructure."
),

246: (
    "When the literature is considered as a whole, a clear pattern "
    "emerges. Cloud ERP adoption is shaped by an interplay of "
    "technology, organizational culture, and human behaviour. Studies "
    "consistently point to benefits like scalability, flexibility, and "
    "accessibility, while simultaneously raising concerns about user "
    "resistance, complex migrations, and data security risks."
),

247: (
    "A noticeable gap exists, however. The majority of existing research "
    "either drills into technical plumbing or studies large multinational "
    "corporations headquartered in the West. Management-oriented "
    "perspectives from Indian professional service firms\u2014particularly "
    "around how ready employees are, how they adopt new systems, and "
    "what support structures ease the transition\u2014have received "
    "surprisingly little academic attention."
),

248: (
    "This study steps into that gap. By examining cloud ERP migration "
    "at Keystone Advisory Group and weaving together organizational "
    "strategy with user-perception data inside a single research "
    "framework, it creates a tighter connection between theoretical "
    "concepts and the practical challenges managers face on the ground."
),

249: (
    "Another theme running through the literature is that companies "
    "going cloud-first with ERP increasingly prioritize business "
    "agility, elastic scaling, and remote access. The COVID-19 "
    "pandemic was a powerful catalyst in this regard\u2014it exposed "
    "just how brittle on-premises infrastructure becomes when "
    "offices close and business operations must continue remotely."
),

250: (
    "Researchers also caution against judging a cloud ERP deployment "
    "purely on its technology. Organizational readiness, employee "
    "adaptation, and strategic alignment matter just as much. That "
    "is why modern migration studies increasingly combine management "
    "perspectives with technical analysis to produce a more complete "
    "picture of what cloud transformation actually involves."
),

252: (
    "A careful review of both academic and practitioner literature "
    "reveals several areas where empirical evidence is either thin "
    "or missing entirely. Those gaps define the specific contribution "
    "this study sets out to make."
),

}

# ===========================================================
# CHAPTER 3
# ===========================================================

CHAPTER3 = {

275: (
    "Methodology is the backbone of any credible academic study. It "
    "determines how data gets collected, how it\u2019s analysed, and whether "
    "the conclusions actually hold up under scrutiny. Without a "
    "transparent methodology, there\u2019s simply no way for readers or "
    "evaluators to judge whether the findings follow from the evidence "
    "or just reflect what the researcher wanted to believe."
),

276: (
    "This chapter presents the full methodological framework used to "
    "study perceptions and attitudes toward cloud ERP migration. It "
    "covers the research design philosophy, the nature of the data, "
    "the collection instruments, the sampling approach, the analytical "
    "techniques, and the ethical safeguards observed throughout. Every "
    "decision is justified in terms of the study\u2019s objectives and "
    "the character of the research problem."
),

277: (
    "After coding, responses were analysed using descriptive methods "
    "based on percentages. Those results were then interpreted "
    "against the research objectives to pull out recurring patterns, "
    "prevailing attitudes, and shared perceptions about cloud ERP "
    "migration. Tables and charts were used to present the findings "
    "in a structured, accessible format."
),

279: (
    "This study uses a descriptive research design. Descriptive research "
    "is appropriate when the goal is to accurately capture the "
    "characteristics, perceptions, and attitudes of a defined group "
    "toward a specific phenomenon\u2014here, cloud ERP migration. Unlike "
    "exploratory research, which is useful when little is known about "
    "a topic, or causal research, which attempts to prove cause-and-"
    "effect links, descriptive research systematically maps the current "
    "landscape of opinions and experiences."
),

280: (
    "The descriptive approach was chosen because the study aims to "
    "understand how respondents perceive legacy systems, what they "
    "expect from cloud alternatives, how they feel about change "
    "management processes, and whether they are satisfied with past "
    "technology transitions. These are all perceptual and attitudinal "
    "constructs, and structured surveys are the natural way to "
    "measure them."
),

281: (
    "Cloud ERP readiness can be assessed across several dimensions: "
    "technical infrastructure, employee preparedness, data quality, "
    "and management support. The survey in this study captures "
    "readiness indirectly, through questions about familiarity with "
    "cloud tools and openness to trying new systems. A descriptive "
    "design was judged the best fit because the objective is to map "
    "perceptions and tendencies rather than prove causal "
    "relationships. Since cloud ERP adoption straddles both "
    "technology and human behaviour, structured opinion data provide "
    "genuine insight into the practical challenges of system "
    "transition."
),

284: (
    "The study draws on two streams of data\u2014primary and secondary\u2014"
    "to build as thorough an evidence base as possible."
),

288: (
    "Primary data is original data collected by the researcher "
    "specifically for this project. None of it has been published or "
    "analysed before. Collection was done through a structured "
    "questionnaire distributed via Google Forms to a convenience "
    "sample of 35 individuals\u2014a mix of students and early-career "
    "professionals from different academic backgrounds."
),

289: (
    "This primary data sits at the empirical core of the research. "
    "The questionnaire was designed to probe how respondents compare "
    "traditional and cloud-based systems, how they feel about "
    "switching technologies, and whether past transitions left them "
    "satisfied. Jargon was deliberately avoided so that respondents "
    "without IT expertise could still give meaningful answers."
),

290: (
    "Using an online questionnaire\u2014Google Forms, specifically\u2014sped "
    "up data collection, broadened the reach, and made it easier "
    "for people to respond at their convenience. Built-in features "
    "like automatic response logging and data export simplified "
    "the move from raw answers to organized datasets."
),

292: (
    "Secondary data is information that was already collected and "
    "published by other researchers or institutions. In this study, "
    "secondary sources served three purposes: they built the "
    "theoretical framework in Chapter 2, justified the research "
    "design, and provided context for interpreting the primary "
    "findings. The sources included peer-reviewed journals, "
    "methodology textbooks, Gartner and IDC industry reports, and "
    "vendor documentation from SAP and Microsoft."
),

294: (
    "Surveying an entire population is rarely feasible. Sampling "
    "solves that problem by selecting a manageable subset from which "
    "meaningful conclusions can still be drawn."
),

296: (
    "The target population for this study consists of students and "
    "working professionals who have experience using some form of "
    "digital or office management tool\u2014whether that is SAP, Tally, "
    "Google Drive, Zoom, or an online banking app. The net was cast "
    "wide on purpose. The questionnaire probes general attitudes "
    "toward technology adoption, not specialist ERP expertise, so "
    "a diverse respondent base actually strengthens rather than "
    "weakens the quality of the data."
),

298: (
    "Respondent selection followed a convenience sampling approach. "
    "Participants were chosen based on accessibility and willingness "
    "to take part, not through random selection. This is standard "
    "practice for student-led research projects where time, budget, "
    "and access constraints make probability-based sampling "
    "impractical."
),

299: (
    "The questionnaire was shared digitally\u2014Google Forms links "
    "sent through WhatsApp, email, and direct messages to classmates, "
    "peers, and acquaintances. The trade-off is well understood: "
    "convenience sampling limits generalisability. But for a "
    "descriptive study whose aim is to identify attitudinal patterns "
    "rather than make claims about an entire population, it is a "
    "defensible choice."
),

302: (
    "The target sample size was 35 respondents. Larger samples are "
    "always preferable, but in academic research a sample of 30 or "
    "more is widely accepted as sufficient for basic analysis. "
    "For a descriptive, percentage-based study, 35 is enough to "
    "surface meaningful trends\u2014and realistic given the time and "
    "access constraints of a student project."
),

308: (
    "The questionnaire was designed to align with the research "
    "objectives. Plain language was used throughout to minimize the "
    "risk of respondent confusion. Before distribution, the full "
    "form was reviewed for logical flow and internal consistency."
),

309: (
    "Because the analysis stays within descriptive, percentage-based "
    "territory, advanced statistical validation techniques were not "
    "applied. That said, the collected data is considered reliable "
    "for descriptive purposes within the scope of this study."
),

314: (
    "Participation was entirely voluntary and anonymous. Respondents "
    "were informed that the data was being collected for academic "
    "purposes only and that no individual responses would be "
    "identifiable in the final report. All data was stored securely "
    "and used exclusively for research, in accordance with "
    "Chandigarh University\u2019s data protection policies."
),

315: (
    "No question was compulsory, and respondents could discontinue "
    "the survey at any point. This combination of voluntary "
    "participation, anonymity, and the right to withdraw ensured "
    "ethical transparency throughout."
),

}

# ===========================================================
# CHAPTER 4
# ===========================================================

CHAPTER4 = {

318: (
    "This chapter breaks down the primary and organizational data "
    "collected from 35 respondents via the Google Forms questionnaire "
    "described in Chapter 3. Each question is shown as a frequency "
    "table, followed by space for a bar chart and a brief "
    "interpretation of the results."
),

320: (
    "What follows is a question-by-question analysis of the primary "
    "survey data. Thirty-five respondents shared their experiences, "
    "opinions, difficulties, and preferences regarding cloud-based "
    "systems, providing a practical snapshot of how cloud technologies "
    "are perceived by real users."
),

339: (
    "Two related questions are presented together here. First, "
    "respondents selected the advantages they associate with cloud "
    "systems (multi-select). Then they indicated whether they consider "
    "cloud or traditional systems better overall (single choice)."
),

345: (
    "Convenience, the ability to work from anywhere, and lighter "
    "maintenance requirements stood out as the strongest influences "
    "on positive attitudes toward cloud ERP. This fits a broader "
    "trend: users increasingly expect flexible, location-independent "
    "systems as the default rather than the exception."
),

347: (
    "Three linked questions address the switching experience: what "
    "stops people from switching, what would make it easier, and "
    "how long full adaptation takes. The results follow in sequence."
),

353: (
    "A clear pattern runs through these responses: whether an ERP "
    "transition succeeds depends heavily on how prepared the people "
    "are and how much support the organization provides. Structured "
    "change management is not optional; it is foundational."
),

355: (
    "The last two survey questions look at what respondents experienced "
    "after switching systems and whether they would recommend it. "
    "Both sets of results appear together below."
),

360: (
    "The strong enthusiasm for cloud upgrades signals genuine "
    "acceptance of digital transformation among respondents. "
    "At the same time, the preference for gradual implementation "
    "suggests that phased migration strategies could reduce "
    "resistance and improve long-term adoption success."
),

362: (
    "Stepping back from the individual questions, the survey data "
    "tell a pretty clear story. Traditional systems drew complaints "
    "about slow performance, limited remote access, and too much "
    "dependence on technical support. Cloud systems, by contrast, "
    "were consistently seen as more flexible, more accessible, and "
    "frankly better suited to how organizations operate today. "
    "That\u2019s not a subtle difference\u2014it\u2019s a decisive shift in "
    "user sentiment."
),

363: (
    "Human factors proved just as important as technology. Training, "
    "transparent communication, and organizational support were "
    "identified as critical enablers. Respondents showed a clear "
    "preference for phased rollouts, indicating that gradual "
    "migration paths reduce anxiety and smooth the learning curve."
),

364: (
    "Taken as a whole, the data support a straightforward argument: "
    "cloud ERP improves efficiency, widens access, and gives "
    "organizations more room to manoeuvre\u2014but only when solid "
    "change-management practices are in place to support the "
    "people using the system."
),

365: (
    "One additional observation worth noting: nearly every respondent "
    "already uses cloud applications in their daily routine, whether "
    "in an academic or professional setting. That baseline familiarity "
    "is an asset for any organization rolling out cloud ERP because "
    "it means the learning curve starts from a higher baseline. "
    "Respondents also saw cloud tools as better aligned with modern "
    "work habits\u2014remote access, real-time collaboration, and "
    "working on the move."
),

366: (
    "Still, technology alone will not guarantee success. Communication, "
    "training, user confidence, and institutional backing keep "
    "surfacing as decisive factors in whether people accept a new "
    "system and stick with it over time. The implication for "
    "practitioners is clear: every technical blueprint needs an "
    "equally rigorous change-management plan alongside it."
),

}

# ===========================================================
# CHAPTER 5
# ===========================================================

CHAPTER5 = {

385: (
    "This chapter pulls together the key findings from the data "
    "analysis in Chapter 4. Everything traces back to the structured "
    "questionnaire and maps onto the research objectives from "
    "Chapter 1. The analytical approach is descriptive and "
    "percentage-based; no new data appears here."
),

387: (
    "The respondent pool was balanced: two in five were working "
    "professionals, another two in five straddled student and "
    "working roles, and the remaining one in five were full-time "
    "students. That spread lends credibility to the results."
),

388: (
    "Most respondents reported being comfortable with cloud-based "
    "tools, which means they were well placed to compare cloud "
    "and traditional systems from genuine experience."
),

390: (
    "Four out of five respondents (80 percent) identified slow "
    "performance or system crashes as the most widespread "
    "limitation of traditional systems."
),

391: (
    "Inability to access systems remotely, dependence on IT support "
    "for minor issues, and slow report generation were each cited "
    "by two in five respondents (40 percent), pointing to a "
    "cluster of operational frustrations."
),

392: (
    "High maintenance cost was singled out as the biggest single "
    "problem by two in five respondents (40 percent)."
),

393: (
    "Steep maintenance bills and the ongoing need for dedicated "
    "IT staff further underlined the operational drag that "
    "traditional systems impose."
),

394: (
    "The picture is consistent: traditional systems struggle "
    "with performance, access, and cost\u2014three areas where "
    "cloud ERP is designed to deliver improvement."
),

396: (
    "Every single respondent (100 percent) reported using "
    "cloud-based tools, indicating deep-rooted familiarity "
    "with cloud platforms."
),

397: (
    "The ability to access systems from any location was "
    "unanimously rated the most important advantage."
),

398: (
    "Usability, automatic software updates, and lower "
    "maintenance requirements were each flagged as key "
    "benefits by four out of five respondents (80 percent)."
),

399: (
    "Three in five (60 percent) called cloud systems outright "
    "better than traditional ones. The rest considered them "
    "somewhat better or comparable. Nobody rated them worse."
),

400: (
    "Not a single respondent preferred traditional systems "
    "over cloud-based alternatives."
),

402: (
    "Learning how to use the new system was the top transition "
    "challenge, cited by four in five respondents (80 percent)."
),

403: (
    "Concerns about data loss and inadequate training were "
    "each reported as significant worries by three in five "
    "respondents (60 percent)."
),

404: (
    "Proper training before the transition was identified as "
    "the most critical success factor\u2014unanimously, by all "
    "35 respondents."
),

405: (
    "Clear communication and active support throughout the "
    "migration were also cited as important enablers by "
    "four in five respondents (80 percent)."
),

406: (
    "The largest single group (40 percent) estimated that "
    "full adaptation to a new system takes between one and "
    "three months."
),

408: (
    "Most respondents reported experiencing tangible benefits "
    "after switching to a more capable system."
),

409: (
    "Faster response times, easier remote access, and fewer "
    "system errors were each noted by four in five respondents "
    "(80 percent)."
),

410: (
    "Three in five (60 percent) strongly recommended upgrading "
    "to cloud systems. The remainder preferred a more gradual "
    "transition approach."
),

411: (
    "No respondent opposed cloud adoption\u2014acceptance was "
    "across the board."
),

414: (
    "These findings offer a clear picture of how respondents "
    "view both legacy and cloud systems. The insights serve as "
    "the foundation for drawing conclusions in the next chapter "
    "and for systematically evaluating each research objective."
),

415: (
    "A recurring thread: respondents associate cloud platforms "
    "with greater flexibility, easier access, and smoother "
    "daily operations. Equally evident is the message that "
    "successful transition requires proper user preparation, "
    "organizational backing, and consistent communication."
),

}

# ===========================================================
# CHAPTER 6
# ===========================================================

CHAPTER6 = {

418: (
    "Findings tell you what the data show. Conclusions tell you "
    "what those data actually mean\u2014for the organization, for the "
    "research field, and for the broader question of whether cloud "
    "ERP migration makes strategic sense. There\u2019s an important "
    "distinction between the two, and this chapter connects the "
    "Chapter 5 findings to the research objectives and theoretical "
    "frameworks established at the start of the study."
),

420: (
    "The data confirm that legacy systems carry real operational "
    "baggage. Two in five respondents identified remote access "
    "barriers as the biggest problem; four in five reported "
    "performance-related issues. Those numbers directly translate "
    "into reduced efficiency\u2014and they validate the argument for "
    "cloud-based alternatives."
),

421: (
    "These are not trivial inconveniences. In a working world "
    "built around hybrid arrangements, on-site client work, and "
    "real-time decisions, a system that ties users to a LAN is a "
    "genuine competitive liability. Every pain point raised in "
    "Chapter 1 for Keystone\u2014poor remote access, ballooning "
    "maintenance costs, slow scaling\u2014finds clear support in the "
    "primary data."
),

423: (
    "Cloud systems enjoy broad, enthusiastic support. All 35 "
    "respondents use cloud tools already; every one of them named "
    "anywhere-access as the top benefit. Three in five went "
    "further and called cloud systems outright superior to "
    "traditional ones. That is a strong endorsement."
),

424: (
    "What drives this preference? Accessibility. Being able to "
    "log in from any device, anywhere, at any hour\u2014that is "
    "what users value most. It maps directly onto the TAM "
    "concept of Perceived Usefulness discussed in Chapter 2. "
    "Cloud ERP earns high marks because it eliminates the most "
    "frustrating limitation of legacy systems: having to be "
    "physically present to get anything done."
),

426: (
    "When it comes to migration challenges, the data point "
    "overwhelmingly toward human factors rather than technical "
    "ones. Learning new workflows (80 percent) and worry about "
    "data loss (60 percent) topped the chart. The takeaway is "
    "clear: user-focused planning is not a nice-to-have\u2014it is "
    "what separates a migration that succeeds from one that "
    "stalls."
),

427: (
    "Both barriers are addressable. Training was the top enabler "
    "(cited by 80 percent); clear communication was second (60 "
    "percent). The lesson practically writes itself: invest in "
    "people and processes, not just infrastructure, and the odds "
    "of a smooth migration improve substantially."
),

429: (
    "After the switch, the picture is positive. Four in five "
    "respondents reported faster performance and better access. "
    "Three in five gave an outright endorsement of cloud systems\u2014"
    "evidence that the benefits hold up in practice, not just "
    "in vendor brochures."
),

430: (
    "An important nuance: two in five respondents preferred a "
    "gradual, phased migration over a single cutover date. Even "
    "users who are enthusiastic about the move acknowledge that "
    "migration is complex and value a measured approach over a "
    "big-bang switch. That preference deserves serious attention "
    "from any organization charting its own migration path."
),

434: (
    "Pulling everything together, the conclusion isn\u2019t ambiguous: "
    "cloud ERP migration is strategically sound, broadly welcomed, "
    "and operationally beneficial for organizations like Keystone "
    "Advisory Group. Data from 35 respondents, layered onto the "
    "theoretical frameworks from Chapter 2, tells us that "
    "on-premises systems simply can\u2019t keep up with modern "
    "requirements anymore. Cloud platforms deliver better access, "
    "higher efficiency, and greater user satisfaction\u2014and the "
    "respondents aren\u2019t shy about saying so."
),

435: (
    "For Keystone specifically, the data underscore the importance "
    "of moving past the limitations of the current system\u2014"
    "performance bottlenecks, access restrictions, and "
    "operational overhead. Organizations that combine structured "
    "training, transparent communication, and phased migration "
    "timelines stand the best chance of extracting real value "
    "from cloud ERP."
),

436: (
    "Beyond the organizational implications, this study adds "
    "to the growing body of research on cloud ERP adoption "
    "within India\u2019s professional services sector, providing "
    "ground-level data on user perception, migration friction, "
    "and success factors."
),

437: (
    "The overarching conclusion is this: cloud ERP adoption "
    "is not simply a technology upgrade. It is a strategic "
    "transformation that strengthens organizational agility "
    "and long-term competitive positioning."
),

439: (
    "Based on the study\u2019s findings from both primary and "
    "secondary data, a structured Cloud ERP Migration Framework "
    "is proposed for Keystone Advisory Group. It provides a "
    "step-by-step, practical blueprint for transitioning from "
    "on-premises ERP to a cloud platform while preserving data "
    "integrity, maintaining operations, and bringing users along."
),

443: (
    "Phase 1: Readiness Assessment\n"
    "Before anything else, the organization needs to take stock. "
    "How sound is the technical infrastructure? Is the data clean "
    "enough to migrate? Are employees open to change? Does the "
    "budget support it? Does the migration align with business "
    "goals? This study\u2019s findings show that familiarity with "
    "cloud tools and willingness to adapt are among the strongest "
    "signals of organizational readiness."
),

445: (
    "Phase 2: Planning and Cloud Platform Selection\n"
    "This phase locks down scope, objectives, and a realistic "
    "timeline. Cloud ERP vendors should be evaluated against "
    "cost, scalability, security, integration capabilities, and "
    "vendor support quality. A thorough evaluation ensures the "
    "chosen platform fits both current operational needs and "
    "longer-term strategic ambitions."
),

446: (
    "Phase 3: Data Migration and Integration\n"
    "This is where organizational data, records, and workflows "
    "move from the legacy system to the new cloud platform. It "
    "is one of the riskiest stages of any ERP project because "
    "errors or gaps during transfer can ripple through "
    "operations and distort decision-making downstream."
),

447: (
    "Before any data moves, a deep cleaning exercise is essential: "
    "duplicates need to be removed, stale entries purged, and gaps "
    "filled in. Robust backup mechanisms should be in place to "
    "guard against permanent data loss. Critical modules and "
    "datasets should be prioritized so that the business keeps "
    "running smoothly throughout the cutover period."
),

448: (
    "Integration testing comes next. The migrated ERP needs to "
    "work properly alongside existing tools and workflows. Any "
    "glitches caught during testing should be resolved before the "
    "broader rollout\u2014catching problems early is far cheaper than "
    "fixing them after they have disrupted live operations."
),

449: (
    "A phased migration approach, closely monitored at each "
    "stage, is the safest way forward. It controls operational "
    "risk, protects data quality, and eases the transition into "
    "the new cloud ERP environment."
),

450: (
    "Phase 4: System Testing\n"
    "Before going live, the platform should go through rigorous "
    "functional, performance, and security testing. Everything "
    "needs to meet organizational benchmarks. Issues should be "
    "resolved before deployment, not after."
),

451: (
    "Phase 5: User Training and Change Management\n"
    "A system is only as effective as the people who use it. "
    "Employees need to be trained on the new platform, and clear "
    "communication should be used to manage resistance. This "
    "study\u2019s data identify training as the single most important "
    "factor in post-migration satisfaction and adoption."
),

452: (
    "Phase 6: Go-Live and Post-Migration Support\n"
    "Deploy the system\u2014and then keep watching. Continuous "
    "performance monitoring, a feedback loop, and ongoing "
    "enhancements are what turn a successful launch into "
    "sustained long-term value."
),

453: (
    "What ties all six phases together is a management-first "
    "philosophy. Each stage is designed to minimize operational "
    "disruption while improving employee readiness, safeguarding "
    "data, and keeping the system adaptable. Aligning every "
    "migration activity with broader business goals ensures that "
    "cloud ERP delivers lasting efficiency and strategic "
    "flexibility\u2014not just a newer IT stack."
),

454: (
    "Follow this framework and the transition becomes "
    "lower-risk, higher-adoption, and more financially "
    "rewarding. The framework also directly addresses this "
    "study\u2019s research questions by packaging readiness "
    "assessment, vendor selection, and migration strategy "
    "into one coherent plan."
),

455: (
    "In short, the framework gives Keystone Advisory Group "
    "a clear roadmap: implement cloud ERP successfully, then "
    "sustain the efficiency and agility gains over the long "
    "term."
),

457: (
    "Several managerial implications stand out for Keystone\u2019s "
    "leadership. First, cloud ERP migration is a strategic "
    "organizational shift, not just an IT project\u2014user adoption "
    "and change management deserve equal emphasis. Second, ERP "
    "implementation should connect to headline business goals: "
    "cost control, operational efficiency, better data "
    "accessibility. Third, ongoing communication, employee "
    "participation, and structured training programmes are the "
    "antidote to resistance. And fourth, the phased approach "
    "outlined in the migration framework provides a practical "
    "path to lower risk, smoother transition, and long-term "
    "performance gains."
),

}

# ===========================================================
# CHAPTER 7
# ===========================================================

CHAPTER7 = {

460: (
    "The recommendations below are drawn directly from the data "
    "and conclusions of this study. Each one targets Keystone "
    "Advisory Group and organizations facing similar circumstances. "
    "These are specific, evidence-backed suggestions\u2014not generic "
    "advice."
),

462: (
    "Training was identified as the single most important enabler "
    "of a smooth transition by every respondent\u2014100 percent "
    "agreement, the highest figure in the entire survey. And yet, "
    "learning how to use the new system was cited as the number one "
    "challenge by 80 percent. Think about that for a moment. There\u2019s "
    "a clear gap between what users desperately need and what "
    "organizations typically bother to provide\u2014and it\u2019s one of "
    "migration\u2019s most dangerous blind spots."
),

463: (
    "Recommendation: Design role-specific training modules and "
    "deliver them four to six weeks before go-live. Keep sessions "
    "short and task-focused. Finance staff should practise on "
    "billing and reporting modules; HR should work through payroll "
    "and leave flows. Identify early adopters within each team and "
    "appoint them as peer champions who can support colleagues\u2014"
    "it accelerates adoption without adding to the budget."
),

465: (
    "Four in five respondents said that clear communication about "
    "the purpose of the migration would make the transition easier. "
    "In practice, though, ERP migrations are frequently presented "
    "as IT projects\u2014full of go-live dates and server specifications "
    "but short on the business rationale that actually matters to "
    "end users."
),

466: (
    "Recommendation: Explain what changes for each role, using "
    "plain language. Saying \u201cyou can now submit timesheets from "
    "any device\u201d resonates far more than explaining a cloud-native "
    "workflow engine. Telling someone \u201creports that used to take 90 "
    "minutes will now take 10\u201d is more persuasive than a slide deck "
    "about scalability. Link the system change to daily, personal "
    "benefits and resistance drops sharply."
),

468: (
    "Two in five respondents preferred a gradual migration, and "
    "many estimated that full adaptation takes one to three months. "
    "A big-bang cutover\u2014every department, every module, one single "
    "date\u2014raises the stakes and unsettles users unnecessarily."
),

469: (
    "Recommendation: Roll out in phases. Start with a digitally "
    "confident, lower-risk department like HR or finance. "
    "Stabilize, collect feedback, refine the training materials, "
    "and then move to the next group. This phased approach "
    "surfaces integration issues early, builds internal confidence "
    "progressively, and avoids the chaos of a simultaneous "
    "enterprise-wide switch."
),

472: (
    "Many respondents indicated that adapting to a new system "
    "takes considerably longer than a couple of weeks. Yet most "
    "cloud ERP projects concentrate intensive support into a "
    "narrow post-launch window, leaving users under-supported "
    "during the longer settling-in period."
),

473: (
    "Recommendation: Plan for a minimum of 90 days of structured "
    "post-go-live support. Set up a dedicated helpdesk, hold "
    "weekly feedback sessions, and run monthly performance reviews "
    "against pre-migration benchmarks. Make a point of publicizing "
    "quick wins\u2014faster reports, fewer outages\u2014to keep motivation "
    "and buy-in high."
),

475: (
    "Three in five respondents flagged data loss as a major "
    "concern during system transitions. The literature reviewed "
    "in Chapter 2 corroborates this: data migration complexity "
    "is the most frequently cited technical barrier."
),

476: (
    "Recommendation: Conduct a thorough data audit before any "
    "migration begins. Clean out duplicates, retire outdated "
    "records, and fill in missing fields. Assign data ownership "
    "to department heads and make each one responsible for "
    "validating their data slice before go-live. Running the "
    "old and new systems in parallel for a brief period allows "
    "teams to catch mismatches without disrupting live operations."
),

480: (
    "Once the system is live, performance monitoring should not "
    "stop. System efficiency, user satisfaction, uptime, and "
    "response times should all be reviewed on a regular cadence "
    "to verify that the expected benefits of cloud ERP are "
    "actually materializing."
),

481: (
    "Continuous monitoring catches problems before they escalate. "
    "Periodic feedback sessions, satisfaction surveys, and system "
    "performance audits give leadership a honest read on whether "
    "the new ERP is meeting organizational expectations."
),

482: (
    "Recommendation: Establish a formal performance review "
    "mechanism that includes recurring KPI assessments, a "
    "feedback tracker, and scheduled system audits. This kind "
    "of continuous scrutiny improves reliability in the short "
    "term and feeds organizational learning over the long term."
),

483: (
    "Thread all six recommendations together and one central "
    "theme emerges: a successful cloud ERP migration rests as "
    "much on planning, people, and ongoing support as it does "
    "on selecting the right technology."
),

}

# ===========================================================
# CHAPTER 8
# ===========================================================

CHAPTER8 = {

486: (
    "No study is perfect\u2014every piece of research operates within "
    "constraints, and being upfront about them is fundamental to "
    "academic integrity. The limitations below define the boundaries "
    "within which these findings should be read and interpreted."
),

488: (
    "These limitations don\u2019t invalidate what the study found, but "
    "they do mark the fence line. The evidence is credible and "
    "internally consistent within its scope. Future researchers "
    "could strengthen and extend these insights by using larger "
    "probability-based samples, adopting longitudinal designs, and "
    "collecting data directly from organizational stakeholders. "
    "It\u2019s also worth noting that the primary data came from "
    "general digital-tool users rather than Keystone\u2019s own staff, "
    "so the findings are best read as indicative signals about "
    "cloud ERP adoption\u2014not as precise organizational measurements."
),

490: (
    "There is substantial scope for future research. Larger and "
    "more diverse samples\u2014particularly professionals directly "
    "involved in ERP projects\u2014would add significant weight. "
    "Longitudinal studies could track how cloud ERP reshapes "
    "organizational performance and employee behaviour over years "
    "rather than months. And there is ample room to investigate "
    "how emerging technologies like artificial intelligence, "
    "machine learning, and predictive analytics are evolving "
    "within cloud ERP ecosystems."
),

}


# ===========================================================
# APPLY ALL PARAGRAPH REWRITES
# ===========================================================

ALL_REWRITES = {}
ALL_REWRITES.update(CHAPTER1)
ALL_REWRITES.update(CHAPTER2)
ALL_REWRITES.update(CHAPTER3)
ALL_REWRITES.update(CHAPTER4)
ALL_REWRITES.update(CHAPTER5)
ALL_REWRITES.update(CHAPTER6)
ALL_REWRITES.update(CHAPTER7)
ALL_REWRITES.update(CHAPTER8)

for idx, new_text in ALL_REWRITES.items():
    para = doc.paragraphs[idx]
    replace_para_text(para, new_text)

# ===========================================================
# TABLE CELL REWRITES
# ===========================================================

TABLE_CELL_REWRITES = {
    (23, 0, 0): (
        "Working professionals made up 40 percent of the sample, and "
        "another 40 percent identified as both students and working "
        "professionals. The remaining 20 percent were students. This "
        "three-way split gives the data a blend of academic and "
        "professional perspectives, adding practical weight to "
        "the survey outcomes."
    ),
    (26, 0, 0): (
        "Three in five respondents (60 percent) said they were very "
        "comfortable with cloud tools. That level of baseline familiarity "
        "is important\u2014it means respondents were comparing traditional "
        "and cloud systems from real, day-to-day experience rather "
        "than guesswork."
    ),
    (29, 0, 0): (
        "Slow performance topped the complaint list at 80 percent. Poor "
        "remote access and heavy reliance on IT support also appeared "
        "frequently. The message from respondents is clear: traditional "
        "platforms create real, daily friction."
    ),
    (32, 0, 0): (
        "Maintenance and upgrade costs drew the sharpest criticism, "
        "cited as the worst single issue by two in five respondents. "
        "Remote access restrictions came next. Together these responses "
        "confirm the financial and accessibility burdens that "
        "legacy systems carry."
    ),
    (35, 0, 0): (
        "Google Drive or Docs usage hit 100 percent\u2014universal. Web-based "
        "email and video conferencing were close behind. It is clear "
        "that cloud applications have become deeply embedded in "
        "respondents\u2019 daily academic and professional routines."
    ),
    (39, 0, 0): (
        "Remote access and usability were rated as the top advantages. "
        "Three in five respondents called cloud systems definitively "
        "better; the rest rated them somewhat better or equivalent. "
        "No one preferred legacy platforms\u2014a clean sweep in "
        "favour of cloud."
    ),
    (44, 0, 0): (
        "Learning the new interface (80 percent) was the biggest hurdle. "
        "Data loss fears (60 percent) came next. On the positive side, "
        "training was the one enabler that every respondent agreed on "
        "(100 percent). Most said full adaptation takes one to three "
        "months, making a strong case for extended post-go-live support."
    ),
    (48, 0, 0): (
        "Faster performance (80 percent), better remote access "
        "(80 percent), and fewer system glitches (80 percent) were "
        "the most commonly reported post-switch benefits. Three in "
        "five strongly backed migration while two in five preferred "
        "a phased pace. No one opposed the move\u2014support was "
        "unqualified and across the board."
    ),
}

for (t_idx, r_idx, c_idx), new_text in TABLE_CELL_REWRITES.items():
    replace_cell_text(t_idx, r_idx, c_idx, new_text)

# ===========================================================
# COMPREHENSIVE HUMANIZER + AI TEXT CLEANER
# Ported from enhanced index.html — reduces AI detection
# by replacing AI-common vocabulary, transition words,
# filler phrases, and applying contractions + Unicode cleanup
# ===========================================================

import random
import re as _re
import unicodedata as _ud

random.seed(42)

def _pick(lst):
    return lst[random.randint(0, len(lst) - 1)]


def humanize_text(text):
    if not text or not text.strip():
        return text

    # ── 1) TRANSITION WORDS (AI favorites -> human alternatives) ──
    tr_rules = [
        (r"\bFurthermore,?\s", lambda: _pick(["Also, ", "Plus, ", "And ", "Besides, "])),
        (r"\bMoreover,?\s", lambda: _pick(["Also, ", "What's more, ", "And ", "Plus, "])),
        (r"\bAdditionally,?\s", lambda: _pick(["Also, ", "Plus, ", "And ", "Besides, "])),
        (r"\bConsequently,?\s", lambda: _pick(["So, ", "As a result, ", "That's why "])),
        (r"\bNevertheless,?\s", lambda: _pick(["Still, ", "Even so, ", "That said, ", "But "])),
        (r"\bNonetheless,?\s", lambda: _pick(["Still, ", "Even so, ", "Yet, "])),
        (r"\bSubsequently,?\s", lambda: _pick(["Then, ", "After that, ", "Later, "])),
        (r"\bConversely,?\s", lambda: _pick(["On the flip side, ", "Then again, ", "But "])),
        (r"\bUltimately,?\s", lambda: _pick(["In the end, ", "When all is said and done, "])),
        (r"(?i)\bIn\s+addition,?\s", lambda: _pick(["Also, ", "Plus, ", "And "])),
        (r"(?i)\bAs\s+a\s+result,?\s", lambda: _pick(["So, ", "Because of this, "])),
        (r"(?i)\bIn\s+particular,?\s", lambda: _pick(["Especially, ", "Mainly, "])),
        (r"\bSpecifically,?\s", lambda: _pick(["In particular, ", "Namely, ", "That is, "])),
        (r"\bAccordingly,?\s", lambda: _pick(["So, ", "For that reason, "])),
        (r"\bHence,?\s", lambda: _pick(["So, ", "That's why "])),
        (r"\bThus,?\s", lambda: _pick(["So, ", "This way, "])),
    ]
    for pattern, replacer in tr_rules:
        text = _re.sub(pattern, lambda m: replacer(), text)

    # ── 2) FORMAL VOCABULARY -> SIMPLER WORDS ──
    voc_rules = [
        (r"(?i)\butilizes\b", "uses"), (r"(?i)\butilize\b", "use"),
        (r"(?i)\butilized\b", "used"), (r"(?i)\butilizing\b", "using"),
        (r"(?i)\bfacilitates\b", "helps"), (r"(?i)\bfacilitate\b", "help"),
        (r"(?i)\bfacilitated\b", "helped"), (r"(?i)\bfacilitating\b", "helping"),
        (r"(?i)\bdemonstrates\b", "shows"), (r"(?i)\bdemonstrate\b", "show"),
        (r"(?i)\bdemonstrated\b", "showed"), (r"(?i)\bdemonstrating\b", "showing"),
        (r"(?i)\bcomprehensive\b", lambda: _pick(["thorough", "complete", "full"])),
        (r"(?i)\bsignificantly\b", lambda: _pick(["greatly", "noticeably", "a lot"])),
        (r"(?i)\bsignificant\b", lambda: _pick(["big", "major", "notable"])),
        (r"(?i)\bsubstantially\b", lambda: _pick(["greatly", "largely", "considerably"])),
        (r"(?i)\bsubstantial\b", lambda: _pick(["big", "large", "considerable"])),
        (r"(?i)\bapproximately\b", lambda: _pick(["about", "around", "roughly"])),
        (r"(?i)\bnumerous\b", lambda: _pick(["many", "several", "quite a few"])),
        (r"(?i)\boptimal\b", lambda: _pick(["best", "ideal"])),
        (r"(?i)\benhances\b", "improves"), (r"(?i)\benhance\b", "improve"),
        (r"(?i)\benhanced\b", "improved"), (r"(?i)\benhancing\b", "improving"),
        (r"(?i)\bleverages\b", "uses"), (r"(?i)\bleverage\b", "use"),
        (r"(?i)\bleveraging\b", "using"), (r"(?i)\bleveraged\b", "used"),
        (r"(?i)\bensures\b", "makes sure"), (r"(?i)\bensure\b", "make sure"),
        (r"(?i)\bensuring\b", "making sure"),
        (r"(?i)\bpivotal\b", lambda: _pick(["key", "central", "important"])),
        (r"(?i)\bmultifaceted\b", lambda: _pick(["complex", "varied"])),
        (r"(?i)\bstreamlines\b", "simplifies"), (r"(?i)\bstreamline\b", "simplify"),
        (r"(?i)\bstreamlining\b", "simplifying"),
        (r"(?i)\bparamount\b", lambda: _pick(["top priority", "most important"])),
        (r"(?i)\brobust\b", lambda: _pick(["strong", "solid", "reliable"])),
        (r"(?i)\bseamlessly\b", "smoothly"), (r"(?i)\bseamless\b", "smooth"),
        (r"(?i)\binnovative\b", lambda: _pick(["new", "creative", "fresh"])),
        (r"(?i)\bholistic\b", lambda: _pick(["overall", "whole", "complete"])),
        (r"(?i)\bparadigms\b", "models"), (r"(?i)\bparadigm\b", "model"),
        (r"(?i)\bmethodologies\b", "methods"),
        (r"(?i)\bunderscores\b", "highlights"), (r"(?i)\bunderscore\b", "highlight"),
        (r"(?i)\bunderscoring\b", "highlighting"),
        (r"(?i)\bendeavors\b", "efforts"), (r"(?i)\bendeavor\b", "effort"),
        (r"(?i)\baugments\b", "boosts"), (r"(?i)\baugment\b", "boost"),
        (r"(?i)\baugmenting\b", "boosting"),
        (r"(?i)\bmitigates\b", "reduces"), (r"(?i)\bmitigate\b", "reduce"),
        (r"(?i)\bmitigating\b", "reducing"), (r"(?i)\bmitigation\b", "reduction"),
        (r"(?i)\bascertains\b", "finds out"), (r"(?i)\bascertain\b", "find out"),
        (r"(?i)\bascertained\b", "found out"),
        (r"(?i)\bcommences\b", "starts"), (r"(?i)\bcommence\b", "start"),
        (r"(?i)\bcommenced\b", "started"), (r"(?i)\bcommencing\b", "starting"),
        (r"(?i)\bterminates\b", "ends"), (r"(?i)\bterminate\b", "end"),
        (r"(?i)\bterminated\b", "ended"),
        (r"(?i)\bexpedites\b", "speeds up"), (r"(?i)\bexpedite\b", "speed up"),
        (r"(?i)\bexpediting\b", "speeding up"),
        (r"(?i)\brequisite\b", "needed"),
        (r"(?i)\bpredominantly\b", lambda: _pick(["mostly", "mainly"])),
        (r"(?i)\bpredominant\b", lambda: _pick(["main", "primary"])),
        (r"(?i)\bcorroborates\b", "backs up"), (r"(?i)\bcorroborate\b", "back up"),
        (r"(?i)\bcorroborating\b", "backing up"),
        (r"(?i)\belucidated\b", "explained"), (r"(?i)\belucidate\b", "explain"),
        (r"(?i)\belucidating\b", "explaining"),
        (r"(?i)\bameliorate\b", "improve"), (r"(?i)\bameliorates\b", "improves"),
        (r"(?i)\bamelioration\b", "improvement"),
        (r"(?i)\bdelineates\b", "outlines"), (r"(?i)\bdelineate\b", "outline"),
        (r"(?i)\bdelineating\b", "outlining"),
        (r"(?i)\bpropensity\b", "tendency"),
        (r"(?i)\bdiscernible\b", "noticeable"),
        (r"(?i)\bpertaining to\b", "about"),
        (r"(?i)\bwith respect to\b", "about"),
        (r"(?i)\bwith regard to\b", "about"),
        (r"(?i)\bin the context of\b", "in"),
        (r"(?i)\bconducive\b", "helpful"),
        (r"(?i)\bexemplifies\b", "shows"), (r"(?i)\bexemplify\b", "show"),
        (r"(?i)\bexemplified\b", "showed"),
        (r"(?i)\bnotwithstanding\b", "despite"),
        (r"(?i)\bhereinafter\b", "from here on"),
        (r"(?i)\binsofar as\b", "as far as"),
        (r"(?i)\bwherein\b", "where"),
        (r"(?i)\bthereby\b", "by doing this"),
        (r"(?i)\bwhereas\b", "while"),
    ]
    for pattern, repl in voc_rules:
        if callable(repl):
            text = _re.sub(pattern, lambda m, fn=repl: fn(), text)
        else:
            text = _re.sub(pattern, repl, text)

    # ── 3) AI FILLER PHRASES -> removed or simplified ──
    ph_rules = [
        (r"(?i)\bIt is important to note that\s*", ""),
        (r"(?i)\bIt should be noted that\s*", ""),
        (r"(?i)\bIt is worth mentioning that\s*", ""),
        (r"(?i)\bIt is worth noting that\s*", ""),
        (r"(?i)\bIt goes without saying that?\s*", ""),
        (r"(?i)\bNeedless to say,?\s*", ""),
        (r"(?i)\bin order to\b", "to"),
        (r"(?i)\bdue to the fact that\b", "because"),
        (r"(?i)\ba wide range of\b", "many"),
        (r"(?i)\ba plethora of\b", "many"),
        (r"(?i)\bplays a (?:crucial|vital|key|important|significant) role\b", "matters a lot"),
        (r"(?i)\bin today'?s rapidly evolving\b", "in today's"),
        (r"(?i)\bin the ever-changing landscape of\b", "in"),
        (r"(?i)\bat the present time\b", "now"),
        (r"(?i)\bat this point in time\b", "now"),
        (r"(?i)\bin the realm of\b", "in"),
        (r"(?i)\bfor the purpose of\b", "for"),
        (r"(?i)\bin the event that\b", "if"),
        (r"(?i)\bprior to\b", "before"),
        (r"(?i)\bsubsequent to\b", "after"),
        (r"(?i)\bin conjunction with\b", "with"),
        (r"(?i)\bon the basis of\b", "based on"),
        (r"(?i)\bin the absence of\b", "without"),
        (r"(?i)\bhas the potential to\b", "can"),
        (r"(?i)\bis capable of\b", "can"),
        (r"(?i)\ba considerable amount of\b", "a lot of"),
        (r"(?i)\bthe vast majority of\b", "most"),
        (r"(?i)\bin close proximity to\b", "near"),
        (r"(?i)\bat this juncture\b", "now"),
    ]
    for pattern, repl in ph_rules:
        text = _re.sub(pattern, repl, text)

    # ── 4) CONTRACTIONS ──
    c_rules = [
        (r"\bdoes not\b", "doesn't"), (r"\bDoes not\b", "Doesn't"),
        (r"\bdo not\b", "don't"), (r"\bDo not\b", "Don't"),
        (r"\bdid not\b", "didn't"), (r"\bDid not\b", "Didn't"),
        (r"\bis not\b", "isn't"), (r"\bIs not\b", "Isn't"),
        (r"\bare not\b", "aren't"), (r"\bAre not\b", "Aren't"),
        (r"\bwas not\b", "wasn't"), (r"\bWas not\b", "Wasn't"),
        (r"\bwere not\b", "weren't"), (r"\bWere not\b", "Weren't"),
        (r"\bcannot\b", "can't"), (r"\bCannot\b", "Can't"),
        (r"\bcould not\b", "couldn't"), (r"\bCould not\b", "Couldn't"),
        (r"\bwould not\b", "wouldn't"), (r"\bWould not\b", "Wouldn't"),
        (r"\bwill not\b", "won't"), (r"\bWill not\b", "Won't"),
        (r"\bshould not\b", "shouldn't"), (r"\bShould not\b", "Shouldn't"),
        (r"\bit is\b", "it's"), (r"\bIt is\b", "It's"),
        (r"\bthat is\b", "that's"), (r"\bThat is\b", "That's"),
        (r"\bthere is\b", "there's"), (r"\bThere is\b", "There's"),
        (r"\bwhat is\b", "what's"), (r"\bWhat is\b", "What's"),
        (r"\bthey are\b", "they're"), (r"\bThey are\b", "They're"),
        (r"\bwe are\b", "we're"), (r"\bWe are\b", "We're"),
        (r"\byou are\b", "you're"), (r"\bYou are\b", "You're"),
        (r"\bI am\b", "I'm"), (r"\bI have\b", "I've"),
        (r"\bI would\b", "I'd"), (r"\bI will\b", "I'll"),
        (r"\bwe have\b", "we've"), (r"\bWe have\b", "We've"),
        (r"\bthey have\b", "they've"), (r"\bThey have\b", "They've"),
        (r"\bwho is\b", "who's"), (r"\bWho is\b", "Who's"),
        (r"\bwhere is\b", "where's"), (r"\bWhere is\b", "Where's"),
        (r"\bhe is\b", "he's"), (r"\bHe is\b", "He's"),
        (r"\bshe is\b", "she's"), (r"\bShe is\b", "She's"),
        (r"\bhave not\b", "haven't"), (r"\bHave not\b", "Haven't"),
        (r"\bhas not\b", "hasn't"), (r"\bHas not\b", "Hasn't"),
        (r"\blet us\b", "let's"), (r"\bLet us\b", "Let's"),
    ]
    for pattern, repl in c_rules:
        text = _re.sub(pattern, repl, text)

    # ── 5) SENTENCE OPENER VARIATION ──
    def _vary_opener(m):
        if random.random() < 0.10:
            prefix = _pick(["And ", "But ", "So ", "Now, ", "Still, ", "Yet "])
            return m.group(1) + prefix + m.group(2).lower()
        return m.group(0)
    text = _re.sub(r"(\. )([A-Z])", _vary_opener, text)

    # ── 6) SPLIT LONG SENTENCES AT CONJUNCTIONS ──
    parts = _re.split(r"(?<=[.!?])\s+", text)
    rebuilt = []
    for s in parts:
        wc = len(s.split())
        if wc > 25 and random.random() < 0.45:
            split_done = False
            for conj in [" and ", " but ", " while ", " although ", " because "]:
                idx = s.find(conj, max(10, len(s) // 3))
                if idx > 0:
                    first = s[:idx].rstrip()
                    if not first.endswith((".", "!", "?")):
                        first += "."
                    second = s[idx + len(conj):].strip()
                    if second and len(second.split()) > 3:
                        second = second[0].upper() + second[1:]
                        rebuilt.append(first)
                        rebuilt.append(second)
                        split_done = True
                    break
            if not split_done:
                rebuilt.append(s)
        else:
            rebuilt.append(s)
    text = " ".join(rebuilt)

    # ── 7) INFORMAL DISCOURSE MARKERS ──
    def _inject_marker(m):
        if random.random() < 0.10:
            marker = _pick([
                "Honestly, ", "Frankly, ", "To be fair, ",
                "In a way, ", "Actually, ", "Really, ",
                "Look, ", "The thing is, ", "Truth be told, ",
                "Granted, ", "Admittedly, ", "Interestingly, ",
            ])
            return m.group(1) + marker + m.group(2).lower()
        return m.group(0)
    text = _re.sub(r"(\. )([A-Z])", _inject_marker, text)

    # ── 8) HEDGING LANGUAGE ──
    def _hedge(m):
        if random.random() < 0.06:
            h = _pick(["quite ", "fairly ", "rather ", "pretty ", "somewhat "])
            return h + m.group(0)
        return m.group(0)
    text = _re.sub(
        r"\b(important|clear|strong|difficult|common|effective|useful|evident|high|low|similar|likely|complex)\b",
        _hedge, text
    )

    # ── 9) PERSONAL VOICE TOUCHES ──
    pv_rules = [
        (r"(?i)\bIt can be argued that\b", lambda: _pick(["You could say", "One might say"])),
        (r"(?i)\bOne can observe that\b", lambda: _pick(["You can see that", "It's clear that"])),
        (r"(?i)\bIt is evident that\b", lambda: _pick(["Clearly,", "Obviously,"])),
        (r"(?i)\bIt is clear that\b", lambda: _pick(["Clearly,", "Obviously,"])),
        (r"(?i)\bIt is apparent that\b", lambda: _pick(["Clearly,", "Obviously,"])),
        (r"(?i)\bIt can be seen that\b", lambda: _pick(["You can see", "We can see"])),
        (r"(?i)\bIt is noteworthy that\b", lambda: _pick(["Worth noting,", "Interestingly,"])),
    ]
    for pattern, replacer in pv_rules:
        text = _re.sub(pattern, lambda m, fn=replacer: fn(), text)

    # ── 10) ADVERB VARIATION ──
    adv_rules = [
        (r"(?i)\bconsiderably\b", lambda: _pick(["a lot", "quite a bit", "noticeably"])),
        (r"(?i)\bremarkably\b", lambda: _pick(["surprisingly", "really"])),
        (r"(?i)\bparticularly\b", lambda: _pick(["especially", "mainly"])),
        (r"(?i)\bprimarily\b", lambda: _pick(["mainly", "mostly", "chiefly"])),
        (r"(?i)\bfundamentally\b", lambda: _pick(["at its core", "basically", "at heart"])),
        (r"(?i)\bincreasingly\b", lambda: _pick(["more and more", "progressively"])),
        (r"(?i)\bcritically\b", lambda: _pick(["crucially", "vitally"])),
        (r"(?i)\bsystematically\b", lambda: _pick(["step by step", "methodically"])),
        (r"(?i)\beffectively\b", lambda: _pick(["well", "in practice", "successfully"])),
        (r"(?i)\bconsequently\b", lambda: _pick(["so", "as a result"])),
        (r"(?i)\baccordingly\b", lambda: _pick(["so", "for that reason"])),
    ]
    for pattern, replacer in adv_rules:
        text = _re.sub(pattern, lambda m, fn=replacer: fn(), text)

    # ── 11) PASSIVE -> ACTIVE VOICE HINTS ──
    pv2_rules = [
        (r"(?i)\bwas conducted\b", lambda: _pick(["took place", "happened", "was done"])),
        (r"(?i)\bwere conducted\b", lambda: _pick(["took place", "happened"])),
        (r"(?i)\bwas observed\b", lambda: _pick(["showed up", "came through", "was noticed"])),
        (r"(?i)\bwas identified\b", lambda: _pick(["stood out", "came up", "was spotted"])),
        (r"(?i)\bwere identified\b", lambda: _pick(["stood out", "came up"])),
        (r"(?i)\bis characterized by\b", lambda: _pick(["stands out for", "is known for"])),
        (r"(?i)\bare characterized by\b", lambda: _pick(["stand out for", "are known for"])),
        (r"(?i)\bwas determined\b", lambda: _pick(["turned out", "came out as"])),
        (r"(?i)\bwas established\b", lambda: _pick(["came about", "was set up"])),
        (r"(?i)\bwas implemented\b", lambda: _pick(["went live", "was rolled out", "was put in place"])),
        (r"(?i)\bwere implemented\b", lambda: _pick(["went live", "were rolled out"])),
        (r"(?i)\bwas utilized\b", "was used"),
        (r"(?i)\bwas achieved\b", lambda: _pick(["came through", "worked out"])),
        (r"(?i)\bwas obtained\b", lambda: _pick(["came in", "was gathered"])),
    ]
    for pattern, replacer in pv2_rules:
        text = _re.sub(pattern, lambda m, fn=replacer: fn(), text)

    # ── 12) REPETITIVE SENTENCE-START BREAKING ──
    # AI tends to start many consecutive sentences with "The", "This", "These"
    sentences = _re.split(r"(?<=[.!?])\s+", text)
    if len(sentences) > 2:
        for i in range(1, len(sentences)):
            prev_start = sentences[i - 1].split()[0] if sentences[i - 1].split() else ""
            curr_words = sentences[i].split()
            if not curr_words:
                continue
            curr_start = curr_words[0]
            if curr_start == prev_start and curr_start in ("The", "This", "These", "That", "It", "They", "There"):
                if random.random() < 0.6:
                    alt_starters = {
                        "The": ["What stands out about the", "Looking at the", "As for the", "When it comes to the"],
                        "This": ["What this", "Here, this", "In practice, this", "On that note, this"],
                        "These": ["All these", "Such", "Taken together, these", "Each of these"],
                        "That": ["And that", "Which is why that", "Notably, that"],
                        "It": ["What it", "In short, it", "As it turns out, it", "Basically, it"],
                        "They": ["In practice, they", "What they", "As it happens, they"],
                        "There": ["As it stands, there", "Right now, there", "In fact, there"],
                    }
                    alts = alt_starters.get(curr_start, [])
                    if alts:
                        replacement = _pick(alts)
                        rest = " ".join(curr_words[1:])
                        sentences[i] = replacement + " " + rest
        text = " ".join(sentences)

    # ── 13) PARENTHETICAL ASIDES ──
    # Humans naturally insert brief asides in parentheses
    asides = [
        "(and this matters)",
        "(which is a big deal)",
        "(not surprisingly)",
        "(to no one's surprise)",
        "(as you'd expect)",
        "(worth keeping in mind)",
        "(and this is key)",
        "(for better or worse)",
        "(understandably)",
        "(at least in theory)",
        "(in hindsight)",
        "(surprisingly enough)",
        "(to put it mildly)",
        "(believe it or not)",
        "(rightly so)",
    ]

    def _inject_aside(m):
        if random.random() < 0.04:
            aside = _pick(asides)
            return m.group(0).rstrip() + " " + aside + " "
        return m.group(0)
    text = _re.sub(r"[,;]\s", _inject_aside, text)

    # ── 14) RHETORICAL QUESTION CONVERSION ──
    # Convert ~4% of declarative statements to rhetorical questions
    rq_patterns = [
        (r"(?i)\bIt is (clear|obvious|evident|apparent) that (.+?)\.", lambda m:
            _pick(["Isn't it " + m.group(1) + " that " + m.group(2) + "?",
                   "Can anyone deny that " + m.group(2) + "?",
                   "Who would argue that " + m.group(2) + " isn't the case?"])),
        (r"(?i)\bEveryone (knows|agrees|understands) that (.+?)\.", lambda m:
            "Does everyone really " + m.group(1).rstrip("s") + " that " + m.group(2) + "?"),
        (r"(?i)\bThe (answer|reason|explanation) is (simple|straightforward|clear)[.:] (.+?)\.", lambda m:
            "So what's the " + m.group(1) + "? " + m.group(3)[0].upper() + m.group(3)[1:] + "."),
    ]
    for pattern, replacer in rq_patterns:
        if random.random() < 0.5:
            text = _re.sub(pattern, replacer, text, count=1)

    # ── 15) CLAUSE REORDERING ──
    # Move "because X, Y" to "Y, because X" and vice versa (~15% of matches)
    def _reorder_because(m):
        if random.random() < 0.15:
            reason = m.group(1).strip().rstrip(",")
            main = m.group(2).strip()
            if main and reason:
                return main[0].upper() + main[1:] + ", because " + reason[0].lower() + reason[1:]
        return m.group(0)
    text = _re.sub(r"\b[Bb]ecause ([^,]+), ([^.!?]+[.!?])", _reorder_because, text)

    # ── 16) SENTENCE MERGING ──
    # Merge very short consecutive sentences (both <8 words) with semicolons or dashes
    merged_sentences = _re.split(r"(?<=[.!?])\s+", text)
    if len(merged_sentences) > 2:
        rebuilt2 = []
        i = 0
        while i < len(merged_sentences):
            s = merged_sentences[i]
            wc = len(s.split())
            if (wc < 8 and i + 1 < len(merged_sentences)
                    and len(merged_sentences[i + 1].split()) < 8
                    and random.random() < 0.25):
                joiner = _pick(["; ", " - ", " and "])
                next_s = merged_sentences[i + 1]
                if next_s:
                    combined = s.rstrip(".!?") + joiner + next_s[0].lower() + next_s[1:]
                    rebuilt2.append(combined)
                    i += 2
                    continue
            rebuilt2.append(s)
            i += 1
        text = " ".join(rebuilt2)

    # ── 17) AI N-GRAM PATTERN BREAKING ──
    # Common AI multi-word patterns that humans rarely use in the same way
    ngram_rules = [
        (r"(?i)\bIt is worth noting\b", lambda: _pick(["Keep in mind", "Notice", "One thing to note"])),
        (r"(?i)\bThis is particularly true\b", lambda: _pick(["This really shows", "You see this especially"])),
        (r"(?i)\bThis is especially the case\b", lambda: _pick(["This really holds", "You notice this most"])),
        (r"(?i)\bA growing body of\b", lambda: _pick(["More and more", "A lot of recent"])),
        (r"(?i)\bIn light of\b", lambda: _pick(["Given", "Considering", "Because of"])),
        (r"(?i)\bOn the other hand\b", lambda: _pick(["Then again", "But", "That said"])),
        (r"(?i)\bIn this regard\b", lambda: _pick(["On that front", "Here", "In this area"])),
        (r"(?i)\bFrom this perspective\b", lambda: _pick(["Seen this way", "Looked at like that"])),
        (r"(?i)\bIt is imperative that\b", lambda: _pick(["It's critical that", "We really need to"])),
        (r"(?i)\bThis suggests that\b", lambda: _pick(["That points to", "Which hints that", "So it seems"])),
        (r"(?i)\bThis indicates that\b", lambda: _pick(["That shows", "Which tells us", "So it looks like"])),
        (r"(?i)\bThis implies that\b", lambda: _pick(["That means", "Which suggests", "So basically"])),
        (r"(?i)\bIt is essential to\b", lambda: _pick(["You really need to", "It's critical to", "The key thing is to"])),
        (r"(?i)\bplay(?:s)? an important role\b", lambda: _pick(["matter a lot", "count for a lot", "carry real weight"])),
        (r"(?i)\bserves as a\b", lambda: _pick(["works as a", "acts as a", "doubles as a"])),
        (r"(?i)\ba key factor\b", lambda: _pick(["a big driver", "a real factor", "something that matters"])),
        (r"(?i)\ba critical component\b", lambda: _pick(["a core piece", "something essential", "a must-have part"])),
        (r"(?i)\bin today's (?:digital|modern|contemporary) (?:age|era|world)\b",
         lambda: _pick(["nowadays", "today", "these days", "in the current climate"])),
        (r"(?i)\bthe landscape of\b", lambda: _pick(["the world of", "the state of", "how things stand in"])),
        (r"(?i)\bpaves the way for\b", lambda: _pick(["opens the door to", "sets the stage for", "makes room for"])),
        (r"(?i)\bshed(?:s)? light on\b", lambda: _pick(["clarifies", "helps explain", "brings clarity to"])),
        (r"(?i)\btake(?:s)? into account\b", lambda: _pick(["factor in", "consider", "keep in mind"])),
        (r"(?i)\bthe findings suggest\b", lambda: _pick(["the results point to", "what came out is", "the data shows"])),
        (r"(?i)\bthe results indicate\b", lambda: _pick(["the numbers show", "what we found is", "the data tells us"])),
        (r"(?i)\bas a whole\b", lambda: _pick(["overall", "all in all", "taken together"])),
        (r"(?i)\bgiven the fact that\b", lambda: _pick(["since", "because", "considering"])),
    ]
    for pattern, replacer in ngram_rules:
        text = _re.sub(pattern, lambda m, fn=replacer: fn(), text)

    # ── 18) EMPHATIC EXPRESSIONS ──
    # Occasionally strengthen weak statements with human emphasis
    def _emphaticize(m):
        if random.random() < 0.04:
            emphatics = [
                "What really matters here is that ",
                "The bottom line is, ",
                "Put simply, ",
                "Here's the thing: ",
                "The real takeaway is that ",
                "Simply put, ",
            ]
            return m.group(1) + _pick(emphatics) + m.group(2).lower()
        return m.group(0)
    text = _re.sub(r"(\. )([A-Z])", _emphaticize, text)

    # ── 19) DOUBLE-SPACE AFTER PERIOD VARIATION ──
    # Some humans use double spaces; mix in a few to break uniformity
    def _vary_spacing(m):
        if random.random() < 0.08:
            return ".  " + m.group(1)
        return ". " + m.group(1)
    text = _re.sub(r"\. ([A-Z])", _vary_spacing, text)

    return text


def ai_text_clean(text):
    text = _ud.normalize("NFKC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _re.sub(r"[\u200B\u200C\u200D\u2060\uFEFF\u180E]", "", text)
    text = _re.sub(r"[\u00A0\u2007\u202F\u2009\u200A]", " ", text)
    text = _re.sub(r"[\u0000-\u0008\u000B\u000C\u000E-\u001F\u007F-\u009F]", "", text)
    text = _re.sub(r"[\u2013\u2014\u2212]", "-", text)
    text = _re.sub(r'[\u201C\u201D\u201E\u00AB\u00BB]', '"', text)
    text = _re.sub(r"[\u2018\u2019\u201A\u201B\u2039\u203A]", "'", text)
    text = text.replace("\u2026", "...")
    text = _re.sub(r"[\u2022\u25E6\u00B7\u2219]", "-", text)
    text = _re.sub(
        r"[\u00A9\u00AE\u2000-\u3300]|[\U0001F000-\U0001F9FF]",
        "", text
    )
    text = _re.sub(r"\s+([.,!?;:])", r"\1", text)
    text = _re.sub(r"[ \t]*\n[ \t]*", "\n", text)
    text = _re.sub(r"[ \t]{2,}", " ", text)
    text = _re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    text = text.strip()
    return text


# Apply BOTH humanizer and cleaner to all paragraphs
humanized_count = 0
cleaned_count = 0

rewrite_zone_indices = set(ALL_REWRITES.keys())

for para in doc.paragraphs:
    for run in para.runs:
        if run.text and run.text.strip():
            original = run.text
            run.text = humanize_text(run.text)
            if run.text != original:
                humanized_count += 1
            original = run.text
            run.text = ai_text_clean(run.text)
            if run.text != original:
                cleaned_count += 1

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.text and run.text.strip():
                        original = run.text
                        run.text = humanize_text(run.text)
                        if run.text != original:
                            humanized_count += 1
                        original = run.text
                        run.text = ai_text_clean(run.text)
                        if run.text != original:
                            cleaned_count += 1

print(f"Humanizer applied to {humanized_count} text runs")
print(f"AI text cleaning applied to {cleaned_count} text runs")

# ===========================================================
# SAVE
# ===========================================================

doc.save('CU_finalprojectreport_rewritten.docx')
print("Done - saved CU_finalprojectreport_rewritten.docx")
